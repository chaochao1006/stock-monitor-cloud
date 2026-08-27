"""
美股 BOLL 布林线监控脚本

运行方式：
1. 建议依赖：
   pip install pandas numpy yfinance requests python-docx pillow pandas-market-calendars

2. 在 Spyder 中直接运行本文件，或用 Windows 任务计划程序每天美股收盘后运行。

3. 输出目录：
   C:\\Users\\81975\\Desktop\\每日报告\\BOLL

说明：
- 默认使用 Yahoo Chart API 获取最近 1 年日线数据。
- 如需优先尝试 yfinance，可把 USE_YFINANCE_FIRST 改为 True。
- 图表只嵌入 Word 报告，临时图片会在生成报告后删除。
- 本脚本仅做技术指标监控，不构成投资建议。
"""

from __future__ import annotations

import contextlib
import math
import os
import re
import shutil
import sys
import tempfile
import time
import traceback
from dataclasses import dataclass, field
from datetime import date, datetime, time as dtime, timedelta
from io import StringIO
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from zoneinfo import ZoneInfo

os.environ.setdefault("PANDAS_USE_NUMEXPR", "0")
os.environ.setdefault("PANDAS_USE_BOTTLENECK", "0")

with contextlib.redirect_stderr(StringIO()):
    import numpy as np
    import pandas as pd

import requests
from PIL import Image, ImageDraw, ImageFont

try:
    import yfinance as yf
except Exception:
    yf = None

try:
    import pandas_market_calendars as mcal
except Exception:
    mcal = None

try:
    from docx import Document
    from docx.shared import Inches, Pt
except Exception:
    Document = None
    Inches = None
    Pt = None

try:
    from openpyxl import Workbook
    from openpyxl.chart import LineChart, Reference
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter
except Exception:
    Workbook = None
    LineChart = None
    Reference = None
    Alignment = None
    Font = None
    PatternFill = None
    get_column_letter = None


# =========================
# 用户配置
# =========================

STOCK_POOL = [
    "NOK", "XE", "RKLB", "NBIS", "ARM", "OPTX", "GLW", "COHR", "LITE",
    "MRVL", "INTC", "ALAB", "FLNC", "CRWV", "ORCL", "NVTS", "POET",
    "DELL", "BB", "ONDS", "SMR", "OKLO", "IONQ", "QBTS", "NVDA", "AMZN",
    "AMD", "CRCL", "FLY", "CBRS", "CIEN", "SIVEF", "SHAZ", "LUNR", "ASTS",
    "AAOI", "INOD","SPCX","RDW","QCOM","COIN","CORZ","IREN","VELO","MSFT","GOOG","AVGO",
    "AMAT","AMKR","LRCX","SNDK","MU","VICR","NNE","CCJ","BWAY","MANE",
    "QSI","CRSP","IBRX","INSP",
]

REPORT_BASE_DIR = Path(os.environ.get("REPORT_BASE_DIR", Path(__file__).resolve().parents[1] / "data"))
OUTPUT_DIR = REPORT_BASE_DIR / "BOLL"
EXCEL_TRACKING_DIR = OUTPUT_DIR / "EXCLE"

BOLL_PERIOD = 20
BOLL_STD_MULTIPLIER = 2
BOLL_LOW_LOOKBACK_DAYS = 168
BOLL_SQUEEZE_MIN_DAYS = 6
BOLL_WEAK_SQUEEZE_MIN_DAYS = 8
TRACKING_DAYS = 30
LOOKBACK_PERIOD = "1y"
REQUEST_TIMEOUT = 15
YFINANCE_TIMEOUT = 4
SLEEP_BETWEEN_SYMBOLS = 0.2
USE_YFINANCE_FIRST = False


# =========================
# 数据结构
# =========================

@dataclass
class BollResult:
    symbol: str
    company_name: str = ""
    data_source: str = ""
    data_error: str = ""
    df: Optional[pd.DataFrame] = None
    latest_date: Optional[pd.Timestamp] = None
    close: Optional[float] = None
    latest_volume: Optional[float] = None
    avg_volume20: Optional[float] = None
    volume_ratio: Optional[float] = None
    bandwidth: Optional[float] = None
    bandwidth_percentile: Optional[float] = None
    percent_b: Optional[float] = None
    price_position: str = "无法判断"
    signals: List[str] = field(default_factory=list)
    signal_grade: str = "无明显信号"
    conclusion: str = ""
    risk_notes: List[str] = field(default_factory=list)


# =========================
# 基础工具
# =========================

def eastern_now() -> datetime:
    return datetime.now(ZoneInfo("America/New_York"))


def fmt_date(value) -> str:
    if value is None or pd.isna(value):
        return "N/A"
    if isinstance(value, pd.Timestamp):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    return str(value)


def safe_float(value) -> Optional[float]:
    try:
        if value is None or pd.isna(value):
            return None
        value = float(value)
        if not math.isfinite(value):
            return None
        return value
    except Exception:
        return None


def fmt_num(value, digits: int = 4) -> str:
    value = safe_float(value)
    if value is None:
        return "N/A"
    return f"{value:.{digits}f}"


def fmt_pct(value, digits: int = 2) -> str:
    value = safe_float(value)
    if value is None:
        return "N/A"
    return f"{value * 100:.{digits}f}%"


def fmt_price(value) -> str:
    value = safe_float(value)
    if value is None:
        return "N/A"
    return f"${value:.2f}"


def clean_text(text: str) -> str:
    return " ".join(str(text or "").split())


def get_company_name(symbol: str) -> str:
    fallback = {
        "NOK": "Nokia Oyj",
        "XE": "X-energy",
        "RKLB": "Rocket Lab USA",
        "NBIS": "Nebius Group",
        "ARM": "Arm Holdings",
        "OPTX": "Syntec Optics Holdings",
        "GLW": "Corning",
        "COHR": "Coherent",
        "LITE": "Lumentum Holdings",
        "MRVL": "Marvell Technology",
        "INTC": "Intel",
        "ALAB": "Astera Labs",
        "FLNC": "Fluence Energy",
        "CRWV": "CoreWeave",
        "ORCL": "Oracle",
        "NVTS": "Navitas Semiconductor",
        "POET": "POET Technologies",
        "DELL": "Dell Technologies",
        "BB": "BlackBerry",
        "ONDS": "Ondas Holdings",
        "SMR": "NuScale Power",
        "OKLO": "Oklo",
        "IONQ": "IonQ",
        "QBTS": "D-Wave Quantum",
        "NVDA": "NVIDIA",
        "AMZN": "Amazon.com",
        "AMD": "Advanced Micro Devices",
        "CRCL": "Circle Internet Group",
        "FLY": "Firefly Aerospace",
        "CBRS": "Chain Bridge Bancorp",
        "CIEN": "Ciena",
        "SIVEF": "Sivers Semiconductors",
        "SHAZ": "Shaz",
        "LUNR": "Intuitive Machines",
        "ASTS": "AST SpaceMobile",
        "AAOI": "Applied Optoelectronics",
        "INOD": "Innodata",
    }
    if symbol in fallback:
        return fallback[symbol]
    if yf is not None:
        try:
            info = yf.Ticker(symbol).get_info()
            name = info.get("shortName") or info.get("longName")
            if name:
                return clean_text(name)
        except Exception:
            pass
    return symbol


# =========================
# 交易日判断
# =========================

def get_market_info() -> Dict[str, object]:
    now_et = eastern_now()
    today_et = now_et.date()
    after_close = now_et.time() >= dtime(16, 10)
    info = {
        "now_et": now_et,
        "today_et": today_et,
        "after_close": after_close,
        "is_trading_day": None,
        "next_trading_day": None,
        "calendar_source": "weekday fallback",
    }
    if mcal is not None:
        try:
            nyse = mcal.get_calendar("NYSE")
            schedule = nyse.schedule(
                start_date=pd.Timestamp(today_et - timedelta(days=7)),
                end_date=pd.Timestamp(today_et + timedelta(days=14)),
            )
            trading_dates = [d.date() for d in schedule.index]
            info["is_trading_day"] = today_et in trading_dates
            future = [d for d in trading_dates if d > today_et]
            info["next_trading_day"] = future[0] if future else None
            info["calendar_source"] = "pandas_market_calendars:NYSE"
            return info
        except Exception:
            pass

    info["is_trading_day"] = today_et.weekday() < 5
    next_day = today_et + timedelta(days=1)
    while next_day.weekday() >= 5:
        next_day += timedelta(days=1)
    info["next_trading_day"] = next_day
    return info


# =========================
# 数据获取
# =========================

def ensure_ohlcv(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()
    df = df.copy()
    if isinstance(df.columns, pd.MultiIndex):
        df.columns = [c[0] if isinstance(c, tuple) else c for c in df.columns]

    rename_map = {}
    for col in df.columns:
        key = str(col).strip().lower()
        if key in {"open", "high", "low", "close", "volume"}:
            rename_map[col] = key.capitalize()
        elif key in {"adj close", "adjusted close"}:
            rename_map[col] = "Adjusted Close"
    df = df.rename(columns=rename_map)

    required = ["Open", "High", "Low", "Close", "Volume"]
    if any(col not in df.columns for col in required):
        return pd.DataFrame()
    if "Adjusted Close" not in df.columns:
        df["Adjusted Close"] = np.nan

    df = df[["Open", "High", "Low", "Close", "Adjusted Close", "Volume"]].copy()
    df.index = pd.to_datetime(df.index)
    if getattr(df.index, "tz", None) is not None:
        df.index = df.index.tz_convert("America/New_York").tz_localize(None)
    df.index = df.index.normalize()
    df = df.sort_index()
    df = df[~df.index.duplicated(keep="last")]
    for col in ["Open", "High", "Low", "Close", "Adjusted Close", "Volume"]:
        df[col] = pd.to_numeric(df[col], errors="coerce")
    df = df.dropna(subset=["Open", "High", "Low", "Close"])
    df = df[df["Close"] > 0]
    return df


def fetch_yfinance(symbol: str) -> Tuple[pd.DataFrame, str]:
    if yf is None:
        return pd.DataFrame(), "未安装或无法导入 yfinance"
    try:
        with contextlib.redirect_stdout(StringIO()), contextlib.redirect_stderr(StringIO()):
            df = yf.download(
                symbol,
                period=LOOKBACK_PERIOD,
                interval="1d",
                auto_adjust=False,
                progress=False,
                threads=False,
                timeout=YFINANCE_TIMEOUT,
            )
        df = ensure_ohlcv(df)
        if df.empty:
            return df, "yfinance 返回空数据或缺少 OHLCV 字段"
        return df, ""
    except Exception as exc:
        return pd.DataFrame(), f"yfinance 异常：{exc}"


def fetch_yahoo_chart(symbol: str) -> Tuple[pd.DataFrame, str]:
    url = f"https://query1.finance.yahoo.com/v8/finance/chart/{symbol}"
    params = {
        "range": LOOKBACK_PERIOD,
        "interval": "1d",
        "events": "history",
        "includeAdjustedClose": "true",
    }
    try:
        response = requests.get(
            url,
            params=params,
            headers={"User-Agent": "Mozilla/5.0 (compatible; boll-monitor/1.0)"},
            timeout=REQUEST_TIMEOUT,
        )
        response.raise_for_status()
        payload = response.json()
        chart = payload.get("chart", {})
        if chart.get("error"):
            return pd.DataFrame(), str(chart["error"])
        result = chart.get("result") or []
        if not result:
            return pd.DataFrame(), "Yahoo Chart API 未返回 result"
        block = result[0]
        timestamps = block.get("timestamp") or []
        quote = (block.get("indicators", {}).get("quote") or [{}])[0]
        adjclose = (block.get("indicators", {}).get("adjclose") or [{}])[0].get("adjclose")
        if not timestamps or not quote:
            return pd.DataFrame(), "Yahoo Chart API 返回数据不完整"
        df = pd.DataFrame({
            "Date": pd.to_datetime(timestamps, unit="s", utc=True)
                .tz_convert("America/New_York")
                .tz_localize(None)
                .normalize(),
            "Open": quote.get("open"),
            "High": quote.get("high"),
            "Low": quote.get("low"),
            "Close": quote.get("close"),
            "Adjusted Close": adjclose if adjclose is not None else quote.get("close"),
            "Volume": quote.get("volume"),
        }).set_index("Date")
        df = ensure_ohlcv(df)
        if df.empty:
            return df, "Yahoo Chart API 数据清洗后为空"
        return df, ""
    except Exception as exc:
        return pd.DataFrame(), f"Yahoo Chart API 异常：{exc}"


def fetch_ohlcv(symbol: str) -> Tuple[pd.DataFrame, str, str]:
    errors = []
    sources = [("Yahoo Chart API", fetch_yahoo_chart), ("yfinance", fetch_yfinance)]
    if USE_YFINANCE_FIRST:
        sources = [("yfinance", fetch_yfinance), ("Yahoo Chart API", fetch_yahoo_chart)]
    for source, func in sources:
        df, err = func(symbol)
        if not df.empty:
            return df, source, ""
        errors.append(f"{source}: {err}")
    return pd.DataFrame(), "", "；".join(errors)


# =========================
# BOLL 计算与信号
# =========================

def add_bollinger(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    middle = df["Close"].rolling(BOLL_PERIOD, min_periods=BOLL_PERIOD).mean()
    std = df["Close"].rolling(BOLL_PERIOD, min_periods=BOLL_PERIOD).std(ddof=0)
    upper = middle + BOLL_STD_MULTIPLIER * std
    lower = middle - BOLL_STD_MULTIPLIER * std
    width = upper - lower
    df["Middle Band"] = middle
    df["Upper Band"] = upper
    df["Lower Band"] = lower
    df["Bandwidth"] = width / middle.replace(0, np.nan)
    df["Percent B"] = (df["Close"] - lower) / width.replace(0, np.nan)
    df["Volume MA20"] = df["Volume"].rolling(20, min_periods=5).mean()
    df["Daily Return"] = df["Close"].pct_change()
    return df


def latest_data_date(results: List[BollResult]) -> Optional[pd.Timestamp]:
    dates = [r.latest_date for r in results if r.latest_date is not None]
    if not dates:
        return None
    return max(pd.Timestamp(d).normalize() for d in dates)


def is_bandwidth_declining(df: pd.DataFrame, days: int = 7) -> bool:
    bw = df["Bandwidth"].dropna().tail(days)
    if len(bw) < max(5, days - 1):
        return False
    diffs = bw.diff().dropna()
    return bool((bw.iloc[-1] < bw.iloc[0]) and ((diffs < 0).sum() >= max(3, len(diffs) - 1)))


def is_bandwidth_declining_strict(df: pd.DataFrame, days: int = 10) -> bool:
    bw = df["Bandwidth"].dropna().tail(days)
    if len(bw) < days:
        return False
    diffs = bw.diff().dropna()
    return bool((diffs < 0).all())


def is_bandwidth_expanding(df: pd.DataFrame, days: int = 5) -> bool:
    bw = df["Bandwidth"].dropna().tail(days)
    if len(bw) < days:
        return False
    diffs = bw.diff().dropna()
    expansion_ratio = bw.iloc[-1] / bw.iloc[0] if bw.iloc[0] else np.nan
    required_ratio = 1.04 if days <= 3 else 1.08
    return bool((bw.iloc[-1] > bw.iloc[0]) and ((diffs > 0).sum() >= max(1, len(diffs) - 1)) and expansion_ratio >= required_ratio)


def count_squeeze_days(df: pd.DataFrame, percentile_threshold: float = 0.20) -> int:
    """统计截至最新交易日，Bandwidth 连续处于历史低分位的天数。"""
    bw = df["Bandwidth"].dropna()
    if len(bw) < 60:
        return 0
    ranks = bw.rank(pct=True)
    flags = ranks <= percentile_threshold
    count = 0
    for value in flags.iloc[::-1]:
        if bool(value):
            count += 1
        else:
            break
    return count


def count_recent_expansion_days(df: pd.DataFrame) -> int:
    """统计最近连续开口天数；只看最新一段，最多用于 5 天内启动判断。"""
    bw = df["Bandwidth"].dropna()
    if len(bw) < 10:
        return 0
    diffs = bw.diff()
    count = 0
    for value in diffs.iloc[::-1]:
        if pd.notna(value) and value > 0:
            count += 1
        else:
            break
    return count


def had_squeeze_before_recent_expansion(df: pd.DataFrame, expansion_days: int, percentile_threshold: float = 0.20) -> bool:
    if expansion_days <= 0:
        return False
    bw = df["Bandwidth"].dropna()
    if len(bw) < 60:
        return False
    ranks = bw.rank(pct=True)
    pre_expansion = ranks.iloc[: -expansion_days]
    if len(pre_expansion) < 5:
        return False
    return bool((pre_expansion.tail(10) <= percentile_threshold).sum() >= 5)


def is_upward_boll_breakout_after_squeeze(df: pd.DataFrame, expansion_days: int) -> bool:
    if expansion_days <= 0 or expansion_days > 5:
        return False
    latest = df.iloc[-1]
    recent = df.tail(expansion_days + 1)
    price_up = latest["Close"] > recent["Close"].iloc[0]
    upper_up = latest["Upper Band"] > recent["Upper Band"].iloc[0]
    lower_down_or_flat = latest["Lower Band"] <= recent["Lower Band"].iloc[0] * 1.01
    above_middle = latest["Close"] > latest["Middle Band"]
    near_or_above_upper = latest["Percent B"] >= 0.85
    return bool(price_up and upper_up and lower_down_or_flat and above_middle and near_or_above_upper)


def is_low_price_area(df: pd.DataFrame) -> bool:
    closes = df["Close"].dropna()
    if len(closes) < 60:
        return False
    latest_close = closes.iloc[-1]
    close_percentile = float((closes <= latest_close).mean())
    one_year_low = closes.min()
    one_year_high = closes.max()
    if one_year_high <= one_year_low:
        return False
    low_zone_ratio = (latest_close - one_year_low) / (one_year_high - one_year_low)
    return bool(close_percentile <= 0.35 or low_zone_ratio <= 0.35)


def bandwidth_holds_low(df: pd.DataFrame, days: int = 5, percentile_threshold: float = 0.20) -> bool:
    bw = df["Bandwidth"].dropna()
    if len(bw) < 60:
        return False
    ranks = bw.rank(pct=True)
    recent = ranks.tail(days)
    return bool(len(recent) >= days and (recent <= percentile_threshold).all())


def recent_bandwidth_percentile(df: pd.DataFrame, lookback: int = 126) -> Optional[float]:
    bw = df["Bandwidth"].dropna().tail(lookback)
    if len(bw) < 60:
        return None
    latest_bw = bw.iloc[-1]
    return float((bw <= latest_bw).mean())


def count_recent_low_bandwidth_days(df: pd.DataFrame, percentile_threshold: float, lookback: int = 126) -> int:
    bw = df["Bandwidth"].dropna().tail(lookback)
    if len(bw) < 60:
        return 0
    ranks = bw.rank(pct=True)
    flags = ranks <= percentile_threshold
    count = 0
    for value in flags.iloc[::-1]:
        if bool(value):
            count += 1
        else:
            break
    return count


def is_bandwidth_100d_low(df: pd.DataFrame) -> bool:
    bw = df["Bandwidth"].dropna().tail(100)
    if len(bw) < 60:
        return False
    return bool(bw.iloc[-1] <= bw.min())


def had_recent_obvious_squeeze(df: pd.DataFrame, recent_days: int = 10, lookback: int = BOLL_LOW_LOOKBACK_DAYS) -> bool:
    bw = df["Bandwidth"].dropna().tail(lookback)
    if len(bw) < 80:
        return False
    ranks = bw.rank(pct=True)
    recent_ranks = ranks.tail(recent_days)
    recent_bw = bw.tail(recent_days)
    low_enough = (recent_ranks <= 0.20).sum() >= 3 or recent_ranks.min() <= 0.10
    recent_100d_low = len(recent_bw) >= 10 and recent_bw.min() <= bw.tail(100).min()
    return bool(low_enough or recent_100d_low)


def is_near_middle_narrow_range(df: pd.DataFrame, days: int = 5) -> bool:
    recent = df.dropna(subset=["Close", "Middle Band", "Upper Band", "Lower Band", "Percent B"]).tail(days)
    if len(recent) < days:
        return False
    in_middle_zone = recent["Percent B"].between(0.32, 0.68).mean() >= 0.6
    channel = (recent["Upper Band"] - recent["Lower Band"]).replace(0, np.nan)
    max_range_ratio = ((recent["High"] - recent["Low"]) / channel).dropna()
    narrow_range = not max_range_ratio.empty and max_range_ratio.mean() <= 0.45
    return bool(in_middle_zone and narrow_range)


def line_slope_positive(series: pd.Series, days: int = 3) -> bool:
    values = series.dropna().tail(days)
    if len(values) < days:
        return False
    return bool(values.iloc[-1] > values.iloc[0])


def line_slope_negative(series: pd.Series, days: int = 3) -> bool:
    values = series.dropna().tail(days)
    if len(values) < days:
        return False
    return bool(values.iloc[-1] < values.iloc[0])


def evaluate_boll(result: BollResult) -> BollResult:
    df = result.df
    if df is None or df.empty or len(df) < 60:
        result.data_error = result.data_error or "数据长度不足，无法稳定计算 BOLL"
        return result

    df = add_bollinger(df)
    df = df.dropna(subset=["Middle Band", "Upper Band", "Lower Band", "Bandwidth", "Percent B"])
    if df.empty or len(df) < 30:
        result.data_error = "BOLL 指标有效数据不足"
        return result

    result.df = df
    latest = df.iloc[-1]
    prev = df.iloc[-2]
    result.latest_date = df.index[-1]
    result.close = safe_float(latest["Close"])
    result.latest_volume = safe_float(latest["Volume"])
    result.avg_volume20 = safe_float(latest["Volume MA20"])
    if result.latest_volume is not None and result.avg_volume20 and result.avg_volume20 > 0:
        result.volume_ratio = result.latest_volume / result.avg_volume20
    result.bandwidth = safe_float(latest["Bandwidth"])
    result.percent_b = safe_float(latest["Percent B"])

    result.bandwidth_percentile = recent_bandwidth_percentile(df, BOLL_LOW_LOOKBACK_DAYS)

    close = latest["Close"]
    middle = latest["Middle Band"]
    upper = latest["Upper Band"]
    lower = latest["Lower Band"]
    percent_b = latest["Percent B"]
    return_today = safe_float(latest["Daily Return"])
    volume_ratio = result.volume_ratio or 0

    near_middle = abs(percent_b - 0.5) <= 0.18
    near_upper = 0.85 <= percent_b <= 1.0
    near_lower = 0.0 <= percent_b <= 0.15
    above_upper = close > upper and percent_b > 1
    below_lower = close < lower and percent_b < 0
    between_lower_middle = 0.0 <= percent_b < 0.50
    between_middle_upper = 0.50 <= percent_b < 1.0
    near_middle_from_below = 0.35 <= percent_b < 0.50
    crossed_middle_from_below = close > middle and prev["Close"] <= prev["Middle Band"]
    rebounded_from_lower = df["Percent B"].tail(6).min() <= 0.1 and close > middle

    if above_upper:
        result.price_position = "突破上轨"
    elif below_lower:
        result.price_position = "跌破下轨"
    elif near_middle:
        result.price_position = "中轨附近"
    elif near_upper:
        result.price_position = "靠近上轨"
    elif near_lower:
        result.price_position = "靠近下轨"
    elif percent_b > 0.5:
        result.price_position = "中轨上方"
    else:
        result.price_position = "中轨下方"

    squeeze_days_10 = count_recent_low_bandwidth_days(df, 0.10, BOLL_LOW_LOOKBACK_DAYS)
    squeeze_days_20 = count_recent_low_bandwidth_days(df, 0.20, BOLL_LOW_LOOKBACK_DAYS)
    squeeze_days_30 = count_recent_low_bandwidth_days(df, 0.30, BOLL_LOW_LOOKBACK_DAYS)
    low_price_area = is_low_price_area(df)
    bandwidth_100d_low = is_bandwidth_100d_low(df)
    low_bandwidth_hold_30 = squeeze_days_30 >= BOLL_WEAK_SQUEEZE_MIN_DAYS
    squeeze_narrowing = (
        is_bandwidth_declining_strict(df, 15)
        or is_bandwidth_declining_strict(df, 12)
        or is_bandwidth_declining_strict(df, 10)
        or is_bandwidth_declining_strict(df, 8)
    )
    middle_narrow_range = near_middle and is_near_middle_narrow_range(df, 5)
    strong_squeeze = ((squeeze_days_10 >= BOLL_SQUEEZE_MIN_DAYS) or bandwidth_100d_low) and middle_narrow_range
    medium_squeeze = squeeze_days_20 >= BOLL_SQUEEZE_MIN_DAYS and middle_narrow_range
    bandwidth_low_30 = result.bandwidth_percentile is not None and result.bandwidth_percentile <= 0.30
    weak_squeeze = bandwidth_low_30 and (squeeze_narrowing or low_bandwidth_hold_30) and middle_narrow_range
    recent_obvious_squeeze = had_recent_obvious_squeeze(df)

    expanding = is_bandwidth_expanding(df, 5) or is_bandwidth_expanding(df, 3) or is_bandwidth_expanding(df, 2)
    expansion_days = count_recent_expansion_days(df)
    upper_up = line_slope_positive(df["Upper Band"], 3)
    lower_down = line_slope_negative(df["Lower Band"], 3)
    volume_big = volume_ratio >= 1.3
    volume_huge = volume_ratio >= 1.8
    volume_mild = volume_ratio >= 1.1
    two_day_above_upper = bool((df["Close"].tail(2) > df["Upper Band"].tail(2)).all())
    strong_return = return_today is not None and return_today >= 0.03
    mild_return = return_today is not None and return_today >= 0.015
    middle_rebound = close > middle and percent_b >= 0.75 and volume_big and (
        crossed_middle_from_below or df["Percent B"].tail(6).min() <= 0.55
    )
    lower_rebound_watch = (
        bandwidth_low_30
        and df["Percent B"].tail(8).min() <= 0.12
        and between_middle_upper
        and volume_ratio >= 1.1
    )
    near_middle_watch = medium_squeeze and near_middle_from_below and volume_mild
    expansion_break = expanding and upper_up and lower_down and volume_big and (above_upper or near_upper)

    squeeze_triggered = False
    expansion_triggered = False
    upper_breakout_triggered = False
    middle_rebound_triggered = False
    lower_rebound_triggered = False

    if strong_squeeze:
        if bandwidth_100d_low:
            result.signals.append("强收口：Bandwidth创近100日新低")
        else:
            result.signals.append(f"强收口：Bandwidth低于近8个月10%分位且持续{squeeze_days_10}天")
        result.signals.append("中轨附近窄幅震荡")
        squeeze_triggered = True
    elif medium_squeeze:
        result.signals.append(f"中等收口：Bandwidth低于近8个月20%分位且持续{squeeze_days_20}天")
        result.signals.append("中轨附近窄幅震荡")
        squeeze_triggered = True
    elif weak_squeeze and near_middle:
        if squeeze_narrowing and low_bandwidth_hold_30:
            result.signals.append(f"弱收口：Bandwidth低于近8个月30%分位且最近8个交易日及以上连续下降，并已连续{squeeze_days_30}天维持低分位")
        elif squeeze_narrowing:
            result.signals.append("弱收口：Bandwidth低于近8个月30%分位且最近8个交易日及以上连续下降")
        else:
            result.signals.append(f"弱收口：Bandwidth连续{squeeze_days_30}天维持低于近8个月30%分位")
        result.signals.append("中轨附近窄幅震荡")
        squeeze_triggered = True

    if expansion_break:
        result.signals.append(f"布林带向上开口（最近{min(expansion_days, 5)}天）")
        result.signals.append("成交量放大")
        expansion_triggered = True

    if above_upper:
        upper_breakout_triggered = True
        if volume_huge and expanding:
            result.signals.append("放量突破上轨-强")
        elif volume_ratio >= 1.3:
            result.signals.append("放量突破上轨-中")
        else:
            result.signals.append("突破上轨但量能不足-弱")
        if two_day_above_upper:
            result.signals.append("连续2日站上上轨")

    if middle_rebound and volume_big:
        middle_rebound_triggered = True
        result.signals.append("中轨反弹并接近上轨")
        result.signals.append("成交量放大")

    if lower_rebound_watch:
        lower_rebound_triggered = True
        result.signals.append("下轨反弹：处于中轨与上轨之间")
        result.signals.append("成交量温和改善")

    if near_middle_watch and not middle_rebound_triggered:
        result.signals.append("接近突破中轨")

    if low_price_area and result.signals:
        result.signals.append("低位形态")

    medium_squeeze_setup = (strong_squeeze or medium_squeeze) and (crossed_middle_from_below or between_middle_upper) and volume_mild

    if upper_breakout_triggered and volume_huge and expanding:
        result.signal_grade = "A 强信号"
    elif expansion_triggered and above_upper and volume_big:
        result.signal_grade = "A 强信号" if (strong_return or two_day_above_upper or low_price_area) else "B 中等信号"
    elif middle_rebound_triggered and volume_big and close > middle:
        result.signal_grade = "A 强信号" if (low_price_area or expanding or strong_return) else "B 中等信号"
    elif upper_breakout_triggered and volume_ratio >= 1.3:
        result.signal_grade = "B 中等信号"
    elif medium_squeeze_setup:
        result.signal_grade = "B 中等信号"
    elif lower_rebound_triggered and between_middle_upper and volume_ratio >= 1.1:
        result.signal_grade = "B 中等信号"
    elif near_middle_watch:
        result.signal_grade = "B 中等信号"
    elif expansion_triggered:
        result.signal_grade = "B 中等信号"
    elif lower_rebound_triggered:
        result.signal_grade = "B 中等信号"
    elif strong_squeeze:
        result.signal_grade = "C 观察信号"
    elif weak_squeeze and near_middle and not (strong_squeeze or medium_squeeze):
        result.signal_grade = "C 观察信号"
    elif upper_breakout_triggered:
        result.signal_grade = "C 观察信号"
    else:
        result.signal_grade = "无明显信号"
        result.signals = []

    result.conclusion = build_technical_conclusion(
        result, squeeze_triggered, expansion_triggered, upper_breakout_triggered, middle_rebound_triggered,
        lower_rebound_triggered, above_upper, below_lower, low_price_area, volume_big
    )
    result.risk_notes = build_risk_notes(result, above_upper, below_lower, expansion_triggered, volume_big, two_day_above_upper)
    return result


def build_technical_conclusion(
    result: BollResult,
    squeeze_triggered: bool,
    expansion_triggered: bool,
    upper_breakout_triggered: bool,
    middle_rebound_triggered: bool,
    lower_rebound_triggered: bool,
    above_upper: bool,
    below_lower: bool,
    low_price_area: bool,
    volume_big: bool,
) -> str:
    if upper_breakout_triggered and above_upper and volume_big:
        return "股价放量突破布林线上轨，说明短线多头动能较强，可能进入强势上涨阶段，但也要注意短线过热和回踩风险。"
    if upper_breakout_triggered and above_upper:
        return "股价突破布林线上轨，但成交量确认不足，可能是假突破，需要下一交易日确认。"
    if expansion_triggered and above_upper:
        return "布林带开口，说明波动率开始放大，可能进入趋势启动阶段，需要结合突破方向判断。当前方向偏向上。"
    if expansion_triggered:
        return "布林带向上开口，说明波动率开始放大，股价接近上轨，可能进入趋势启动观察阶段，需要继续结合突破方向和成交量判断。"
    if middle_rebound_triggered:
        return "股价从中轨附近反弹并接近上轨，且成交量放大，说明短线修复增强；若后续能维持中轨支撑并突破上轨，形态有望继续改善。"
    if lower_rebound_triggered:
        if result.percent_b is not None and result.percent_b >= 0.50:
            return "股价从下轨附近反弹并回到中轨上方，目前处于中轨与上轨之间，说明短线修复有所增强，但仍需要成交量和中轨支撑继续确认。"
        return "股价从下轨附近反弹，但仍处于下轨与中轨之间，属于观察级别的修复信号，需要等待重新站上中轨确认。"
    if squeeze_triggered:
        return "布林带收口，说明波动率下降，股价可能处于蓄势阶段，但方向尚未确认。"
    return "未满足当前 BOLL 触发条件，不列入触发。"


def build_risk_notes(
    result: BollResult,
    above_upper: bool,
    below_lower: bool,
    expansion_triggered: bool,
    volume_big: bool,
    two_day_above_upper: bool,
) -> List[str]:
    risks = []
    if above_upper and not volume_big:
        risks.append("突破上轨但成交量不足，存在假突破风险。")
    if above_upper and result.percent_b is not None and result.percent_b > 1.15:
        risks.append("股价明显高于上轨，短线可能过热，需防回踩。")
    if above_upper and not two_day_above_upper:
        risks.append("尚未连续两日站上上轨，需要下一交易日确认。")
    if not expansion_triggered and not above_upper:
        risks.append("当前方向仍未完全确认，需要观察成交量和中轨支撑。")
    if below_lower:
        risks.append("价格跌破下轨且方向偏弱，需警惕趋势继续下行。")
    if not risks:
        risks.append("重点观察成交量变化以及中轨支撑/压力是否有效。")
    return risks


# =========================
# 图表
# =========================

def find_font() -> Optional[str]:
    for path in [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\arial.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]:
        if Path(path).exists():
            return path
    return None


def create_boll_chart(result: BollResult, image_path: Path) -> None:
    if result.df is None or result.df.empty:
        return
    df = result.df.tail(90).copy()
    width, height = 1280, 860
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_path = find_font()
    title_font = ImageFont.truetype(font_path, 25) if font_path else ImageFont.load_default()
    label_font = ImageFont.truetype(font_path, 16) if font_path else ImageFont.load_default()
    small_font = ImageFont.truetype(font_path, 13) if font_path else ImageFont.load_default()

    draw.text((70, 24), f"{result.symbol} - {result.company_name} BOLL 布林线", fill="#111111", font=title_font)
    signal_text = "；".join(result.signals[:4]) if result.signals else "无明显信号"
    draw.text((70, 58), f"最新交易日：{fmt_date(result.latest_date)}    信号：{signal_text}", fill="#333333", font=small_font)

    price_area = (80, 100, 1220, 590)
    vol_area = (80, 645, 1220, 805)
    draw_price_panel(draw, df, result, price_area, label_font, small_font)
    draw_volume_panel(draw, df, vol_area, label_font, small_font)
    image.save(image_path)


def draw_price_panel(draw, df: pd.DataFrame, result: BollResult, area, label_font, small_font) -> None:
    x0, y0, x1, y1 = area
    cols = ["Close", "Upper Band", "Middle Band", "Lower Band"]
    values = []
    for col in cols:
        values.extend([float(v) for v in df[col].dropna().tolist()])
    if not values:
        return
    min_v, max_v = min(values), max(values)
    pad = (max_v - min_v) * 0.08 if max_v > min_v else max_v * 0.05
    min_v -= pad
    max_v += pad

    def px(i: int) -> int:
        if len(df) <= 1:
            return x0
        return int(x0 + i * (x1 - x0) / (len(df) - 1))

    def py(v: float) -> int:
        return int(y1 - (v - min_v) * (y1 - y0) / (max_v - min_v))

    draw.rectangle([x0, y0, x1, y1], outline="#cccccc", width=1)
    draw.text((x0, y0 - 24), "收盘价与 BOLL 通道", fill="#111111", font=label_font)
    draw.text((x0 - 68, y0 - 4), f"{max_v:.2f}", fill="#555555", font=small_font)
    draw.text((x0 - 68, y1 - 10), f"{min_v:.2f}", fill="#555555", font=small_font)

    series = [
        ("Close", "Close", "#111111", 3),
        ("Upper", "Upper Band", "#d62728", 2),
        ("Middle", "Middle Band", "#1f77b4", 2),
        ("Lower", "Lower Band", "#2ca02c", 2),
    ]
    legend_x = x0 + 120
    for name, _, color, lw in series:
        draw.line([legend_x, y0 - 14, legend_x + 28, y0 - 14], fill=color, width=lw)
        draw.text((legend_x + 34, y0 - 22), name, fill="#333333", font=small_font)
        legend_x += 115

    for _, col, color, lw in series:
        points = []
        for i, v in enumerate(df[col].tolist()):
            if pd.isna(v):
                continue
            points.append((px(i), py(float(v))))
        if len(points) >= 2:
            draw.line(points, fill=color, width=lw)

    latest_x = px(len(df) - 1)
    latest_close = float(df["Close"].iloc[-1])
    latest_y = py(latest_close)
    draw.ellipse([latest_x - 5, latest_y - 5, latest_x + 5, latest_y + 5], fill="#f2c200", outline="#111111")
    draw.line([latest_x, y0, latest_x, y1], fill="#999999", width=1)
    draw.text((latest_x - 82, latest_y - 28), f"{fmt_date(df.index[-1])} {latest_close:.2f}", fill="#111111", font=small_font)

    first = df.index[0].strftime("%m-%d")
    last = df.index[-1].strftime("%m-%d")
    draw.text((x0, y1 + 8), first, fill="#555555", font=small_font)
    draw.text((x1 - 45, y1 + 8), last, fill="#555555", font=small_font)


def draw_volume_panel(draw, df: pd.DataFrame, area, label_font, small_font) -> None:
    x0, y0, x1, y1 = area
    volumes = df["Volume"].fillna(0).tolist()
    max_v = max(volumes) if volumes else 0
    if max_v <= 0:
        return

    def px(i: int) -> int:
        if len(df) <= 1:
            return x0
        return int(x0 + i * (x1 - x0) / (len(df) - 1))

    draw.rectangle([x0, y0, x1, y1], outline="#cccccc", width=1)
    draw.text((x0, y0 - 24), "成交量", fill="#111111", font=label_font)
    draw.text((x0 - 72, y0 - 4), f"{max_v / 1_000_000:.1f}M", fill="#555555", font=small_font)

    bar_w = max(2, int((x1 - x0) / max(len(df), 1) * 0.55))
    for i, vol in enumerate(volumes):
        x = px(i)
        top = int(y1 - float(vol) / max_v * (y1 - y0))
        close = df["Close"].iloc[i]
        prev_close = df["Close"].iloc[i - 1] if i > 0 else close
        color = "#2ca02c" if close >= prev_close else "#d62728"
        draw.rectangle([x - bar_w // 2, top, x + bar_w // 2, y1], fill=color)


def generate_charts_for_word(watch: List[BollResult], temp_dir: Path) -> Dict[str, Path]:
    paths: Dict[str, Path] = {}
    for result in watch:
        image_path = temp_dir / f"{result.symbol}_BOLL.png"
        try:
            create_boll_chart(result, image_path)
            if image_path.exists():
                paths[result.symbol] = image_path
        except Exception as exc:
            result.risk_notes.append(f"图表生成失败：{exc}")
    return paths


# =========================
# 报告
# =========================

def triggered_results(results: List[BollResult]) -> List[BollResult]:
    order = {"A 强信号": 0, "B 中等信号": 1, "C 观察信号": 2, "无明显信号": 3}
    items = [r for r in results if not r.data_error and r.signal_grade != "无明显信号"]
    return sorted(items, key=lambda r: (order.get(r.signal_grade, 9), -(r.volume_ratio or 0), r.symbol))


def no_signal_results(results: List[BollResult]) -> List[BollResult]:
    return [r for r in results if not r.data_error and r.signal_grade == "无明显信号"]


def failed_results(results: List[BollResult]) -> List[BollResult]:
    return [r for r in results if r.data_error]


def build_markdown_report(
    results: List[BollResult],
    report_date: date,
    market_info: Dict[str, object],
    data_latest_day: Optional[pd.Timestamp],
) -> str:
    triggered = triggered_results(results)
    failed = failed_results(results)
    no_signal = no_signal_results(results)

    lines = []
    lines.append(f"# 美股 BOLL 布林线监控报告 - {report_date.strftime('%Y-%m-%d')}")
    lines.append("")
    lines.append(f"- 生成时间（美东）：{market_info['now_et'].strftime('%Y-%m-%d %H:%M:%S %Z')}")
    lines.append(f"- 最新可用交易日：{fmt_date(data_latest_day)}")
    lines.append("- 数据源优先级：Yahoo Chart API -> yfinance（可选备用，无需额外配置）")
    if market_info.get("is_trading_day") is False:
        lines.append("- 提醒：当前美东日期不是美股交易日，本次使用最近可用交易日的收盘数据计算。")
    elif not market_info.get("after_close"):
        data_date = data_latest_day.date() if data_latest_day is not None else None
        if data_date is not None and data_date < market_info["today_et"]:
            lines.append("- 提醒：当前美东交易日尚未收盘，本次使用最近可用交易日的收盘数据计算。")
        else:
            lines.append("- 提醒：当前时间早于美股常规收盘后，今日数据可能尚未完整更新。")
    lines.append("- 策略说明：本报告按近8个月 Bandwidth 分位判断收口；低于10%且持续5天以上或创近100日新低为强收口，低于20%且持续5天以上为中等收口。弱收口必须满足 Bandwidth 低于近8个月30%分位，并且最近8个交易日及以上连续下降，或连续8个交易日及以上维持低于近8个月30%分位。观察信号要求弱收口且股价接近中轨。向上开口要求 Bandwidth 最近2-5日明显扩大、上轨向上、下轨向下、股价突破或接近上轨并伴随成交量放大；向下开口不列入触发。放量突破上轨按量能、涨幅和带宽扩张分级。")
    lines.append("- 说明：本报告只监控 BOLL 布林线，不构成投资建议。")
    lines.append("")

    lines.append("## 一、今日触发名单")
    lines.append("")
    if triggered:
        lines.append("| 股票代码 | 公司名称 | 等级 | 触发信号 | 位置 | Bandwidth | %B | 量比 | 简要结论 |")
        lines.append("|---|---|---|---|---|---|---|---|---|")
        for r in triggered:
            lines.append(
                f"| {r.symbol} | {r.company_name} | {r.signal_grade} | {'；'.join(r.signals[:4])} "
                f"| {r.price_position} | {fmt_pct(r.bandwidth)} | {fmt_num(r.percent_b, 3)} "
                f"| {fmt_num(r.volume_ratio, 2)} | {r.conclusion[:45]} |"
            )
    else:
        lines.append("今日未发现符合条件的 BOLL 触发信号。")
    lines.append("")

    lines.append("## 二、重点结论")
    lines.append("")
    strong = [r for r in triggered if r.signal_grade == "A 强信号"]
    medium = [r for r in triggered if r.signal_grade == "B 中等信号"]
    watch = [r for r in triggered if r.signal_grade == "C 观察信号"]
    lines.append(f"- 强信号：{', '.join(r.symbol for r in strong) if strong else '无'}。")
    lines.append(f"- 中等信号：{', '.join(r.symbol for r in medium) if medium else '无'}。")
    lines.append(f"- 观察信号：{', '.join(r.symbol for r in watch) if watch else '无'}。")
    if triggered:
        top = triggered[:5]
        lines.append(f"- 今日优先复核：{', '.join(r.symbol for r in top)}。优先看是否放量、能否站稳上轨或中轨。")
    lines.append("- 风险提示：持续收窄只代表波动率降低，不代表方向已经确认；开口和突破也可能是假突破，需要观察成交量和中轨支撑。")
    lines.append("")

    lines.append("## 三、个股分析：技术面解读")
    lines.append("")
    if triggered:
        for r in triggered:
            latest = r.df.iloc[-1] if r.df is not None and not r.df.empty else None
            lines.append(f"### {r.symbol} - {r.company_name}")
            lines.append(f"- 信号等级：{r.signal_grade}")
            lines.append(f"- 最新收盘价：{fmt_price(r.close)}")
            if latest is not None:
                lines.append(
                    f"- BOLL：上轨 {fmt_price(latest['Upper Band'])}；中轨 {fmt_price(latest['Middle Band'])}；下轨 {fmt_price(latest['Lower Band'])}"
                )
            lines.append(f"- Bandwidth：{fmt_pct(r.bandwidth)}；历史分位：{fmt_pct(r.bandwidth_percentile)}；%B：{fmt_num(r.percent_b, 3)}")
            lines.append(f"- 成交量：当日 {fmt_num(r.latest_volume, 0)}；20日均量 {fmt_num(r.avg_volume20, 0)}；量比 {fmt_num(r.volume_ratio, 2)}")
            lines.append(f"- 触发信号：{'；'.join(r.signals)}")
            lines.append(f"- 技术面解读：{r.conclusion}")
            lines.append(f"- 是否需要下一交易日确认：是。重点观察收盘价是否继续位于关键轨道一侧，以及成交量是否延续。")
            lines.append(f"- 风险提示：{'；'.join(r.risk_notes)}")
            lines.append("")
    else:
        lines.append("无触发个股。")
        lines.append("")

    lines.append("## 四、数据异常或无法判断")
    lines.append("")
    if failed:
        lines.append("| 股票代码 | 公司名称 | 原因 |")
        lines.append("|---|---|---|")
        for r in failed:
            lines.append(f"| {r.symbol} | {r.company_name or r.symbol} | {r.data_error} |")
    else:
        lines.append("未发现数据异常或无法判断的股票。")
    lines.append("")

    if no_signal:
        lines.append("## 五、无明显 BOLL 信号")
        lines.append("")
        lines.append(", ".join(f"{r.symbol}（{r.company_name}）" for r in no_signal))
        lines.append("")

    return "\n".join(lines)


def write_markdown(path: Path, content: str) -> None:
    path.write_text(content, encoding="utf-8-sig")


def cleanup_markdown_reports(output_dir: Path) -> None:
    for md_path in output_dir.glob("美股BOLL布林线监控报告-*.md"):
        try:
            md_path.unlink()
        except Exception:
            pass


def append_triggered_symbols_txt(output_dir: Path, report_date: date, triggered: List[BollResult]) -> None:
    if not triggered:
        return
    txt_path = output_dir / "BOLL.txt"
    existing = set()
    existing_lines: List[str] = []
    key_to_count: Dict[str, int] = {}
    symbol_counts: Dict[str, int] = {}

    def format_signal_grade_for_txt(signal_grade: str) -> str:
        grade = (signal_grade or "").strip()
        if len(grade) >= 2 and grade[0] in {"A", "B", "C"} and grade[1].isspace():
            return f"{grade[0]}. {grade[2:].strip()}"
        if len(grade) >= 2 and grade[0] in {"A", "B", "C"} and grade[1] == ".":
            return grade
        return grade

    def format_percentile_for_txt(value) -> str:
        percentile = safe_float(value)
        if percentile is None:
            return "历史分位：N/A"
        return f"历史分位：{percentile * 100:.2f}%"

    def parse_symbol_from_line(line: str) -> str:
        parts = line.strip().split()
        if len(parts) < 2:
            return ""
        symbol_text = parts[1].strip()
        if "（" in symbol_text:
            symbol_text = symbol_text.split("（", 1)[0].strip()
        elif "(" in symbol_text:
            symbol_text = symbol_text.split("(", 1)[0].strip()
        return symbol_text

    def parse_count_from_line(line: str) -> int:
        parts = line.strip().split()
        if len(parts) < 2:
            return 1
        symbol_text = parts[1].strip()
        match = re.search(r"[（(](\d+)[）)]", symbol_text)
        return int(match.group(1)) if match else 1

    def parse_date_symbol_key(line: str) -> str:
        parts = line.strip().split()
        if len(parts) < 2:
            return line.strip()
        symbol_text = parse_symbol_from_line(line)
        return f"{parts[0]}\t{symbol_text}"

    if txt_path.exists():
        try:
            for raw_line in txt_path.read_text(encoding="utf-8-sig").splitlines():
                line = raw_line.strip()
                if not line:
                    continue
                existing_lines.append(line)
                existing.add(parse_date_symbol_key(line))
                symbol = parse_symbol_from_line(line)
                if symbol:
                    key_to_count[parse_date_symbol_key(line)] = parse_count_from_line(line)
                    symbol_counts[symbol] = symbol_counts.get(symbol, 0) + 1
        except Exception:
            existing = set()
            existing_lines = []
            key_to_count = {}
            symbol_counts = {}

    new_lines = []
    changed = False
    date_text = report_date.strftime("%Y-%m-%d")

    def build_line(result: BollResult, trigger_count: int) -> str:
        symbol_text = f"{result.symbol}（{trigger_count}）" if trigger_count > 1 else result.symbol
        signal_grade = format_signal_grade_for_txt(result.signal_grade)
        percentile_text = format_percentile_for_txt(result.bandwidth_percentile)
        detail_parts = [part for part in (signal_grade, percentile_text) if part]
        return f"{date_text}\t{symbol_text}\t" + "\t".join(detail_parts) if detail_parts else f"{date_text}\t{symbol_text}"

    for result in triggered:
        key = f"{date_text}\t{result.symbol}"
        if key in existing:
            trigger_count = key_to_count.get(key, 1)
            updated_line = build_line(result, trigger_count)
            for idx, line in enumerate(existing_lines):
                if parse_date_symbol_key(line) == key and line != updated_line:
                    existing_lines[idx] = updated_line
                    changed = True
                    break
        else:
            next_count = symbol_counts.get(result.symbol, 0) + 1
            new_lines.append(build_line(result, next_count))
            existing.add(key)
            symbol_counts[result.symbol] = next_count
            changed = True

    if changed:
        txt_path.parent.mkdir(parents=True, exist_ok=True)
        all_lines = existing_lines + new_lines
        txt_path.write_text("\n".join(all_lines) + "\n", encoding="utf-8-sig")


def parse_plain_symbol(symbol_text: str) -> str:
    symbol = symbol_text.strip().upper()
    if "（" in symbol:
        symbol = symbol.split("（", 1)[0].strip()
    elif "(" in symbol:
        symbol = symbol.split("(", 1)[0].strip()
    return symbol


def read_boll_trigger_history(output_dir: Path) -> Dict[str, List[date]]:
    txt_path = output_dir / "BOLL.txt"
    history: Dict[str, List[date]] = {}
    if not txt_path.exists():
        return history
    try:
        lines = txt_path.read_text(encoding="utf-8-sig").splitlines()
    except Exception as exc:
        print(f"读取 BOLL.txt 失败，无法更新 Excel 跟踪文件：{exc}")
        return history

    for raw_line in lines:
        parts = raw_line.strip().split()
        if len(parts) < 2:
            continue
        try:
            trigger_date = date.fromisoformat(parts[0])
        except ValueError:
            continue
        symbol = parse_plain_symbol(parts[1])
        if not symbol:
            continue
        history.setdefault(symbol, []).append(trigger_date)
    for symbol in list(history):
        history[symbol] = sorted(set(history[symbol]))
    return history


def safe_tracking_filename(symbol: str) -> str:
    safe = re.sub(r"[^A-Z0-9._-]+", "_", symbol.upper()).strip("._-")
    return safe or "UNKNOWN"


def tracking_workbook_path(tracking_dir: Path, symbol: str, episode_number: int, trigger_date: date) -> Path:
    base_name = safe_tracking_filename(symbol)
    dated_path = tracking_dir / f"{base_name}_{episode_number}_{trigger_date.strftime('%Y%m%d')}.xlsx"
    legacy_path = tracking_dir / f"{base_name}.xlsx" if episode_number == 1 else tracking_dir / f"{base_name}_{episode_number}.xlsx"
    if dated_path.exists():
        return dated_path
    if legacy_path.exists():
        return legacy_path
    return dated_path


def tracking_prices_from_df(df: pd.DataFrame, trigger_date: date) -> pd.DataFrame:
    df = df.copy()
    df["TrackClose"] = pd.to_numeric(df["Close"], errors="coerce")
    df["TrackDailyChangePct"] = df["TrackClose"].pct_change() * 100
    df = df[df.index.date >= trigger_date].copy()
    if not df.empty:
        df = df.head(TRACKING_DAYS)
        df = df.dropna(subset=["TrackClose"])
    return df


def is_trigger_inside_episode(df: pd.DataFrame, episode_start: date, trigger_date: date) -> bool:
    episode_window = df[(df.index.date >= episode_start) & (df.index.date <= trigger_date)]
    if episode_window.empty:
        return True
    return len(episode_window) <= TRACKING_DAYS


def build_tracking_episodes(df: pd.DataFrame, trigger_dates: List[date]) -> List[date]:
    episodes: List[date] = []
    for trigger_date in sorted(set(trigger_dates)):
        if not episodes:
            episodes.append(trigger_date)
            continue
        if not is_trigger_inside_episode(df, episodes[-1], trigger_date):
            episodes.append(trigger_date)
    return episodes


def write_tracking_workbook(path: Path, symbol: str, trigger_date: date, df: pd.DataFrame) -> None:
    if Workbook is None:
        raise RuntimeError("未安装 openpyxl，无法生成 Excel 跟踪文件")
    path.parent.mkdir(parents=True, exist_ok=True)

    wb = Workbook()
    ws = wb.active
    ws.title = "30日跟踪"
    ws.freeze_panes = "A8"

    title = f"{symbol} 触发后30个交易日股价跟踪"
    ws.merge_cells("A1:D1")
    ws["A1"] = title
    ws["A1"].font = Font(bold=True, size=14)
    ws["A1"].alignment = Alignment(horizontal="center")

    base_close = float(df.iloc[0]["TrackClose"])
    latest_date = df.index[-1].date()
    meta_rows = [
        ("股票代码", symbol),
        ("首次触发日", trigger_date),
        ("触发日收盘价", base_close),
        ("更新到交易日", latest_date),
        ("已记录交易日", f"{len(df)}/{TRACKING_DAYS}"),
    ]
    for idx, (label, value) in enumerate(meta_rows, start=2):
        ws.cell(row=idx, column=1, value=label)
        ws.cell(row=idx, column=2, value=value)
        ws.cell(row=idx, column=1).font = Font(bold=True)
    ws["B3"].number_format = "yyyy-mm-dd"
    ws["B4"].number_format = "0.00"
    ws["B5"].number_format = "yyyy-mm-dd"

    headers = ["交易日", "收盘价", "当日涨跌幅%", "相对触发日涨跌幅", "进度"]
    header_row = 7
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=header_row, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E78")
        cell.alignment = Alignment(horizontal="center")

    first_data_row = header_row + 1
    for offset, (idx, row) in enumerate(df.iterrows(), start=0):
        excel_row = first_data_row + offset
        trade_date = idx.date() if hasattr(idx, "date") else idx
        ws.cell(row=excel_row, column=1, value=trade_date)
        ws.cell(row=excel_row, column=2, value=float(row["TrackClose"]))
        daily_change = row.get("TrackDailyChangePct")
        ws.cell(row=excel_row, column=3, value=None if pd.isna(daily_change) else float(daily_change))
        ws.cell(row=excel_row, column=4, value=f"=B{excel_row}/$B${first_data_row}-1")
        ws.cell(row=excel_row, column=5, value=f"第{offset + 1}/{TRACKING_DAYS}个交易日")
        ws.cell(row=excel_row, column=1).number_format = "yyyy-mm-dd"
        ws.cell(row=excel_row, column=2).number_format = "0.00"
        ws.cell(row=excel_row, column=3).number_format = "0.00"
        ws.cell(row=excel_row, column=4).number_format = "0.00%"

    last_data_row = first_data_row + len(df) - 1
    chart = LineChart()
    chart.title = f"{symbol} 相对触发日涨跌幅"
    chart.y_axis.title = "涨跌幅"
    chart.x_axis.title = "交易日"
    chart.y_axis.numFmt = "0%"
    chart.height = 10
    chart.width = 18
    data = Reference(ws, min_col=4, min_row=header_row, max_row=last_data_row)
    cats = Reference(ws, min_col=1, min_row=first_data_row, max_row=last_data_row)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    chart.legend = None
    ws.add_chart(chart, "F2")

    widths = {"A": 14, "B": 12, "C": 18, "D": 18}
    for col, width in widths.items():
        ws.column_dimensions[col].width = width
    for col in range(5, 15):
        ws.column_dimensions[get_column_letter(col)].width = 12

    wb.save(path)


def update_boll_tracking_excels(output_dir: Path) -> List[Path]:
    history = read_boll_trigger_history(output_dir)
    if not history:
        print("BOLL.txt 暂无触发记录，未生成 Excel 跟踪文件。")
        return []
    if Workbook is None:
        print("未安装 openpyxl，跳过 Excel 跟踪文件生成。")
        return []

    tracking_dir = EXCEL_TRACKING_DIR
    tracking_dir.mkdir(parents=True, exist_ok=True)
    updated_paths: List[Path] = []

    for symbol, trigger_dates in sorted(history.items()):
        try:
            full_df, source, err = fetch_ohlcv(symbol)
            if full_df.empty:
                print(f"{symbol} Excel 跟踪更新失败：{err or '无法获取行情数据'}")
                continue
            episodes = build_tracking_episodes(full_df, trigger_dates)
            for episode_number, trigger_date in enumerate(episodes, start=1):
                df = tracking_prices_from_df(full_df, trigger_date)
                if df.empty:
                    print(f"{symbol}_{episode_number} Excel 跟踪更新失败：没有 {trigger_date} 之后的收盘数据")
                    continue
                workbook_path = tracking_workbook_path(tracking_dir, symbol, episode_number, trigger_date)
                write_tracking_workbook(workbook_path, symbol, trigger_date, df)
                updated_paths.append(workbook_path)
                print(f"已更新 {symbol}_{episode_number} 跟踪 Excel：{workbook_path}，记录 {len(df)}/{TRACKING_DAYS} 个交易日。")
        except Exception as exc:
            print(f"{symbol} Excel 跟踪更新失败：{exc}")
    return updated_paths


def add_markdown_table_to_doc(doc, lines: List[str], start_idx: int) -> int:
    header = [c.strip() for c in lines[start_idx].strip("|").split("|")]
    rows = []
    idx = start_idx + 2
    while idx < len(lines) and lines[idx].startswith("|"):
        rows.append([c.strip() for c in lines[idx].strip("|").split("|")])
        idx += 1
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, header):
        cell.text = text
    for row_values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_values):
            cell.text = text
    return idx


def write_word(path: Path, markdown_content: str, triggered: List[BollResult], chart_paths: Dict[str, Path]) -> None:
    if Document is None:
        raise RuntimeError("未安装 python-docx，无法生成 Word 文件。请运行：pip install python-docx")
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)

    lines = markdown_content.splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()
        if not stripped:
            idx += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("|") and idx + 1 < len(lines) and set(lines[idx + 1].replace("|", "").strip()) <= {"-", ":"}:
            idx = add_markdown_table_to_doc(doc, lines, idx)
            continue
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)
        idx += 1

    if triggered and chart_paths:
        doc.add_page_break()
        doc.add_heading("六、触发股票 BOLL 图表", level=2)
        for result in triggered:
            image_path = chart_paths.get(result.symbol)
            if image_path and Path(image_path).exists():
                doc.add_heading(f"{result.symbol} - {result.company_name}", level=3)
                doc.add_picture(str(image_path), width=Inches(6.6))

    doc.save(path)


def timestamped_path(path: Path) -> Path:
    stamp = datetime.now().strftime("%H%M%S")
    return path.with_name(f"{path.stem}-{stamp}{path.suffix}")


def write_market_closed_report(output_dir: Path, market_info: Dict[str, object]) -> Tuple[Path, Path]:
    today = market_info["today_et"]
    content = (
        f"# 美股 BOLL 布林线监控报告 - {today.strftime('%Y-%m-%d')}\n\n"
        f"今日美股休市，无需监控。\n\n"
        f"- 当前美东时间：{market_info['now_et'].strftime('%Y-%m-%d %H:%M:%S %Z')}\n"
        f"- 下一个交易日：{fmt_date(market_info.get('next_trading_day'))}\n"
    )
    docx_path = output_dir / f"美股BOLL布林线监控报告-{today.strftime('%Y-%m-%d')}.docx"
    try:
        write_word(docx_path, content, [], {})
    except PermissionError:
        docx_path = timestamped_path(docx_path)
        write_word(docx_path, content, [], {})
    return docx_path, docx_path


# =========================
# 主流程
# =========================

def run_monitor() -> Tuple[Path, Path]:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    cleanup_markdown_reports(OUTPUT_DIR)

    market_info = get_market_info()

    results: List[BollResult] = []
    print("开始获取数据并计算 BOLL...")
    for idx, symbol in enumerate(STOCK_POOL, start=1):
        print(f"[{idx}/{len(STOCK_POOL)}] {symbol}")
        result = BollResult(symbol=symbol, company_name=get_company_name(symbol))
        df, source, err = fetch_ohlcv(symbol)
        if df.empty:
            result.data_error = err or "无法获取日线 OHLCV 数据"
            results.append(result)
            continue
        result.data_source = source
        result.df = df
        result = evaluate_boll(result)
        results.append(result)
        time.sleep(SLEEP_BETWEEN_SYMBOLS)

    data_day = latest_data_date(results)
    report_day = data_day.date() if data_day is not None else market_info["today_et"]
    triggered = triggered_results(results)

    temp_dir = Path(tempfile.mkdtemp(prefix=f"boll_charts_{report_day.strftime('%Y%m%d')}_"))
    try:
        chart_paths = generate_charts_for_word(triggered, temp_dir)
        content = build_markdown_report(results, report_day, market_info, data_day)
        docx_path = OUTPUT_DIR / f"美股BOLL布林线监控报告-{report_day.strftime('%Y-%m-%d')}.docx"
        try:
            write_word(docx_path, content, triggered, chart_paths)
        except PermissionError:
            docx_path = timestamped_path(docx_path)
            write_word(docx_path, content, triggered, chart_paths)
        append_triggered_symbols_txt(OUTPUT_DIR, report_day, triggered)
        update_boll_tracking_excels(OUTPUT_DIR)
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)

    return docx_path, docx_path


def main() -> None:
    try:
        _, docx_path = run_monitor()
        print("")
        print("报告生成完成：")
        print(f"Word:     {docx_path}")
    except KeyboardInterrupt:
        print("用户中断。")
    except Exception as exc:
        print("运行失败：", exc)
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()

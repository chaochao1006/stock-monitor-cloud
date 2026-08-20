from __future__ import annotations

import datetime as dt
import json
import logging
import math
import os
import time
import urllib.request
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from zoneinfo import ZoneInfo

import numpy as np
import pandas as pd
import yfinance as yf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


# =========================
# 基础配置
# =========================

TICKERS = [
    "OPTX",
    "ASTS",
    "CRWV",
    "FLY",
    "AAOI",
    "LUNR",
    "CRCL",
    "ONDS",
]

REPORT_BASE_DIR = Path(os.environ.get("REPORT_BASE_DIR", Path(__file__).resolve().parents[1] / "data"))
OUTPUT_DIR = REPORT_BASE_DIR / "drop" / "short-risk"
REPORT_PREFIX = "美股短线反弹风险监控"
TRIGGER_TXT_NAME = "RISK.txt"

HISTORY_PERIOD = "18mo"
INTERVAL = "1d"
SHORT_RISK_TRIGGER_SCORE = 6
REQUEST_PAUSE_SECONDS = 0.3


@dataclass
class ShortRiskResult:
    """保存单只股票的短线反弹状态、过热风险和报告文字。"""

    ticker: str
    date: Optional[dt.date] = None
    close: Optional[float] = None
    ma5: Optional[float] = None
    ma10: Optional[float] = None
    ma20: Optional[float] = None
    ma50: Optional[float] = None
    rsi6: Optional[float] = None
    rsi12: Optional[float] = None
    k: Optional[float] = None
    d: Optional[float] = None
    j: Optional[float] = None
    deviation20: Optional[float] = None
    return5: Optional[float] = None
    return10: Optional[float] = None
    volume_ratio: Optional[float] = None
    rebound_score: int = 0
    rebound_state: str = "无法判断"
    short_risk_score: int = 0
    risk_level: str = "无法判断"
    risk_sources: List[str] = field(default_factory=list)
    rebound_reasons: List[str] = field(default_factory=list)
    indicator_summary: str = ""
    observation: str = ""
    error: str = ""
    data_warning: str = ""


def setup_logging() -> None:
    """初始化日志输出。"""

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    log_path = OUTPUT_DIR / "short-risk.log"
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[
            logging.FileHandler(log_path, encoding="utf-8"),
            logging.StreamHandler(),
        ],
    )


def fmt_num(value: Optional[float], digits: int = 2) -> str:
    """格式化普通数值。"""

    if value is None:
        return "-"
    try:
        if pd.isna(value) or math.isinf(float(value)):
            return "-"
        return f"{float(value):.{digits}f}"
    except Exception:
        return "-"


def fmt_price(value: Optional[float]) -> str:
    """格式化价格。"""

    return fmt_num(value, 2)


def safe_float(value) -> Optional[float]:
    """安全转换为浮点数。"""

    try:
        if value is None or pd.isna(value):
            return None
        return float(value)
    except Exception:
        return None


def eastern_now() -> dt.datetime:
    """获取当前美东时间。"""

    return dt.datetime.now(ZoneInfo("America/New_York"))


def previous_weekday(day: dt.date) -> dt.date:
    """返回指定日期之前最近的工作日。"""

    day = day - dt.timedelta(days=1)
    while day.weekday() >= 5:
        day = day - dt.timedelta(days=1)
    return day


def latest_closed_us_trading_day(now_et: Optional[dt.datetime] = None) -> dt.date:
    """估算最近一个已收盘美股交易日，节假日由行情数据自动回退到上一个可用交易日。"""

    now_et = now_et or eastern_now()
    today = now_et.date()
    if today.weekday() >= 5:
        return previous_weekday(today + dt.timedelta(days=1))
    if now_et.time() < dt.time(17, 0):
        return previous_weekday(today)
    return today


def flatten_yfinance_columns(df: pd.DataFrame) -> pd.DataFrame:
    """兼容 yfinance 可能返回的多级列名。"""

    if isinstance(df.columns, pd.MultiIndex):
        df.columns = df.columns.get_level_values(0)
    return df


def normalize_ohlcv(df: pd.DataFrame, target_day: dt.date) -> Tuple[pd.DataFrame, str]:
    """清洗行情数据，并限制到最近已收盘交易日及以前。"""

    if df is None or df.empty:
        return pd.DataFrame(), "未返回数据"
    df = flatten_yfinance_columns(df.copy())
    required = ["Open", "High", "Low", "Close", "Volume"]
    missing = [col for col in required if col not in df.columns]
    if missing:
        return pd.DataFrame(), f"缺少字段：{', '.join(missing)}"
    df = df[required].copy()
    for col in required:
        df[col] = pd.to_numeric(df[col], errors="coerce")
    df = df.dropna(subset=required).sort_index()
    df = df[df.index.date <= target_day]
    if df.empty:
        return pd.DataFrame(), f"没有 {target_day} 及以前的有效收盘数据"
    return df, ""


def fetch_yahoo_chart(ticker: str, target_day: dt.date) -> Tuple[pd.DataFrame, str]:
    """yfinance限流时使用Yahoo Chart API备用获取数据。"""

    end_dt = dt.datetime.combine(target_day + dt.timedelta(days=2), dt.time(0, 0), tzinfo=ZoneInfo("America/New_York"))
    start_dt = end_dt - dt.timedelta(days=560)
    url = (
        f"https://query1.finance.yahoo.com/v8/finance/chart/{ticker}"
        f"?period1={int(start_dt.timestamp())}&period2={int(end_dt.timestamp())}"
        f"&interval=1d&events=history&includeAdjustedClose=true"
    )
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=20) as resp:
            payload = json.loads(resp.read().decode("utf-8"))
        result = payload["chart"]["result"][0]
        timestamps = result.get("timestamp") or []
        quote = result["indicators"]["quote"][0]
        rows = []
        for idx, stamp in enumerate(timestamps):
            rows.append(
                {
                    "Date": dt.datetime.fromtimestamp(stamp, ZoneInfo("America/New_York")).date(),
                    "Open": quote.get("open", [None] * len(timestamps))[idx],
                    "High": quote.get("high", [None] * len(timestamps))[idx],
                    "Low": quote.get("low", [None] * len(timestamps))[idx],
                    "Close": quote.get("close", [None] * len(timestamps))[idx],
                    "Volume": quote.get("volume", [None] * len(timestamps))[idx],
                }
            )
        df = pd.DataFrame(rows)
        if df.empty:
            return pd.DataFrame(), "Yahoo Chart API未返回数据"
        df = df.drop_duplicates("Date").sort_values("Date")
        df.index = pd.to_datetime(df["Date"])
        return normalize_ohlcv(df, target_day)
    except Exception as exc:
        return pd.DataFrame(), f"Yahoo Chart API失败：{exc}"


def fetch_stock_data(ticker: str, target_day: dt.date) -> Tuple[pd.DataFrame, str]:
    """优先用 yfinance 获取数据，失败时自动使用Yahoo Chart API备用。"""

    try:
        df = yf.download(
            ticker,
            period=HISTORY_PERIOD,
            interval=INTERVAL,
            auto_adjust=False,
            progress=False,
            threads=False,
            timeout=15,
        )
    except Exception as exc:
        df = pd.DataFrame()
        yfinance_error = f"yfinance下载失败：{exc}"
    else:
        df, yfinance_error = normalize_ohlcv(df, target_day)
        if not df.empty:
            return df, ""

    fallback_df, fallback_error = fetch_yahoo_chart(ticker, target_day)
    if not fallback_df.empty:
        return fallback_df, ""
    return pd.DataFrame(), f"{yfinance_error}；备用数据源：{fallback_error}"


def calculate_rsi(series: pd.Series, period: int) -> pd.Series:
    """使用Wilder平滑方法计算RSI。"""

    delta = series.diff()
    gain = delta.clip(lower=0)
    loss = -delta.clip(upper=0)
    avg_gain = gain.ewm(alpha=1 / period, adjust=False, min_periods=period).mean()
    avg_loss = loss.ewm(alpha=1 / period, adjust=False, min_periods=period).mean()
    rs = avg_gain / avg_loss.replace(0, np.nan)
    return (100 - 100 / (1 + rs)).fillna(50)


def calculate_kdj(df: pd.DataFrame) -> pd.DataFrame:
    """计算KDJ指标，初始K/D设为50。"""

    df = df.copy()
    low9 = df["Low"].rolling(9).min()
    high9 = df["High"].rolling(9).max()
    rsv = (df["Close"] - low9) / (high9 - low9).replace(0, np.nan) * 100
    rsv = rsv.fillna(50)
    k_values = []
    d_values = []
    k_prev = 50.0
    d_prev = 50.0
    for value in rsv:
        k_today = k_prev * 2 / 3 + float(value) / 3
        d_today = d_prev * 2 / 3 + k_today / 3
        k_values.append(k_today)
        d_values.append(d_today)
        k_prev = k_today
        d_prev = d_today
    df["K"] = k_values
    df["D"] = d_values
    df["J"] = 3 * df["K"] - 2 * df["D"]
    return df


def calculate_indicators(df: pd.DataFrame) -> pd.DataFrame:
    """计算短线反弹风险所需的均线、RSI、KDJ、成交量和涨幅。"""

    df = df.copy()
    for window in [5, 10, 20, 50]:
        df[f"MA{window}"] = df["Close"].rolling(window).mean()
    df["Deviation20"] = (df["Close"] - df["MA20"]) / df["MA20"] * 100
    df["RSI6"] = calculate_rsi(df["Close"], 6)
    df["RSI12"] = calculate_rsi(df["Close"], 12)
    df = calculate_kdj(df)
    df["Volume_MA20"] = df["Volume"].rolling(20).mean()
    df["Volume_Ratio"] = df["Volume"] / df["Volume_MA20"]
    df["Return_5D"] = df["Close"].pct_change(5) * 100
    df["Return_10D"] = df["Close"].pct_change(10) * 100
    return df


def is_short_rebound(df: pd.DataFrame) -> Tuple[int, str, List[str]]:
    """判断是否处于短线反弹状态，满足4项及以上才进入过热评分。"""

    latest = df.iloc[-1]
    reasons: List[str] = []
    score = 0

    checks = [
        (pd.notna(latest["MA5"]) and latest["Close"] > latest["MA5"], "收盘价站上MA5"),
        (pd.notna(latest["MA10"]) and latest["Close"] > latest["MA10"], "收盘价站上MA10"),
        (pd.notna(latest["MA20"]) and latest["Close"] > latest["MA20"], "收盘价站上MA20"),
        (pd.notna(latest["MA5"]) and pd.notna(latest["MA10"]) and latest["MA5"] > latest["MA10"], "MA5高于MA10"),
        (pd.notna(latest["RSI12"]) and latest["RSI12"] > 50, "RSI12站上50"),
        (pd.notna(latest["Return_5D"]) and latest["Return_5D"] > 0, "近5日收益为正"),
        (pd.notna(latest["Return_10D"]) and latest["Return_10D"] > 0, "近10日收益为正"),
    ]
    for passed, reason in checks:
        if passed:
            score += 1
            reasons.append(reason)

    if score >= 5:
        state = "短线反弹确认"
    elif score >= 4:
        state = "短线反弹初步成立"
    else:
        state = "未形成短线反弹"
    return score, state, reasons


def score_short_rebound_risk(df: pd.DataFrame, result: ShortRiskResult) -> int:
    """在短线反弹状态下，判断是否过热或衰竭。"""

    latest = df.iloc[-1]
    prev = df.iloc[-2]
    score = 0
    risks: List[str] = []

    deviation20 = latest["Deviation20"]
    if pd.notna(deviation20):
        if deviation20 > 15:
            score += 2
            risks.append("偏离MA20超过15%")
        elif deviation20 > 8:
            score += 1
            risks.append("偏离MA20超过8%")

    if latest["RSI6"] > 85:
        score += 2
        risks.append("RSI6大于85")
    elif latest["RSI6"] > 75:
        score += 1
        risks.append("RSI6大于75")

    if latest["RSI12"] > 75:
        score += 2
        risks.append("RSI12大于75")
    elif latest["RSI12"] > 65:
        score += 1
        risks.append("RSI12大于65")

    if latest["K"] > 80 and latest["D"] > 80:
        score += 1
        risks.append("KDJ处于高位")

    if latest["J"] > 100:
        score += 2
        risks.append("J线大于100")

    if pd.notna(latest["Return_5D"]):
        if latest["Return_5D"] > 25:
            score += 2
            risks.append("近5日涨幅超过25%")
        elif latest["Return_5D"] > 15:
            score += 1
            risks.append("近5日涨幅超过15%")

    # 放量冲高后缩量：最近5日曾明显放量，最新成交量低于前一日且仍在高位区域。
    recent_volume_ratio = df["Volume_Ratio"].dropna().tail(5)
    if not recent_volume_ratio.empty and recent_volume_ratio.max() >= 1.8:
        if pd.notna(latest["Volume_Ratio"]) and pd.notna(prev["Volume_Ratio"]) and latest["Volume_Ratio"] < prev["Volume_Ratio"]:
            score += 1
            risks.append("放量冲高后缩量")

    if pd.notna(latest["MA5"]) and latest["Close"] < latest["MA5"]:
        score += 1
        risks.append("跌破MA5")

    if latest["RSI6"] < prev["RSI6"]:
        score += 1
        risks.append("RSI6掉头向下")

    if latest["K"] < latest["D"] and prev["K"] >= prev["D"] and prev["K"] > 80:
        score += 2
        risks.append("KDJ高位死叉")

    result.risk_sources = risks if risks else ["短线反弹正常，暂无明显过热衰竭信号"]
    return score


def risk_level_from_score(score: int) -> str:
    """根据短线风险分数给出风险等级。"""

    if score >= 8:
        return "高位反弹衰竭风险"
    if score >= 6:
        return "短线过热，容易回踩"
    if score >= 4:
        return "短线偏热"
    return "短线反弹正常"


def build_observation(result: ShortRiskResult) -> str:
    """生成后续观察位置。"""

    return (
        f"当前价格 {fmt_price(result.close)}；MA5 {fmt_price(result.ma5)}；"
        f"MA10 {fmt_price(result.ma10)}；MA20 {fmt_price(result.ma20)}。"
        "若跌破MA5，短线反弹降温；若跌破MA20，反弹结构可能明显转弱。"
    )


def analyze_stock(ticker: str, target_day: dt.date) -> ShortRiskResult:
    """分析单只股票的短线反弹过热风险。"""

    result = ShortRiskResult(ticker=ticker)
    df, err = fetch_stock_data(ticker, target_day)
    if df.empty:
        result.error = err
        result.risk_level = "数据获取失败"
        return result
    if len(df) < 60:
        result.error = f"有效数据过短：{len(df)}个交易日"
        result.risk_level = "数据不足"
        return result

    if len(df) < 120:
        result.data_warning = f"历史数据偏短，当前仅{len(df)}个交易日。"

    df = calculate_indicators(df)
    latest = df.iloc[-1]
    result.date = latest.name.date() if hasattr(latest.name, "date") else None
    result.close = safe_float(latest["Close"])
    result.ma5 = safe_float(latest["MA5"])
    result.ma10 = safe_float(latest["MA10"])
    result.ma20 = safe_float(latest["MA20"])
    result.ma50 = safe_float(latest["MA50"])
    result.rsi6 = safe_float(latest["RSI6"])
    result.rsi12 = safe_float(latest["RSI12"])
    result.k = safe_float(latest["K"])
    result.d = safe_float(latest["D"])
    result.j = safe_float(latest["J"])
    result.deviation20 = safe_float(latest["Deviation20"])
    result.return5 = safe_float(latest["Return_5D"])
    result.return10 = safe_float(latest["Return_10D"])
    result.volume_ratio = safe_float(latest["Volume_Ratio"])

    result.rebound_score, result.rebound_state, result.rebound_reasons = is_short_rebound(df)
    if result.rebound_score < 4:
        result.short_risk_score = 0
        result.risk_level = "未形成短线反弹，不进入过热评分"
        result.risk_sources = ["未满足短线反弹过滤条件"]
    else:
        result.short_risk_score = score_short_rebound_risk(df, result)
        result.risk_level = risk_level_from_score(result.short_risk_score)

    result.indicator_summary = (
        f"RSI6/RSI12={fmt_num(result.rsi6, 1)}/{fmt_num(result.rsi12, 1)}；"
        f"K/D/J={fmt_num(result.k, 1)}/{fmt_num(result.d, 1)}/{fmt_num(result.j, 1)}；"
        f"偏离MA20={fmt_num(result.deviation20, 2)}%；"
        f"近5日涨幅={fmt_num(result.return5, 2)}%；"
        f"量能={fmt_num(result.volume_ratio, 2)}倍。"
    )
    result.observation = build_observation(result)
    return result


def report_date_from_results(results: List[ShortRiskResult]) -> dt.date:
    """使用最新有效交易日作为报告日期。"""

    dates = [r.date for r in results if r.date is not None]
    return max(dates) if dates else dt.date.today()


def parse_trigger_symbol(symbol_text: str) -> str:
    """从RISK.txt的股票字段中提取纯股票代码。"""

    symbol = symbol_text.strip().upper()
    if "(" in symbol:
        symbol = symbol.split("(", 1)[0].strip()
    if "（" in symbol:
        symbol = symbol.split("（", 1)[0].strip()
    return symbol


def append_risk_txt(results: List[ShortRiskResult], report_date: dt.date) -> Path:
    """分数大于6时写入RISK.txt，格式：日期 股票代码(次数) 得分。"""

    txt_path = OUTPUT_DIR / TRIGGER_TXT_NAME
    existing_keys = set()
    symbol_counts: Dict[str, int] = {}

    if txt_path.exists():
        try:
            for raw_line in txt_path.read_text(encoding="utf-8-sig").splitlines():
                parts = raw_line.strip().split()
                if len(parts) < 2:
                    continue
                symbol = parse_trigger_symbol(parts[1])
                if not symbol:
                    continue
                existing_keys.add(f"{parts[0]}\t{symbol}")
                symbol_counts[symbol] = symbol_counts.get(symbol, 0) + 1
        except Exception as exc:
            logging.warning("读取RISK.txt失败，将仅按本次结果写入：%s", exc)
            existing_keys = set()
            symbol_counts = {}

    date_text = report_date.strftime("%Y-%m-%d")
    triggered = [
        r for r in results
        if not r.error and r.date == report_date and r.short_risk_score > SHORT_RISK_TRIGGER_SCORE
    ]
    triggered.sort(key=lambda r: (-r.short_risk_score, r.ticker))

    new_lines = []
    for result in triggered:
        key = f"{date_text}\t{result.ticker}"
        if key in existing_keys:
            continue
        next_count = symbol_counts.get(result.ticker, 0) + 1
        new_lines.append(f"{date_text}\t{result.ticker}({next_count})\t{result.short_risk_score}分")
        existing_keys.add(key)
        symbol_counts[result.ticker] = next_count

    if new_lines:
        with txt_path.open("a", encoding="utf-8-sig") as file:
            for line in new_lines:
                file.write(line + "\n")
        logging.info("已写入RISK.txt：%s", "；".join(new_lines))
    else:
        logging.info("本交易日无新增短线风险触发记录。")
    return txt_path


def add_header(table, headers: List[str]) -> None:
    """设置Word表格表头。"""

    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def write_word_report(results: List[ShortRiskResult], report_date: dt.date) -> Path:
    """生成短线反弹风险Word报告。"""

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / f"{REPORT_PREFIX}_{report_date.strftime('%Y-%m-%d')}.docx"
    doc = Document()
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"].font.size = Pt(10)

    title = doc.add_heading("美股短线反弹风险监控报告", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(f"报告日期：{report_date.strftime('%Y-%m-%d')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sorted_results = sorted(results, key=lambda r: (r.short_risk_score, r.rebound_score), reverse=True)

    doc.add_heading("一、短线风险排名", level=2)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    add_header(table, ["股票", "短线状态", "反弹分", "风险分", "风险等级", "收盘价", "主要风险"])
    for r in sorted_results:
        cells = table.add_row().cells
        cells[0].text = r.ticker
        cells[1].text = r.rebound_state
        cells[2].text = str(r.rebound_score)
        cells[3].text = str(r.short_risk_score)
        cells[4].text = r.risk_level
        cells[5].text = fmt_price(r.close)
        cells[6].text = "；".join(r.risk_sources[:4])

    doc.add_heading("二、逐股票分析", level=2)
    for r in sorted_results:
        doc.add_heading(r.ticker, level=3)
        if r.error:
            doc.add_paragraph(f"数据异常：{r.error}")
            continue
        if r.data_warning:
            doc.add_paragraph(f"数据提示：{r.data_warning}")
        doc.add_paragraph(f"当前价格：{fmt_price(r.close)}")
        doc.add_paragraph(f"短线状态：{r.rebound_state}；反弹确认分：{r.rebound_score}")
        doc.add_paragraph(f"短线风险分：{r.short_risk_score}；风险等级：{r.risk_level}")
        doc.add_paragraph(f"指标摘要：{r.indicator_summary}")

        doc.add_paragraph("反弹成立条件：")
        if r.rebound_reasons:
            for reason in r.rebound_reasons:
                doc.add_paragraph(reason, style="List Bullet")
        else:
            doc.add_paragraph("暂无", style="List Bullet")

        doc.add_paragraph("主要风险来源：")
        for risk in r.risk_sources:
            doc.add_paragraph(risk, style="List Bullet")
        doc.add_paragraph(f"建议观察位置：{r.observation}")

    doc.add_paragraph("本报告仅用于短线技术状态监控，不构成任何投资建议。")
    doc.save(path)
    return path


def run_monitor() -> Tuple[Path, Path, List[ShortRiskResult]]:
    """执行完整短线风险监控流程。"""

    setup_logging()
    target_day = latest_closed_us_trading_day()
    logging.info("开始短线反弹风险监控，共%s只股票；目标交易日：%s。", len(TICKERS), target_day)
    results: List[ShortRiskResult] = []
    for idx, ticker in enumerate(TICKERS, start=1):
        logging.info("[%s/%s] 分析 %s", idx, len(TICKERS), ticker)
        try:
            result = analyze_stock(ticker, target_day)
        except Exception as exc:
            logging.exception("%s 分析失败：%s", ticker, exc)
            result = ShortRiskResult(ticker=ticker, error=str(exc), risk_level="分析失败")
        results.append(result)
        time.sleep(REQUEST_PAUSE_SECONDS)

    report_date = report_date_from_results(results)
    word_path = write_word_report(results, report_date)
    txt_path = append_risk_txt(results, report_date)
    logging.info("监控完成：Word=%s TXT=%s", word_path, txt_path)
    return word_path, txt_path, results


def main() -> None:
    """程序入口，供Spyder、命令行和计划任务直接运行。"""

    word_path, txt_path, _ = run_monitor()
    print("")
    print("美股短线反弹风险监控完成：")
    print(f"Word报告：{word_path}")
    print(f"触发记录：{txt_path}")


if __name__ == "__main__":
    main()

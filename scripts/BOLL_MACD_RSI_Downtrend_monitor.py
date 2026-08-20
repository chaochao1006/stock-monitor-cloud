"""
美股 BOLL + MACD + RSI 长期下跌风险自动监控程序

运行逻辑：
1. 使用 yfinance 获取股票最近日线数据。
2. 只使用最近一个已经结束的美股交易日数据。
3. 计算 BOLL、MACD、RSI、均线和成交量指标。
4. 出现 A/B/C 长期下跌风险信号时才生成 Word 报告。
5. 不生成 CSV，只保留 Word 报告和运行日志。

本程序仅用于技术指标监控，不构成投资建议。
"""

from __future__ import annotations

import logging
import math
import os
import sys
import time
from dataclasses import dataclass, field
from datetime import date, datetime, time as dtime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from zoneinfo import ZoneInfo

import numpy as np
import pandas as pd
import requests
import yfinance as yf
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.chart import LineChart, Reference
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


# =========================
# 用户配置
# =========================

STOCK_POOL = [
    "NOK", "XE", "RKLB", "NBIS", "ARM", "OPTX", "GLW", "COHR", "LITE",
    "MRVL", "INTC", "ALAB", "FLNC", "CRWV", "ORCL", "NVTS", "POET",
    "DELL", "BB", "ONDS", "SMR", "OKLO", "IONQ", "QBTS", "NVDA",
    "AMZN", "AMD", "CRCL", "FLY", "CBRS", "CIEN", "SIVEF", "SHAZ",
    "LUNR", "ASTS", "AAOI", "INOD", "SPCX", "RDW", "QCOM", "COIN",
    "CORZ", "IREN", "VELO", "MSFT", "GOOG", "AVGO", "AMAT", "AMKR",
    "LRCX", "SNDK", "MU", "VICR", "NNE", "CCJ", "VST", "BWAY",
    "MANE", "QSI", "ERAS", "CRSP", "IBRX", "INSP",
]

REPORT_BASE_DIR = Path(os.environ.get("REPORT_BASE_DIR", Path(__file__).resolve().parents[1] / "data"))
OUTPUT_DIR = REPORT_BASE_DIR / "drop" / "drop-BOLL"
LOG_PATH = OUTPUT_DIR / "BOLL_MACD_RSI_Downtrend_monitor.log"
EXCEL_TRACKING_DIR = OUTPUT_DIR / "EXCLE"
TRIGGER_TXT_PATH = OUTPUT_DIR / "BOLL-dip.txt"

FETCH_PERIOD = "220d"
MIN_ROWS_REQUIRED = 100
TRACKING_DAYS = 30
REPORT_PREFIX = "BOLL_MACD_RSI_Downtrend_Report"
GRADE_RANK = {"A. 强信号": 3, "B. 中等信号": 2, "C. 观察信号": 1, "无信号": 0}
FRESH_BREAK_MAX_DAYS = 8
MAX_DAYS_BELOW_MID_FOR_FRESH_SIGNAL = 8
MAX_BELOW_MID_DEVIATION_PCT = -10.0
START_SIGNAL_LOOKBACK_DAYS = 8
RECENT_HIGH_LOOKBACK_DAYS = 40
MAX_HIGH_TO_BREAK_DAYS = 20
MIN_HIGH_ABOVE_BREAK_MID_PCT = 8.0
MIN_PREBREAK_DAYS_ABOVE_MID = 3
MIN_START_EVENT_ROWS = 40


# =========================
# 数据结构
# =========================

@dataclass
class RiskResult:
    """保存单只股票的指标、评分和风险解释。"""

    ticker: str
    date: Optional[date] = None
    close: Optional[float] = None
    score: int = 0
    risk_level: str = "无信号"
    signal_grade: str = "无信号"
    risk_sources: List[str] = field(default_factory=list)
    detail: Dict[str, float] = field(default_factory=dict)
    macd_dead_cross: bool = False
    rsi12_below_50: bool = False
    error: str = ""


# =========================
# 基础工具
# =========================

def setup_logging() -> None:
    """初始化日志，写入文件并同时输出到控制台。"""

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[
            logging.FileHandler(LOG_PATH, encoding="utf-8"),
            logging.StreamHandler(sys.stdout),
        ],
    )


def safe_float(value) -> Optional[float]:
    """安全转换为 float，遇到空值或无穷值时返回 None。"""

    try:
        if value is None or pd.isna(value):
            return None
        value = float(value)
        if not math.isfinite(value):
            return None
        return value
    except Exception:
        return None


def fmt_num(value, digits: int = 2) -> str:
    """格式化普通数值。"""

    value = safe_float(value)
    if value is None:
        return "N/A"
    return f"{value:.{digits}f}"


def fmt_pct(value, digits: int = 2) -> str:
    """格式化百分比，输入按百分数口径，例如 12.3 表示 12.3%。"""

    value = safe_float(value)
    if value is None:
        return "N/A"
    return f"{value:.{digits}f}%"


def latest_complete_us_date() -> date:
    """按纽约时间判断最近一个可能已经结束的美股交易日日期。"""

    now_ny = datetime.now(ZoneInfo("America/New_York"))
    if now_ny.time() >= dtime(16, 15):
        candidate = now_ny.date()
    else:
        candidate = now_ny.date() - timedelta(days=1)
    while candidate.weekday() >= 5:
        candidate -= timedelta(days=1)
    return candidate


# =========================
# 数据获取
# =========================

def normalize_yfinance_columns(df: pd.DataFrame) -> pd.DataFrame:
    """兼容 yfinance 单票和多层列格式。"""

    if isinstance(df.columns, pd.MultiIndex):
        df.columns = [col[0] if isinstance(col, tuple) else col for col in df.columns]
    return df


def fetch_ohlcv(ticker: str) -> Tuple[pd.DataFrame, str]:
    """获取单只股票日线数据，并过滤到最近一个已结束交易日。"""

    last_error = ""
    for attempt in range(1, 2):
        try:
            df = yf.download(
                ticker,
                period=FETCH_PERIOD,
                interval="1d",
                auto_adjust=False,
                progress=False,
                threads=False,
            )
            df = normalize_yfinance_columns(df)
            if df.empty:
                raise ValueError("yfinance 返回空数据")

            required = ["Open", "High", "Low", "Close", "Volume"]
            missing = [col for col in required if col not in df.columns]
            if missing:
                raise ValueError(f"缺少字段：{missing}")

            df = df[required].copy()
            df.index = pd.to_datetime(df.index).tz_localize(None)
            for col in required:
                df[col] = pd.to_numeric(df[col], errors="coerce")
            df = df.dropna(subset=["Open", "High", "Low", "Close", "Volume"])

            cutoff = latest_complete_us_date()
            df = df[df.index.date <= cutoff]
            df = df.tail(120)
            if len(df) < MIN_ROWS_REQUIRED:
                raise ValueError(f"有效日线不足：{len(df)} 行")
            return df, ""
        except Exception as exc:
            last_error = f"第 {attempt} 次获取失败：{exc}"
            logging.warning("%s %s", ticker, last_error)

    logging.warning("%s yfinance 获取失败，尝试 Yahoo Chart API 备用源。", ticker)
    fallback_df, fallback_error = fetch_ohlcv_yahoo_chart(ticker)
    if not fallback_df.empty:
        return fallback_df, ""
    return pd.DataFrame(), f"{last_error}；备用源失败：{fallback_error}"


def fetch_ohlcv_yahoo_chart(ticker: str) -> Tuple[pd.DataFrame, str]:
    """yfinance 限流时使用 Yahoo Chart API 作为备用日线数据源。"""

    end_dt = datetime.now(ZoneInfo("America/New_York")) + timedelta(days=1)
    start_dt = end_dt - timedelta(days=260)
    url = f"https://query1.finance.yahoo.com/v8/finance/chart/{ticker}"
    params = {
        "period1": int(start_dt.timestamp()),
        "period2": int(end_dt.timestamp()),
        "interval": "1d",
        "events": "history",
        "includeAdjustedClose": "true",
    }
    headers = {"User-Agent": "Mozilla/5.0"}
    last_error = ""

    for attempt in range(1, 4):
        try:
            response = requests.get(url, params=params, headers=headers, timeout=20)
            response.raise_for_status()
            payload = response.json()
            chart_error = payload.get("chart", {}).get("error")
            if chart_error:
                raise ValueError(chart_error)
            result = payload["chart"]["result"][0]
            timestamps = result.get("timestamp") or []
            quote = result["indicators"]["quote"][0]
            rows = []
            for idx, ts in enumerate(timestamps):
                rows.append(
                    {
                        "Date": datetime.fromtimestamp(ts, ZoneInfo("UTC")).replace(tzinfo=None),
                        "Open": quote["open"][idx],
                        "High": quote["high"][idx],
                        "Low": quote["low"][idx],
                        "Close": quote["close"][idx],
                        "Volume": quote["volume"][idx],
                    }
                )
            df = pd.DataFrame(rows)
            if df.empty:
                raise ValueError("Yahoo Chart API 返回空数据")
            df["Date"] = pd.to_datetime(df["Date"]).dt.normalize()
            df = df.set_index("Date").sort_index()
            for col in ["Open", "High", "Low", "Close", "Volume"]:
                df[col] = pd.to_numeric(df[col], errors="coerce")
            df = df.dropna(subset=["Open", "High", "Low", "Close", "Volume"])
            cutoff = latest_complete_us_date()
            df = df[df.index.date <= cutoff].tail(120)
            if len(df) < MIN_ROWS_REQUIRED:
                raise ValueError(f"备用源有效日线不足：{len(df)} 行")
            return df[["Open", "High", "Low", "Close", "Volume"]], ""
        except Exception as exc:
            last_error = f"第 {attempt} 次备用获取失败：{exc}"
            logging.warning("%s %s", ticker, last_error)
            time.sleep(2)

    return pd.DataFrame(), last_error


# =========================
# 指标计算
# =========================

def calc_rsi(close: pd.Series, period: int) -> pd.Series:
    """使用 Wilder 平滑方法计算 RSI。"""

    delta = close.diff()
    gain = delta.clip(lower=0)
    loss = -delta.clip(upper=0)
    avg_gain = gain.ewm(alpha=1 / period, min_periods=period, adjust=False).mean()
    avg_loss = loss.ewm(alpha=1 / period, min_periods=period, adjust=False).mean()
    rs = avg_gain / avg_loss.replace(0, np.nan)
    rsi = 100 - (100 / (1 + rs))
    return rsi.fillna(100)


def add_indicators(df: pd.DataFrame) -> pd.DataFrame:
    """计算 BOLL、MACD、RSI、均线和成交量指标。"""

    df = df.copy()
    close = df["Close"]

    for period in [5, 10, 20, 50, 200]:
        df[f"MA{period}"] = close.rolling(period).mean()

    df["BOLL_MID"] = df["MA20"]
    df["BOLL_STD20"] = close.rolling(20).std(ddof=0)
    df["BOLL_UPPER"] = df["BOLL_MID"] + 2 * df["BOLL_STD20"]
    df["BOLL_LOWER"] = df["BOLL_MID"] - 2 * df["BOLL_STD20"]
    df["Bandwidth"] = (df["BOLL_UPPER"] - df["BOLL_LOWER"]) / df["BOLL_MID"] * 100
    df["PercentB"] = (close - df["BOLL_LOWER"]) / (df["BOLL_UPPER"] - df["BOLL_LOWER"])

    ema12 = close.ewm(span=12, adjust=False).mean()
    ema26 = close.ewm(span=26, adjust=False).mean()
    df["DIF"] = ema12 - ema26
    df["DEA"] = df["DIF"].ewm(span=9, adjust=False).mean()
    df["MACD_Hist"] = df["DIF"] - df["DEA"]

    df["RSI6"] = calc_rsi(close, 6)
    df["RSI12"] = calc_rsi(close, 12)
    df["RSI24"] = calc_rsi(close, 24)

    df["Volume_MA20"] = df["Volume"].rolling(20).mean()
    df["Daily_Return_Pct"] = close.pct_change() * 100
    return df


# =========================
# 风险评分
# =========================

def risk_level(score: int) -> str:
    """根据总分输出风险等级。"""

    if score <= 5:
        return "正常"
    if score <= 10:
        return "趋势转弱"
    if score <= 15:
        return "下降趋势形成"
    if score <= 20:
        return "强下降风险"
    return "长期下降风险"


def has_macd_bearish_divergence(df: pd.DataFrame) -> bool:
    """判断过去60交易日是否出现价格新高但 MACD 未创新高。"""

    window = df.tail(60)
    if len(window) < 60:
        return False
    latest = window.iloc[-1]
    previous = window.iloc[:-1]
    if previous.empty:
        return False
    price_new_high = latest["Close"] >= previous["Close"].max()
    macd_not_new_high = latest["DIF"] < previous["DIF"].max()
    return bool(price_new_high and macd_not_new_high)


def consecutive_down_days(series: pd.Series) -> int:
    """统计序列末端连续下跌的交易日数量。"""

    values = series.dropna().tolist()
    count = 0
    for idx in range(len(values) - 1, 0, -1):
        if values[idx] < values[idx - 1]:
            count += 1
        else:
            break
    return count


def recent_macd_dead_cross(df: pd.DataFrame, lookback_days: int = 6) -> bool:
    """判断最近若干交易日内是否出现 MACD 死叉。"""

    tail = df.tail(lookback_days + 1)
    for idx in range(1, len(tail)):
        today = tail.iloc[idx]
        yesterday = tail.iloc[idx - 1]
        if today["DIF"] < today["DEA"] and yesterday["DIF"] >= yesterday["DEA"]:
            return True
    return False


def macd_dead_cross_around_break(df: pd.DataFrame, break_pos: Optional[int], window_days: int = 6) -> bool:
    """判断 MACD 死叉是否发生在跌破中轨日前后指定交易日窗口内。"""

    if break_pos is None:
        return False
    start_pos = max(1, break_pos - window_days)
    end_pos = min(len(df) - 1, break_pos + window_days)
    for pos in range(start_pos, end_pos + 1):
        today = df.iloc[pos]
        yesterday = df.iloc[pos - 1]
        if today["DIF"] < today["DEA"] and yesterday["DIF"] >= yesterday["DEA"]:
            return True
    return False


def days_since_last_mid_break(df: pd.DataFrame, lookback_days: int = 20) -> Optional[int]:
    """返回最近一次跌破 BOLL 中轨距今的交易日数，未出现则返回 None。"""

    tail = df.tail(lookback_days + 1)
    for idx in range(len(tail) - 1, 0, -1):
        today = tail.iloc[idx]
        yesterday = tail.iloc[idx - 1]
        if today["Close"] < today["BOLL_MID"] and yesterday["Close"] >= yesterday["BOLL_MID"]:
            return len(tail) - idx - 1
    return None


def consecutive_days_below_mid(df: pd.DataFrame) -> int:
    """统计末端连续收盘位于 BOLL 中轨下方的交易日数量。"""

    count = 0
    for _, row in df.dropna(subset=["Close", "BOLL_MID"]).iloc[::-1].iterrows():
        if row["Close"] < row["BOLL_MID"]:
            count += 1
        else:
            break
    return count


def fresh_downtrend_filter(df: pd.DataFrame, break_pos: Optional[int] = None) -> Tuple[bool, Dict[str, float]]:
    """过滤已经明显下跌一段时间的形态，只保留刚跌破中轨后的早期转弱。"""

    latest = df.iloc[-1]
    days_since_break = len(df) - break_pos - 1 if break_pos is not None else days_since_last_mid_break(df, 20)
    below_mid_days = consecutive_days_below_mid(df)
    mid_deviation_pct = (latest["Close"] - latest["BOLL_MID"]) / latest["BOLL_MID"] * 100

    is_fresh = (
        days_since_break is not None
        and days_since_break <= FRESH_BREAK_MAX_DAYS
        and below_mid_days <= MAX_DAYS_BELOW_MID_FOR_FRESH_SIGNAL
        and mid_deviation_pct >= MAX_BELOW_MID_DEVIATION_PCT
    )
    metrics = {
        "Days_Since_Mid_Break": float(days_since_break) if days_since_break is not None else np.nan,
        "Consecutive_Days_Below_Mid": float(below_mid_days),
        "Mid_Deviation_Pct": float(mid_deviation_pct),
        "Fresh_Downtrend_Filter": float(is_fresh),
    }
    return bool(is_fresh), metrics


def recent_high_mid_break_event(df: pd.DataFrame) -> Tuple[bool, Dict[str, float], Optional[int]]:
    """判断最近是否出现“阶段高点之后收盘跌破 BOLL 中轨”的启动事件。"""

    metrics = {
        "Start_Break_Pos": np.nan,
        "Start_Days_Since_Mid_Break": np.nan,
        "Start_High_To_Break_Days": np.nan,
        "Start_High_Above_Break_Mid_Pct": np.nan,
        "Start_Prebreak_Days_Above_Mid": np.nan,
        "Start_From_Recent_High": 0.0,
    }
    if len(df) < MIN_START_EVENT_ROWS:
        return False, metrics, None

    break_pos = None
    first_pos = max(1, len(df) - START_SIGNAL_LOOKBACK_DAYS - 1)
    for pos in range(len(df) - 1, first_pos - 1, -1):
        today = df.iloc[pos]
        yesterday = df.iloc[pos - 1]
        if today["Close"] < today["BOLL_MID"] and yesterday["Close"] >= yesterday["BOLL_MID"]:
            break_pos = pos
            break
    if break_pos is None:
        return False, metrics, None

    break_row = df.iloc[break_pos]
    high_start = max(0, break_pos - RECENT_HIGH_LOOKBACK_DAYS)
    high_window = df.iloc[high_start : break_pos + 1]
    if high_window.empty:
        return False, metrics, None

    high_relative_pos = int(high_window["High"].values.argmax())
    high_pos = high_start + high_relative_pos
    high_value = float(df.iloc[high_pos]["High"])
    days_from_high_to_break = break_pos - high_pos
    high_above_break_mid_pct = (high_value - break_row["BOLL_MID"]) / break_row["BOLL_MID"] * 100

    prebreak = df.iloc[max(0, break_pos - 10) : break_pos]
    prebreak_days_above_mid = int((prebreak["Close"] > prebreak["BOLL_MID"]).sum())
    prebreak_touched_upper_area = bool(
        ((prebreak["High"] >= prebreak["BOLL_UPPER"] * 0.98) | (prebreak["PercentB"] >= 0.75)).any()
    )

    is_from_recent_high = (
        0 <= days_from_high_to_break <= MAX_HIGH_TO_BREAK_DAYS
        and high_above_break_mid_pct >= MIN_HIGH_ABOVE_BREAK_MID_PCT
        and prebreak_days_above_mid >= MIN_PREBREAK_DAYS_ABOVE_MID
        and (prebreak_touched_upper_area or high_above_break_mid_pct >= MIN_HIGH_ABOVE_BREAK_MID_PCT * 1.5)
    )

    metrics.update(
        {
            "Start_Days_Since_Mid_Break": float(len(df) - break_pos - 1),
            "Start_Break_Pos": float(break_pos),
            "Start_High_To_Break_Days": float(days_from_high_to_break),
            "Start_High_Above_Break_Mid_Pct": float(high_above_break_mid_pct),
            "Start_Prebreak_Days_Above_Mid": float(prebreak_days_above_mid),
            "Start_From_Recent_High": float(is_from_recent_high),
        }
    )
    return bool(is_from_recent_high), metrics, break_pos


def price_broke_mid_from_recent_high(df: pd.DataFrame) -> bool:
    """兼容旧调用：判断股价是否从近期高位跌破 BOLL 中轨。"""

    signal, _, _ = recent_high_mid_break_event(df)
    return signal


def failed_to_reclaim_mid_after_break(df: pd.DataFrame, break_pos: Optional[int] = None) -> bool:
    """判断跌破中轨后，后续3-5个交易日未能重新站上中轨。"""

    if break_pos is None:
        tail = df.tail(12)
        break_positions = []
        for idx in range(1, len(tail)):
            today = tail.iloc[idx]
            yesterday = tail.iloc[idx - 1]
            if today["Close"] < today["BOLL_MID"] and yesterday["Close"] >= yesterday["BOLL_MID"]:
                break_positions.append(idx)
        if not break_positions:
            return False
        break_idx = break_positions[-1]
        days_after_break = len(tail) - break_idx - 1
        after_break = tail.iloc[break_idx + 1 :]
    else:
        days_after_break = len(df) - break_pos - 1
        after_break = df.iloc[break_pos + 1 :]

    if days_after_break < 3 or days_after_break > 5:
        return False
    return bool((after_break["Close"] <= after_break["BOLL_MID"]).all())


def boll_opening_down_pressure(df: pd.DataFrame) -> bool:
    """判断近3-5日布林带张口：上轨向上、下轨向下、中轨向下。"""

    for days in range(3, 6):
        if len(df) <= days:
            continue
        latest = df.iloc[-1]
        previous = df.iloc[-1 - days]
        if (
            latest["BOLL_UPPER"] > previous["BOLL_UPPER"]
            and latest["BOLL_LOWER"] < previous["BOLL_LOWER"]
            and latest["BOLL_MID"] < previous["BOLL_MID"]
            and latest["Bandwidth"] > previous["Bandwidth"]
        ):
            return True
    return False


def boll_three_bands_down(df: pd.DataFrame) -> bool:
    """判断近3-5日上轨、中轨、下轨三个轨道同时向下。"""

    for days in range(3, 6):
        if len(df) <= days:
            continue
        latest = df.iloc[-1]
        previous = df.iloc[-1 - days]
        if (
            latest["BOLL_UPPER"] < previous["BOLL_UPPER"]
            and latest["BOLL_MID"] < previous["BOLL_MID"]
            and latest["BOLL_LOWER"] < previous["BOLL_LOWER"]
        ):
            return True
    return False


def score_ticker(ticker: str, df: pd.DataFrame) -> RiskResult:
    """按 A/B/C 规则判断单只股票的长期下跌风险信号。"""

    df = add_indicators(df).dropna(subset=["BOLL_MID", "BOLL_UPPER", "BOLL_LOWER", "DIF", "DEA", "RSI24"])
    if len(df) < 20:
        return RiskResult(ticker=ticker, error="指标有效数据不足")

    latest = df.iloc[-1]
    result = RiskResult(
        ticker=ticker,
        date=df.index[-1].date(),
        close=float(latest["Close"]),
    )

    condition1, start_metrics, break_pos = recent_high_mid_break_event(df)
    condition2 = failed_to_reclaim_mid_after_break(df, break_pos)
    condition3 = macd_dead_cross_around_break(df, break_pos, 6)
    condition4 = consecutive_down_days(df["BOLL_MID"]) >= 3
    condition5 = boll_opening_down_pressure(df)
    condition6 = boll_three_bands_down(df)
    fresh_filter, fresh_metrics = fresh_downtrend_filter(df, break_pos)
    start_signal = condition1 and condition3

    if condition1:
        result.risk_sources.append("1. 启动前提：股价从近期高位跌破BOLL中轨")
    if condition2:
        result.risk_sources.append("2. 跌破中轨后3-5个交易日未能重新站上中轨")
    if condition3:
        result.risk_sources.append("3. 跌破中轨日前后6个交易日内MACD出现死叉")
    if condition4:
        result.risk_sources.append("4. BOLL中轨连续3日及以上下跌")
    if condition5:
        result.risk_sources.append("5. 近3-5日BOLL带张口，且中轨向下")
    if condition6:
        result.risk_sources.append("6. 近3-5日上轨、中轨、下轨同时向下")

    if start_signal and (condition5 or condition6) and fresh_filter:
        result.signal_grade = "A. 强信号"
        result.risk_sources.append("等级前提：近期高位跌破中轨，并且跌破日前后6个交易日内MACD死叉")
        result.risk_sources.append("早期过滤：仍处于跌破中轨后的早期确认阶段")
    elif start_signal and (condition2 or condition4) and fresh_filter:
        result.signal_grade = "B. 中等信号"
        result.risk_sources.append("等级前提：近期高位跌破中轨，并且跌破日前后6个交易日内MACD死叉")
        result.risk_sources.append("早期过滤：仍处于跌破中轨后的早期确认阶段")
    elif start_signal:
        result.signal_grade = "C. 观察信号"
    else:
        result.signal_grade = "无信号"

    result.risk_level = result.signal_grade
    result.score = GRADE_RANK.get(result.signal_grade, 0)
    result.macd_dead_cross = bool(condition3)
    result.rsi12_below_50 = bool(latest["RSI12"] < 50)
    result.detail = {
        "BOLL_UPPER": latest["BOLL_UPPER"],
        "BOLL_MID": latest["BOLL_MID"],
        "BOLL_LOWER": latest["BOLL_LOWER"],
        "Bandwidth": latest["Bandwidth"],
        "DIF": latest["DIF"],
        "DEA": latest["DEA"],
        "MACD_Hist": latest["MACD_Hist"],
        "RSI6": latest["RSI6"],
        "RSI12": latest["RSI12"],
        "RSI24": latest["RSI24"],
        "MA20": latest["MA20"],
        "MA50": latest["MA50"],
        "MA200": latest["MA200"],
        "Volume": latest["Volume"],
        "Volume_MA20": latest["Volume_MA20"],
        "Volume_Ratio": latest["Volume"] / latest["Volume_MA20"] if latest["Volume_MA20"] else np.nan,
        "Daily_Return_Pct": latest["Daily_Return_Pct"],
        "Condition1": float(condition1),
        "Condition2": float(condition2),
        "Condition3": float(condition3),
        "Condition4": float(condition4),
        "Condition5": float(condition5),
        "Condition6": float(condition6),
        "Start_Signal": float(start_signal),
        **fresh_metrics,
        **start_metrics,
    }
    return result


# =========================
# Word报告
# =========================

def set_default_font(doc: Document) -> None:
    """设置 Word 默认中文字体大小。"""

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10)


def add_table(doc: Document, headers: List[str], rows: List[List[str]]) -> None:
    """向 Word 添加简洁表格。"""

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def build_summary_text(result: RiskResult) -> str:
    """生成单只股票自动风险总结。"""

    risks = "；".join(result.risk_sources[:6]) if result.risk_sources else "暂无明显触发条件"
    if result.signal_grade.startswith("A"):
        stage = "当前属于强信号，需要重点观察下跌趋势是否加速"
    elif result.signal_grade.startswith("B"):
        stage = "当前属于中等信号，说明下跌结构正在形成或延续"
    elif result.signal_grade.startswith("C"):
        stage = "当前属于观察信号，需要后续交易日确认是否继续走弱"
    else:
        stage = "当前未触发明确下跌风险信号"
    ma50 = fmt_num(result.detail.get("MA50"))
    return f"{result.ticker}：{result.signal_grade}。触发条件：{risks}。{stage}，同时关注 MA50 附近 {ma50} 的支撑或压力变化。"


def write_word_report(triggered: List[RiskResult], report_date: date) -> Path:
    """只为触发股票生成 Word 风险报告。"""

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_default_font(doc)

    doc.add_heading("美股长期下跌风险监控报告", level=0)
    doc.add_paragraph(f"报告日期：{report_date.strftime('%Y-%m-%d')}")
    doc.add_paragraph("说明：本报告仅在出现 A/B/C 下跌风险信号时生成，仅用于技术指标监控，不构成投资建议。")

    doc.add_heading("一、风险股票排名", level=1)
    ranking_rows = []
    for result in triggered:
        ranking_rows.append([
            result.ticker,
            result.signal_grade,
            fmt_num(result.close),
            "；".join(result.risk_sources[:4]),
        ])
    add_table(doc, ["股票", "信号等级", "当前价格", "触发条件"], ranking_rows)

    doc.add_heading("二、详细分析", level=1)
    for result in triggered:
        d = result.detail
        doc.add_heading(result.ticker, level=2)
        doc.add_paragraph(f"基础信息：日期 {result.date}；当前价格 {fmt_num(result.close)}；信号等级 {result.signal_grade}。")
        doc.add_paragraph(
            f"BOLL：上轨 {fmt_num(d.get('BOLL_UPPER'))}；中轨 {fmt_num(d.get('BOLL_MID'))}；"
            f"下轨 {fmt_num(d.get('BOLL_LOWER'))}；Bandwidth {fmt_pct(d.get('Bandwidth'))}。"
        )
        doc.add_paragraph(
            f"MACD：DIF {fmt_num(d.get('DIF'), 4)}；DEA {fmt_num(d.get('DEA'), 4)}；"
            f"MACD柱 {fmt_num(d.get('MACD_Hist'), 4)}；是否死叉：{'是' if result.macd_dead_cross else '否'}。"
        )
        doc.add_paragraph(
            f"RSI：RSI6 {fmt_num(d.get('RSI6'))}；RSI12 {fmt_num(d.get('RSI12'))}；"
            f"RSI24 {fmt_num(d.get('RSI24'))}。"
        )
        doc.add_paragraph(
            f"均线：MA20 {fmt_num(d.get('MA20'))}；MA50 {fmt_num(d.get('MA50'))}；MA200 {fmt_num(d.get('MA200'))}。"
        )
        doc.add_paragraph(
            f"成交量：当日成交量 {fmt_num(d.get('Volume'), 0)}；20日平均成交量 {fmt_num(d.get('Volume_MA20'), 0)}；"
            f"成交量倍数 {fmt_num(d.get('Volume_Ratio'))}。"
        )
        doc.add_paragraph(f"触发条件：{'；'.join(result.risk_sources) if result.risk_sources else '无'}。")

    doc.add_heading("三、自动生成风险总结", level=1)
    for result in triggered:
        doc.add_paragraph(build_summary_text(result))

    output_path = OUTPUT_DIR / f"{REPORT_PREFIX}_{report_date.strftime('%Y-%m-%d')}.docx"
    doc.save(output_path)
    return output_path


# =========================
# 触发记录与Excel跟踪
# =========================

def parse_plain_symbol(symbol_text: str) -> str:
    """从 BOLL-dip.txt 的股票文本中提取纯股票代码。"""

    symbol = symbol_text.strip().upper()
    if "（" in symbol:
        symbol = symbol.split("（", 1)[0].strip()
    elif "(" in symbol:
        symbol = symbol.split("(", 1)[0].strip()
    return symbol


def append_triggered_txt(triggered: List[RiskResult], report_date: date) -> None:
    """把本次触发股票写入 BOLL-dip.txt，并标记累计触发次数。"""

    if not triggered:
        return

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    existing_keys = set()
    symbol_counts: Dict[str, int] = {}

    def parse_symbol_from_line(line: str) -> str:
        parts = line.strip().split()
        if len(parts) < 2:
            return ""
        return parse_plain_symbol(parts[1])

    def parse_date_symbol_key(line: str) -> str:
        parts = line.strip().split()
        if len(parts) < 2:
            return line.strip()
        symbol = parse_symbol_from_line(line)
        return f"{parts[0]}\t{symbol}"

    if TRIGGER_TXT_PATH.exists():
        try:
            for raw_line in TRIGGER_TXT_PATH.read_text(encoding="utf-8-sig").splitlines():
                line = raw_line.strip()
                if not line:
                    continue
                existing_keys.add(parse_date_symbol_key(line))
                symbol = parse_symbol_from_line(line)
                if symbol:
                    symbol_counts[symbol] = symbol_counts.get(symbol, 0) + 1
        except Exception as exc:
            logging.warning("读取 BOLL-dip.txt 失败，将重新统计本次触发：%s", exc)
            existing_keys = set()
            symbol_counts = {}

    date_text = report_date.strftime("%Y-%m-%d")
    new_lines = []
    for result in triggered:
        key = f"{date_text}\t{result.ticker}"
        if key in existing_keys:
            continue
        next_count = symbol_counts.get(result.ticker, 0) + 1
        symbol_text = f"{result.ticker}（{next_count}）" if next_count > 1 else result.ticker
        new_lines.append(f"{date_text}\t{symbol_text}\t{result.signal_grade}")
        existing_keys.add(key)
        symbol_counts[result.ticker] = next_count

    if new_lines:
        with TRIGGER_TXT_PATH.open("a", encoding="utf-8-sig") as file:
            for line in new_lines:
                file.write(line + "\n")
        logging.info("已写入 BOLL-dip.txt：%s", "；".join(new_lines))


def read_dip_trigger_history() -> Dict[str, List[date]]:
    """读取 BOLL-dip.txt，返回每只股票的历史触发日期。"""

    history: Dict[str, List[date]] = {}
    if not TRIGGER_TXT_PATH.exists():
        return history
    try:
        lines = TRIGGER_TXT_PATH.read_text(encoding="utf-8-sig").splitlines()
    except Exception as exc:
        logging.warning("读取 BOLL-dip.txt 失败，无法更新 Excel 跟踪文件：%s", exc)
        return history

    for raw_line in lines:
        parts = raw_line.strip().split()
        if len(parts) < 2:
            continue
        try:
            trigger_date = date.fromisoformat(parts[0])
        except ValueError:
            continue
        symbol = parse_plain_symbol(parts[1])
        if not symbol:
            continue
        history.setdefault(symbol, []).append(trigger_date)

    for symbol in list(history):
        history[symbol] = sorted(set(history[symbol]))
    return history


def safe_tracking_filename(symbol: str) -> str:
    """生成安全的 Excel 文件名主体。"""

    safe = "".join(ch if ch.isalnum() or ch in "._-" else "_" for ch in symbol.upper()).strip("._-")
    return safe or "UNKNOWN"


def tracking_workbook_path(tracking_dir: Path, symbol: str, episode_number: int, trigger_date: date) -> Path:
    """生成触发周期的 Excel 文件路径。"""

    base_name = safe_tracking_filename(symbol)
    dated_path = tracking_dir / f"{base_name}_{episode_number}_{trigger_date.strftime('%Y%m%d')}.xlsx"
    legacy_path = tracking_dir / f"{base_name}.xlsx" if episode_number == 1 else tracking_dir / f"{base_name}_{episode_number}.xlsx"
    if dated_path.exists():
        return dated_path
    if legacy_path.exists():
        return legacy_path
    return dated_path


def tracking_prices_from_df(df: pd.DataFrame, trigger_date: date) -> pd.DataFrame:
    """从完整行情中截取触发日起最多30个交易日的收盘价。"""

    df = df.copy()
    df["TrackClose"] = pd.to_numeric(df["Close"], errors="coerce")
    df["TrackDailyChangePct"] = df["TrackClose"].pct_change() * 100
    df = df[df.index.date >= trigger_date].copy()
    if not df.empty:
        df = df.head(TRACKING_DAYS)
        df = df.dropna(subset=["TrackClose"])
    return df


def is_trigger_inside_episode(df: pd.DataFrame, episode_start: date, trigger_date: date) -> bool:
    """判断再次触发是否仍处在上一轮30交易日跟踪周期内。"""

    episode_window = df[(df.index.date >= episode_start) & (df.index.date <= trigger_date)]
    if episode_window.empty:
        return True
    return len(episode_window) <= TRACKING_DAYS


def build_tracking_episodes(df: pd.DataFrame, trigger_dates: List[date]) -> List[date]:
    """把多次触发日期合并成若干个30交易日跟踪周期。"""

    episodes: List[date] = []
    for trigger_date in sorted(set(trigger_dates)):
        if not episodes:
            episodes.append(trigger_date)
            continue
        if not is_trigger_inside_episode(df, episodes[-1], trigger_date):
            episodes.append(trigger_date)
    return episodes


def write_tracking_workbook(path: Path, symbol: str, trigger_date: date, df: pd.DataFrame) -> None:
    """写入单只股票触发后30个交易日涨跌幅跟踪 Excel。"""

    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "30日跟踪"
    ws.freeze_panes = "A8"

    ws.merge_cells("A1:D1")
    ws["A1"] = f"{symbol} 长期下跌风险触发后30个交易日股价跟踪"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A1"].alignment = Alignment(horizontal="center")

    base_close = float(df.iloc[0]["TrackClose"])
    latest_date = df.index[-1].date()
    meta_rows = [
        ("股票代码", symbol),
        ("首次触发日", trigger_date),
        ("触发日收盘价", base_close),
        ("更新到交易日", latest_date),
        ("已记录交易日", f"{len(df)}/{TRACKING_DAYS}"),
    ]
    for idx, (label, value) in enumerate(meta_rows, start=2):
        ws.cell(row=idx, column=1, value=label)
        ws.cell(row=idx, column=2, value=value)
        ws.cell(row=idx, column=1).font = Font(bold=True)
    ws["B3"].number_format = "yyyy-mm-dd"
    ws["B4"].number_format = "0.00"
    ws["B5"].number_format = "yyyy-mm-dd"

    headers = ["交易日", "收盘价", "当日涨跌幅%", "相对触发日涨跌幅", "进度"]
    header_row = 7
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=header_row, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E78")
        cell.alignment = Alignment(horizontal="center")

    first_data_row = header_row + 1
    for offset, (idx, row) in enumerate(df.iterrows()):
        excel_row = first_data_row + offset
        trade_date = idx.date() if hasattr(idx, "date") else idx
        ws.cell(row=excel_row, column=1, value=trade_date)
        ws.cell(row=excel_row, column=2, value=float(row["TrackClose"]))
        daily_change = row.get("TrackDailyChangePct")
        ws.cell(row=excel_row, column=3, value=None if pd.isna(daily_change) else float(daily_change))
        ws.cell(row=excel_row, column=4, value=f"=B{excel_row}/$B${first_data_row}-1")
        ws.cell(row=excel_row, column=5, value=f"第{offset + 1}/{TRACKING_DAYS}个交易日")
        ws.cell(row=excel_row, column=1).number_format = "yyyy-mm-dd"
        ws.cell(row=excel_row, column=2).number_format = "0.00"
        ws.cell(row=excel_row, column=3).number_format = "0.00"
        ws.cell(row=excel_row, column=4).number_format = "0.00%"

    last_data_row = first_data_row + len(df) - 1
    chart = LineChart()
    chart.title = f"{symbol} 相对触发日涨跌幅"
    chart.y_axis.title = "涨跌幅"
    chart.x_axis.title = "交易日"
    chart.y_axis.numFmt = "0%"
    chart.height = 10
    chart.width = 18
    data = Reference(ws, min_col=4, min_row=header_row, max_row=last_data_row)
    cats = Reference(ws, min_col=1, min_row=first_data_row, max_row=last_data_row)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    chart.legend = None
    ws.add_chart(chart, "F2")

    widths = {"A": 14, "B": 12, "C": 18, "D": 18}
    for col, width in widths.items():
        ws.column_dimensions[col].width = width
    for col in range(5, 15):
        ws.column_dimensions[get_column_letter(col)].width = 12

    wb.save(path)


def update_dip_tracking_excels(price_cache: Optional[Dict[str, pd.DataFrame]] = None) -> List[Path]:
    """根据 BOLL-dip.txt 更新所有触发股票的30日跟踪 Excel。"""

    history = read_dip_trigger_history()
    if not history:
        logging.info("BOLL-dip.txt 暂无触发记录，未生成 Excel 跟踪文件。")
        return []

    EXCEL_TRACKING_DIR.mkdir(parents=True, exist_ok=True)
    updated_paths: List[Path] = []
    cache = price_cache or {}

    for symbol, trigger_dates in sorted(history.items()):
        try:
            full_df = cache.get(symbol)
            if full_df is None or full_df.empty:
                full_df, err = fetch_ohlcv(symbol)
                if full_df.empty:
                    logging.warning("%s Excel 跟踪更新失败：%s", symbol, err or "无法获取行情数据")
                    continue

            episodes = build_tracking_episodes(full_df, trigger_dates)
            for episode_number, trigger_date in enumerate(episodes, start=1):
                df = tracking_prices_from_df(full_df, trigger_date)
                if df.empty:
                    logging.warning("%s_%s Excel 跟踪更新失败：没有 %s 之后的收盘数据", symbol, episode_number, trigger_date)
                    continue
                workbook_path = tracking_workbook_path(EXCEL_TRACKING_DIR, symbol, episode_number, trigger_date)
                write_tracking_workbook(workbook_path, symbol, trigger_date, df)
                updated_paths.append(workbook_path)
                logging.info("已更新 %s_%s 跟踪 Excel：%s，记录 %s/%s 个交易日。", symbol, episode_number, workbook_path, len(df), TRACKING_DAYS)
        except Exception as exc:
            logging.exception("%s Excel 跟踪更新失败：%s", symbol, exc)

    return updated_paths


# =========================
# 主流程
# =========================

def run_monitor() -> None:
    """执行完整监控流程。"""

    setup_logging()
    logging.info("开始 BOLL + MACD + RSI 长期下跌风险监控，共 %s 只股票。", len(STOCK_POOL))

    results: List[RiskResult] = []
    errors: List[RiskResult] = []
    price_cache: Dict[str, pd.DataFrame] = {}

    for idx, ticker in enumerate(STOCK_POOL, start=1):
        logging.info("[%s/%s] 分析 %s", idx, len(STOCK_POOL), ticker)
        df, error = fetch_ohlcv(ticker)
        if df.empty:
            result = RiskResult(ticker=ticker, error=error)
            errors.append(result)
            logging.error("%s 数据获取失败：%s", ticker, error)
            continue
        price_cache[ticker] = df
        try:
            result = score_ticker(ticker, df)
            if result.error:
                errors.append(result)
                logging.error("%s 指标计算失败：%s", ticker, result.error)
            else:
                results.append(result)
                logging.info("%s %s 收盘 %.2f 信号 %s 条件 %s", ticker, result.date, result.close, result.signal_grade, "；".join(result.risk_sources) or "无")
        except Exception as exc:
            errors.append(RiskResult(ticker=ticker, error=str(exc)))
            logging.exception("%s 分析失败：%s", ticker, exc)

    triggered = sorted(
        [r for r in results if r.signal_grade != "无信号"],
        key=lambda x: GRADE_RANK.get(x.signal_grade, 0),
        reverse=True,
    )
    if not triggered:
        logging.info("没有触发 A/B/C 长期下跌风险信号的股票，不生成 Word 报告。")
        if errors:
            logging.info("数据异常股票：%s", ", ".join(f"{r.ticker}:{r.error}" for r in errors))
        update_dip_tracking_excels(price_cache)
        return

    report_date = triggered[0].date or latest_complete_us_date()
    append_triggered_txt(triggered, report_date)
    update_dip_tracking_excels(price_cache)
    report_path = write_word_report(triggered, report_date)
    logging.info("已生成 Word 报告：%s", report_path)
    if errors:
        logging.info("数据异常股票：%s", ", ".join(f"{r.ticker}:{r.error}" for r in errors))


def clear_log_file() -> None:
    """运行结束后清空日志文件，避免日志长期累积。"""

    logging.shutdown()
    try:
        LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
        LOG_PATH.write_text("", encoding="utf-8")
    except Exception:
        pass


if __name__ == "__main__":
    try:
        run_monitor()
    finally:
        clear_log_file()

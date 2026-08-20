from __future__ import annotations

import argparse
import contextlib
import json
import logging
import math
import os
import re
import sys
import time
import traceback
import urllib.parse
import xml.etree.ElementTree as ET
from dataclasses import dataclass, field
from datetime import date, datetime, time as dtime, timedelta
from io import StringIO
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from zoneinfo import ZoneInfo

os.environ.setdefault("PANDAS_USE_NUMEXPR", "0")
os.environ.setdefault("PANDAS_USE_BOTTLENECK", "0")

with contextlib.redirect_stderr(StringIO()):
    import numpy as np
    import pandas as pd

import requests

try:
    import yfinance as yf
except Exception:
    yf = None

try:
    import pandas_market_calendars as mcal
except Exception:
    mcal = None

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None
    Pt = None

try:
    from openpyxl import Workbook
    from openpyxl.chart import LineChart, Reference
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter
except Exception:
    Workbook = None
    LineChart = None
    Reference = None
    Alignment = None
    Font = None
    PatternFill = None
    get_column_letter = None


REPORT_BASE_DIR = Path(os.environ.get("REPORT_BASE_DIR", Path(__file__).resolve().parents[1] / "data"))
OUTPUT_DIR = REPORT_BASE_DIR / "CROSS"
EXCEL_TRACKING_DIR = OUTPUT_DIR / "EXCLE"
GOOGLE_SHEET_ID = "1dD2z5C2-f02N-BaZlghyIKDfwmGJf0rdU6mwYqUszak"
GOOGLE_WORKSHEET_NAME = "CROSS"
GOOGLE_SERVICE_ACCOUNT_ENV = "GOOGLE_SERVICE_ACCOUNT_JSON"
GOOGLE_SERVICE_ACCOUNT_FILE = OUTPUT_DIR / "google_service_account.json"
TRACKING_DAYS = 30
TICKERS = [
    "NOK", "XE", "RKLB", "NBIS", "ARM", "OPTX", "GLW", "COHR", "LITE",
    "MRVL", "INTC", "ALAB", "FLNC", "CRWV", "ORCL", "NVTS", "POET",
    "DELL", "BB", "ONDS", "SMR", "OKLO", "IONQ", "QBTS", "NVDA",
    "AMZN", "AMD", "CRCL", "FLY", "CBRS", "CIEN", "SIVEF", "SHAZ",
    "LUNR", "ASTS", "AAOI", "INOD", "SPCX", "RDW", "QCOM", "COIN",
    "CORZ", "IREN", "VELO", "MSFT", "GOOG", "AVGO", "AMAT", "AMKR",
    "LRCX", "SNDK", "MU", "VICR", "NNE", "CCJ", "VST", "NEE", "BWAY",
    "MANE", "QSI", "ERAS", "CRSP", "IBRX", "INSP",
]

HISTORY_PERIOD = "2y"
INTERVAL = "1d"
TRIGGER_SCORE = 60
REQUEST_TIMEOUT = 15
SLEEP_BETWEEN_TICKERS = 0.2
NEWS_LOOKBACK_DAYS = [1, 3, 7]

MACD_FAST = 12
MACD_SLOW = 26
MACD_SIGNAL = 9
KDJ_PERIOD = 9
K_SMOOTH = 3
D_SMOOTH = 3


@dataclass
class NewsItem:
    title: str
    published: str
    source: str
    link: str
    summary_cn: str
    relevance: str


@dataclass
class StockResult:
    ticker: str
    company_name: str = ""
    df: Optional[pd.DataFrame] = None
    data_source: str = ""
    data_error: str = ""
    data_date: Optional[pd.Timestamp] = None
    latest_close: Optional[float] = None
    latest_change_pct: Optional[float] = None
    volume_ratio: Optional[float] = None
    score_parts: Dict[str, float] = field(default_factory=dict)
    total_score: float = 0.0
    signal_level: str = "不触发"
    macd_status: str = "无金叉"
    kdj_status: str = "无金叉"
    rsi_status: str = "无金叉"
    near_cross_notes: List[str] = field(default_factory=list)
    technical_summary: str = ""
    news: List[NewsItem] = field(default_factory=list)
    reason_categories: List[str] = field(default_factory=list)
    confidence: str = "低"
    reason_summary: str = ""
    risks: List[str] = field(default_factory=list)


def setup_logging(output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    log_path = output_dir / "cross_monitor.log"
    logging.getLogger("yfinance").setLevel(logging.CRITICAL)
    logging.getLogger("peewee").setLevel(logging.CRITICAL)
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[
            logging.FileHandler(log_path, encoding="utf-8"),
            logging.StreamHandler(sys.stdout),
        ],
    )


def eastern_now() -> datetime:
    return datetime.now(ZoneInfo("America/New_York"))


def fmt_date(value) -> str:
    if value is None or pd.isna(value):
        return "N/A"
    if isinstance(value, pd.Timestamp):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    return str(value)


def safe_float(value) -> Optional[float]:
    try:
        if value is None or pd.isna(value):
            return None
        value = float(value)
        if not math.isfinite(value):
            return None
        return value
    except Exception:
        return None


def fmt_num(value, digits: int = 2) -> str:
    value = safe_float(value)
    if value is None:
        return "N/A"
    return f"{value:.{digits}f}"


def fmt_price(value) -> str:
    value = safe_float(value)
    if value is None:
        return "N/A"
    return f"${value:.2f}"


def clean_text(text: str) -> str:
    text = re.sub(r"<[^>]+>", " ", str(text or ""))
    return " ".join(text.split())


def get_company_name(ticker: str) -> str:
    fallback = {
        "NOK": "Nokia Oyj", "XE": "X-energy", "RKLB": "Rocket Lab USA",
        "NBIS": "Nebius Group", "ARM": "Arm Holdings", "OPTX": "Syntec Optics Holdings",
        "GLW": "Corning", "COHR": "Coherent", "LITE": "Lumentum Holdings",
        "MRVL": "Marvell Technology", "INTC": "Intel", "ALAB": "Astera Labs",
        "FLNC": "Fluence Energy", "CRWV": "CoreWeave", "ORCL": "Oracle",
        "NVTS": "Navitas Semiconductor", "POET": "POET Technologies",
        "DELL": "Dell Technologies", "BB": "BlackBerry", "ONDS": "Ondas Holdings",
        "SMR": "NuScale Power", "OKLO": "Oklo", "IONQ": "IonQ",
        "QBTS": "D-Wave Quantum", "NVDA": "NVIDIA", "AMZN": "Amazon.com",
        "AMD": "Advanced Micro Devices", "CRCL": "Circle Internet Group",
        "FLY": "Firefly Aerospace", "CBRS": "Chain Bridge Bancorp",
        "CIEN": "Ciena", "SIVEF": "Sivers Semiconductors", "SHAZ": "Shaz",
        "LUNR": "Intuitive Machines", "ASTS": "AST SpaceMobile",
        "AAOI": "Applied Optoelectronics", "INOD": "Innodata",
    }
    if ticker in fallback:
        return fallback[ticker]
    if yf is not None:
        try:
            info = yf.Ticker(ticker).get_info()
            return clean_text(info.get("shortName") or info.get("longName") or ticker)
        except Exception:
            pass
    return ticker


def trading_dates_around(today_et: date) -> List[date]:
    if mcal is not None:
        try:
            nyse = mcal.get_calendar("NYSE")
            sched = nyse.schedule(
                start_date=pd.Timestamp(today_et - timedelta(days=14)),
                end_date=pd.Timestamp(today_et + timedelta(days=3)),
            )
            return [d.date() for d in sched.index]
        except Exception:
            pass
    dates = []
    start = today_et - timedelta(days=14)
    for i in range(18):
        d = start + timedelta(days=i)
        if d.weekday() < 5:
            dates.append(d)
    return dates


def latest_official_trading_day(now_et: Optional[datetime] = None) -> date:
    now_et = now_et or eastern_now()
    today_et = now_et.date()
    after_close = now_et.time() >= dtime(16, 10)
    dates = trading_dates_around(today_et)
    eligible = [d for d in dates if d < today_et or (d == today_et and after_close)]
    if eligible:
        return max(eligible)
    d = today_et - timedelta(days=1)
    while d.weekday() >= 5:
        d -= timedelta(days=1)
    return d


def ensure_ohlcv(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()
    df = df.copy()
    if isinstance(df.columns, pd.MultiIndex):
        df.columns = [c[0] if isinstance(c, tuple) else c for c in df.columns]
    rename_map = {}
    for col in df.columns:
        key = str(col).strip().lower()
        if key in {"open", "high", "low", "close", "volume"}:
            rename_map[col] = key.capitalize()
        elif key in {"adj close", "adjusted close"}:
            rename_map[col] = "Adjusted Close"
    df = df.rename(columns=rename_map)
    required = ["Open", "High", "Low", "Close", "Volume"]
    if any(c not in df.columns for c in required):
        return pd.DataFrame()
    if "Adjusted Close" not in df.columns:
        df["Adjusted Close"] = df["Close"]
    df = df[["Open", "High", "Low", "Close", "Adjusted Close", "Volume"]].copy()
    df.index = pd.to_datetime(df.index)
    if getattr(df.index, "tz", None) is not None:
        df.index = df.index.tz_convert("America/New_York").tz_localize(None)
    df.index = df.index.normalize()
    df = df.sort_index()
    df = df[~df.index.duplicated(keep="last")]
    for col in df.columns:
        df[col] = pd.to_numeric(df[col], errors="coerce")
    df = df.dropna(subset=["Open", "High", "Low", "Close", "Volume"])
    df = df[df["Close"] > 0]
    return df


def fetch_yfinance(ticker: str) -> Tuple[pd.DataFrame, str]:
    if yf is None:
        return pd.DataFrame(), "未安装或无法导入 yfinance"
    try:
        with contextlib.redirect_stdout(StringIO()), contextlib.redirect_stderr(StringIO()):
            df = yf.download(
                ticker,
                period=HISTORY_PERIOD,
                interval=INTERVAL,
                auto_adjust=False,
                progress=False,
                threads=False,
                timeout=10,
            )
        df = ensure_ohlcv(df)
        if df.empty:
            return df, "yfinance 返回空数据或字段不足"
        return df, ""
    except Exception as exc:
        return pd.DataFrame(), f"yfinance 异常：{exc}"


def fetch_yahoo_chart(ticker: str) -> Tuple[pd.DataFrame, str]:
    url = f"https://query1.finance.yahoo.com/v8/finance/chart/{ticker}"
    params = {
        "range": HISTORY_PERIOD,
        "interval": INTERVAL,
        "events": "history",
        "includeAdjustedClose": "true",
    }
    try:
        response = requests.get(
            url,
            params=params,
            headers={"User-Agent": "Mozilla/5.0 (compatible; cross-monitor/1.0)"},
            timeout=REQUEST_TIMEOUT,
        )
        response.raise_for_status()
        payload = response.json()
        chart = payload.get("chart", {})
        if chart.get("error"):
            return pd.DataFrame(), str(chart["error"])
        result = chart.get("result") or []
        if not result:
            return pd.DataFrame(), "Yahoo Chart API 未返回 result"
        block = result[0]
        timestamps = block.get("timestamp") or []
        quote = (block.get("indicators", {}).get("quote") or [{}])[0]
        adjclose = (block.get("indicators", {}).get("adjclose") or [{}])[0].get("adjclose")
        if not timestamps or not quote:
            return pd.DataFrame(), "Yahoo Chart API 返回数据不完整"
        df = pd.DataFrame({
            "Date": pd.to_datetime(timestamps, unit="s", utc=True)
                .tz_convert("America/New_York")
                .tz_localize(None)
                .normalize(),
            "Open": quote.get("open"),
            "High": quote.get("high"),
            "Low": quote.get("low"),
            "Close": quote.get("close"),
            "Adjusted Close": adjclose if adjclose is not None else quote.get("close"),
            "Volume": quote.get("volume"),
        }).set_index("Date")
        df = ensure_ohlcv(df)
        if df.empty:
            return df, "Yahoo Chart API 清洗后为空"
        return df, ""
    except Exception as exc:
        return pd.DataFrame(), f"Yahoo Chart API 异常：{exc}"


def fetch_ohlcv(ticker: str, data_day: date) -> Tuple[pd.DataFrame, str, str]:
    errors = []
    for name, func in [("yfinance", fetch_yfinance), ("Yahoo Chart API", fetch_yahoo_chart)]:
        df, err = func(ticker)
        if not df.empty:
            df = df[df.index.date <= data_day]
            if df.empty:
                errors.append(f"{name}: 没有 {data_day} 及以前的正式收盘数据")
                continue
            return df, name, ""
        errors.append(f"{name}: {err}")
    return pd.DataFrame(), "", "；".join(errors)


def wilder_rsi(series: pd.Series, period: int) -> pd.Series:
    delta = series.diff()
    gain = delta.clip(lower=0)
    loss = -delta.clip(upper=0)
    avg_gain = gain.ewm(alpha=1 / period, adjust=False, min_periods=period).mean()
    avg_loss = loss.ewm(alpha=1 / period, adjust=False, min_periods=period).mean()
    rs = avg_gain / avg_loss.replace(0, np.nan)
    rsi = 100 - (100 / (1 + rs))
    return rsi.fillna(50)


def add_indicators(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    price = df["Adjusted Close"].where(df["Adjusted Close"].notna(), df["Close"])
    df["TrendClose"] = price
    df["MA20"] = price.rolling(20).mean()
    df["MA50"] = price.rolling(50).mean()
    df["Prev20High"] = df["High"].shift(1).rolling(20).max()
    df["VolumeMA20"] = df["Volume"].rolling(20).mean()
    df["VolumeRatio"] = df["Volume"] / df["VolumeMA20"].replace(0, np.nan)
    df["DailyChangePct"] = (price / price.shift(1) - 1) * 100
    daily_range = (df["High"] - df["Low"]).replace(0, np.nan)
    df["CloseLocation"] = ((df["Close"] - df["Low"]) / daily_range).fillna(0.5)

    ema_fast = price.ewm(span=MACD_FAST, adjust=False).mean()
    ema_slow = price.ewm(span=MACD_SLOW, adjust=False).mean()
    df["DIF"] = ema_fast - ema_slow
    df["DEA"] = df["DIF"].ewm(span=MACD_SIGNAL, adjust=False).mean()
    df["MACD_HIST"] = df["DIF"] - df["DEA"]

    low_n = df["Low"].rolling(KDJ_PERIOD).min()
    high_n = df["High"].rolling(KDJ_PERIOD).max()
    denom = (high_n - low_n).replace(0, np.nan)
    rsv = ((df["Close"] - low_n) / denom * 100).clip(0, 100).fillna(50)
    df["K"] = rsv.ewm(alpha=1 / K_SMOOTH, adjust=False).mean()
    df["D"] = df["K"].ewm(alpha=1 / D_SMOOTH, adjust=False).mean()
    df["J"] = 3 * df["K"] - 2 * df["D"]

    df["RSI6"] = wilder_rsi(price, 6)
    df["RSI12"] = wilder_rsi(price, 12)
    df["RSI24"] = wilder_rsi(price, 24)
    return df


def is_cross_up(fast: pd.Series, slow: pd.Series, offset: int = 0) -> bool:
    if len(fast) < offset + 2 or len(slow) < offset + 2:
        return False
    today = -1 - offset
    yesterday = -2 - offset
    return bool(fast.iloc[today] > slow.iloc[today] and fast.iloc[yesterday] <= slow.iloc[yesterday])


def crossed_within_two_days(df: pd.DataFrame, fast_col: str, slow_col: str) -> bool:
    fast = df[fast_col]
    slow = df[slow_col]
    today_cross = is_cross_up(fast, slow, 0)
    prev_cross = is_cross_up(fast, slow, 1)
    still_bullish = fast.iloc[-1] > slow.iloc[-1]
    return bool(still_bullish and (today_cross or prev_cross))


def is_gap_improving(df: pd.DataFrame, fast_col: str, slow_col: str, days: int = 3) -> bool:
    gap = (df[fast_col] - df[slow_col]).dropna().tail(days)
    if len(gap) < days:
        return False
    return bool(gap.iloc[-1] < 0 and (gap.diff().dropna() > 0).sum() >= days - 2)


def macd_zero_score(df: pd.DataFrame) -> float:
    row = df.iloc[-1]
    close = row["TrendClose"]
    if close <= 0:
        return 0
    ndif = row["DIF"] / close
    ndea = row["DEA"] / close
    if row["DIF"] > 0 and row["DEA"] > 0:
        return 10
    hist_norm = (df["DIF"] / df["TrendClose"].replace(0, np.nan)).dropna().abs()
    threshold = hist_norm.quantile(0.35) if len(hist_norm) >= 60 else 0.01
    if abs(ndif) <= threshold and abs(ndea) <= threshold:
        return 6
    return 0


def score_macd(df: pd.DataFrame, result: StockResult) -> float:
    row = df.iloc[-1]
    prev = df.iloc[-2]
    score = 0.0
    if crossed_within_two_days(df, "DIF", "DEA"):
        score += 10
        result.macd_status = "最近两个交易日内有效金叉"
    elif is_gap_improving(df, "DIF", "DEA"):
        result.macd_status = "即将金叉"
        result.near_cross_notes.append("MACD 即将金叉")
    score += macd_zero_score(df)

    hist = df["MACD_HIST"].tail(4)
    hist_diff = hist.diff().dropna()
    if len(hist) >= 4 and hist.iloc[-1] > 0 and (hist_diff.tail(3) > 0).all():
        score += 10
    elif len(hist) >= 3 and hist.iloc[-1] > 0 and (hist_diff.tail(2) > 0).all():
        score += 7
    elif row["MACD_HIST"] > prev["MACD_HIST"] or (row["MACD_HIST"] > 0 and prev["MACD_HIST"] <= 0):
        score += 4

    if row["DIF"] > prev["DIF"] and row["DEA"] > prev["DEA"]:
        score += 5
    return min(score, 35)


def score_kdj(df: pd.DataFrame, result: StockResult) -> float:
    row = df.iloc[-1]
    score = 0.0
    if crossed_within_two_days(df, "K", "D"):
        score += 10
        result.kdj_status = "最近两个交易日内有效金叉"
    elif is_gap_improving(df, "K", "D"):
        result.kdj_status = "即将金叉"
        result.near_cross_notes.append("KDJ 即将金叉")

    k, d, j = row["K"], row["D"], row["J"]
    center = (k + d) / 2
    if 30 <= center <= 65 or 40 <= center <= 70:
        score += 5
    elif 20 <= center <= 75:
        score += 3

    j_flags = [j > k, j > 50, j > df["J"].iloc[-2]]
    count = sum(bool(x) for x in j_flags)
    if count == 3:
        score += 5
    elif count == 2:
        score += 3
    elif count == 1:
        score += 1
    return min(score, 20)


def score_rsi(df: pd.DataFrame, result: StockResult) -> float:
    row = df.iloc[-1]
    prev = df.iloc[-2]
    score = 0.0
    rsi_crosses = []
    if is_cross_up(df["RSI6"], df["RSI12"], 0):
        rsi_crosses.append("RSI6上穿RSI12")
    if is_cross_up(df["RSI12"], df["RSI24"], 0):
        rsi_crosses.append("RSI12上穿RSI24")
    if rsi_crosses:
        result.rsi_status = "；".join(rsi_crosses)
    elif row["RSI6"] > row["RSI12"] > row["RSI24"]:
        result.rsi_status = "RSI多头排列"
    elif is_gap_improving(df, "RSI6", "RSI12") or is_gap_improving(df, "RSI12", "RSI24"):
        result.rsi_status = "即将金叉"
        result.near_cross_notes.append("RSI 即将金叉")

    if row["RSI6"] > row["RSI12"] > row["RSI24"]:
        score += 10
    if row["RSI12"] >= 60:
        score += 5
    elif row["RSI12"] >= 55:
        score += 4
    elif row["RSI12"] >= 50:
        score += 3
    if row["RSI24"] >= 50 and row["RSI24"] > prev["RSI24"]:
        score += 5
    elif row["RSI24"] >= 45 and row["RSI24"] > prev["RSI24"]:
        score += 3
    return min(score, 20)


def score_price_trend(df: pd.DataFrame) -> float:
    row = df.iloc[-1]
    score = 0.0
    above_ma20 = row["TrendClose"] > row["MA20"] if pd.notna(row["MA20"]) else False
    above_ma50 = row["TrendClose"] > row["MA50"] if pd.notna(row["MA50"]) else False
    if above_ma20 and above_ma50:
        score += 10
    elif above_ma20 or above_ma50:
        score += 5
    if pd.notna(row["Prev20High"]) and row["TrendClose"] > row["Prev20High"]:
        score += 5
    return min(score, 15)


def score_volume(df: pd.DataFrame) -> float:
    vr = safe_float(df.iloc[-1]["VolumeRatio"])
    if vr is None:
        return 0
    if vr >= 1.5:
        return 10
    if vr >= 1.3:
        return 8
    if vr >= 1.2:
        return 5
    return 0


def signal_level(score: float) -> str:
    if score >= 80:
        return "强势金叉"
    if score >= 71:
        return "较强金叉"
    if score >= 60:
        return "普通有效信号"
    return "不触发"


def build_technical_summary(result: StockResult) -> str:
    df = result.df
    if df is None or df.empty:
        return ""
    row = df.iloc[-1]
    points = []
    if result.macd_status != "无金叉":
        points.append(f"MACD{result.macd_status}，DIF/DEA为{fmt_num(row['DIF'], 4)}/{fmt_num(row['DEA'], 4)}")
    if result.kdj_status != "无金叉":
        points.append(f"KDJ{result.kdj_status}，K/D/J为{fmt_num(row['K'], 1)}/{fmt_num(row['D'], 1)}/{fmt_num(row['J'], 1)}")
    if result.rsi_status != "无金叉":
        points.append(f"RSI状态为{result.rsi_status}，RSI12为{fmt_num(row['RSI12'], 1)}")
    if row["TrendClose"] > row["MA20"] and row["TrendClose"] > row["MA50"]:
        points.append("股价站上MA20和MA50")
    elif row["TrendClose"] > row["MA20"] or row["TrendClose"] > row["MA50"]:
        points.append("股价站上部分均线")
    if pd.notna(row["Prev20High"]) and row["TrendClose"] > row["Prev20High"]:
        points.append("收盘价突破过去20日高点")
    elif pd.notna(row["Prev20High"]) and (row["Prev20High"] - row["TrendClose"]) / row["Prev20High"] <= 0.02:
        points.append("接近突破过去20日高点")
    if not points:
        points.append("技术指标尚未形成明显强势共振")
    return "；".join(points) + "。"


def parse_rss_items(xml_text: str, max_items: int = 6) -> List[NewsItem]:
    items: List[NewsItem] = []
    try:
        root = ET.fromstring(xml_text)
    except Exception:
        return items
    for item in root.findall(".//item"):
        title = clean_text(item.findtext("title"))
        link = clean_text(item.findtext("link"))
        pub = clean_text(item.findtext("pubDate"))
        desc = clean_text(item.findtext("description"))
        source = "Yahoo Finance"
        source_node = item.find("{http://search.yahoo.com/mrss/}credit")
        if source_node is not None and source_node.text:
            source = clean_text(source_node.text)
        if title and link:
            items.append(NewsItem(title, pub, source, link, desc[:120], ""))
        if len(items) >= max_items:
            break
    return items


def search_yahoo_news(ticker: str) -> List[NewsItem]:
    url = f"https://feeds.finance.yahoo.com/rss/2.0/headline?s={urllib.parse.quote(ticker)}&region=US&lang=en-US"
    try:
        res = requests.get(url, timeout=REQUEST_TIMEOUT, headers={"User-Agent": "Mozilla/5.0"})
        if res.ok:
            return parse_rss_items(res.text, 8)
    except Exception:
        pass
    return []


def search_google_news(ticker: str, company_name: str) -> List[NewsItem]:
    query = urllib.parse.quote(f"{ticker} {company_name} stock")
    url = f"https://news.google.com/rss/search?q={query}&hl=en-US&gl=US&ceid=US:en"
    try:
        res = requests.get(url, timeout=REQUEST_TIMEOUT, headers={"User-Agent": "Mozilla/5.0"})
        if res.ok:
            items = parse_rss_items(res.text, 8)
            for item in items:
                if item.source == "Yahoo Finance":
                    item.source = "Google News"
            return items
    except Exception:
        pass
    return []


def search_news_for_result(result: StockResult) -> List[NewsItem]:
    items = search_yahoo_news(result.ticker)
    if len(items) < 2:
        items.extend(search_google_news(result.ticker, result.company_name))
    seen = set()
    deduped = []
    for item in items:
        key = (item.title.lower(), item.link.split("?")[0])
        if key in seen:
            continue
        seen.add(key)
        item.summary_cn = summarize_news_cn(item.title, item.summary_cn)
        item.relevance = classify_news_relevance(item.title + " " + item.summary_cn)
        deduped.append(item)
        if len(deduped) >= 3:
            break
    return deduped


def summarize_news_cn(title: str, description: str) -> str:
    text = f"{title} {description}".lower()
    if any(k in text for k in ["earnings", "revenue", "guidance", "profit"]):
        return "新闻涉及财报、收入或业绩指引，可能影响市场对基本面的预期。"
    if any(k in text for k in ["contract", "order", "partnership", "deal", "award"]):
        return "新闻涉及订单、合同或合作事项，可能增强市场对后续增长的关注。"
    if any(k in text for k in ["upgrade", "price target", "analyst", "rating"]):
        return "新闻涉及分析师评级或目标价变化，可能影响短线情绪。"
    if any(k in text for k in ["sec", "8-k", "offering", "convertible", "atm"]):
        return "新闻或公告涉及监管文件、融资或资本运作，需要结合细节判断影响。"
    return "新闻与公司股价或行业动态相关，但是否直接解释本次技术信号仍需进一步确认。"


def classify_news_relevance(text: str) -> str:
    text = text.lower()
    if any(k in text for k in ["earnings", "revenue", "guidance"]):
        return "财报或业绩预期改善"
    if any(k in text for k in ["contract", "order", "partnership", "deal", "award"]):
        return "新订单、合同或合作"
    if any(k in text for k in ["upgrade", "price target", "analyst", "rating"]):
        return "分析师评级或目标价上调"
    if any(k in text for k in ["offering", "convertible", "atm", "merger", "acquisition"]):
        return "并购、融资或资本运作"
    return "新闻利好但持续性存疑"


def analyze_reason(result: StockResult) -> None:
    categories = []
    confidence = "低"
    if result.news:
        for item in result.news:
            if item.relevance and item.relevance not in categories:
                categories.append(item.relevance)
        if any(c in categories for c in ["财报或业绩预期改善", "新订单、合同或合作", "分析师评级或目标价上调"]):
            confidence = "中"
    if result.total_score >= TRIGGER_SCORE:
        if "技术性突破" not in categories:
            categories.append("技术性突破")
    if not categories:
        categories = ["暂无明确新闻催化"]
    result.reason_categories = categories[:3]
    result.confidence = confidence
    if result.news:
        result.reason_summary = f"检索到相关新闻，主要关联为{'、'.join(result.reason_categories)}。"
    else:
        result.reason_summary = "未发现足以解释本次信号的明确重大新闻，更可能属于技术性反弹、板块联动或资金推动，判断置信度较低。"


def build_risks(result: StockResult) -> List[str]:
    df = result.df
    if df is None or df.empty:
        return ["数据不足，风险无法完整评估。"]
    row = df.iloc[-1]
    risks = []
    if row["RSI6"] > 80:
        risks.append("RSI6高于80，短线存在过热风险。")
    if row["RSI12"] > 75:
        risks.append("RSI12高于75，需防短线震荡。")
    if row["K"] > 80 and row["D"] > 80:
        risks.append("KDJ处于高位，可能出现高位钝化或回落。")
    if pd.notna(row["MA20"]) and row["MA20"] > 0 and (row["TrendClose"] / row["MA20"] - 1) > 0.10:
        risks.append("股价远离MA20超过10%，短线追高风险较高。")
    if row["CloseLocation"] <= 0.3:
        risks.append("收盘位置靠近当日低位，存在冲高回落风险。")
    if result.volume_ratio is not None and result.volume_ratio < 1.2:
        risks.append("成交量未明显放大，需防信号确认不足。")
    if len([s for s in [result.macd_status, result.kdj_status, result.rsi_status] if s != "无金叉"]) <= 1:
        risks.append("仅少数指标确认，可能是假突破。")
    if not risks:
        risks.append("重点观察次日量能和指标延续性。")
    return risks[:2]


def analyze_ticker(ticker: str, data_day: date) -> StockResult:
    result = StockResult(ticker=ticker, company_name=get_company_name(ticker))
    df, source, err = fetch_ohlcv(ticker, data_day)
    if df.empty:
        result.data_error = err or "无法获取行情数据"
        logging.warning("%s 数据获取失败：%s", ticker, result.data_error)
        return result
    if len(df) < 80:
        result.data_error = "历史行情不足，无法稳定计算指标"
        logging.warning("%s 历史行情不足：%s 行", ticker, len(df))
        return result
    result.data_source = source
    df = add_indicators(df)
    df = df.dropna(subset=["MA20", "MA50", "DIF", "DEA", "MACD_HIST", "K", "D", "J", "RSI6", "RSI12", "RSI24"])
    if len(df) < 30:
        result.data_error = "指标有效数据不足"
        logging.warning("%s 指标有效数据不足", ticker)
        return result
    result.df = df
    row = df.iloc[-1]
    result.data_date = df.index[-1]
    if result.data_date.date() != data_day:
        result.data_error = f"最新行情日期为 {result.data_date.date()}，缺少 {data_day} 的正式收盘数据"
        logging.warning("%s 数据日期滞后：%s", ticker, result.data_error)
        return result
    result.latest_close = safe_float(row["TrendClose"])
    result.latest_change_pct = safe_float(row["DailyChangePct"])
    result.volume_ratio = safe_float(row["VolumeRatio"])

    result.score_parts["MACD"] = score_macd(df, result)
    result.score_parts["KDJ"] = score_kdj(df, result)
    result.score_parts["RSI"] = score_rsi(df, result)
    result.score_parts["PriceTrend"] = score_price_trend(df)
    result.score_parts["Volume"] = score_volume(df)
    result.total_score = max(0, min(100, sum(result.score_parts.values())))
    result.signal_level = signal_level(result.total_score)
    result.technical_summary = build_technical_summary(result)
    result.risks = build_risks(result)
    return result


def report_exists_for_day(output_dir: Path, data_day: date) -> Optional[Path]:
    path = output_dir / f"强势金叉监控报告_{data_day.strftime('%Y-%m-%d')}.docx"
    return path if path.exists() else None


def append_triggered_symbols_txt(output_dir: Path, report_date: date, triggered: List[StockResult]) -> List[Tuple[str, str, str]]:
    if not triggered:
        return []
    txt_path = output_dir / "CROSS.txt"
    existing = set()
    symbol_counts: Dict[str, int] = {}

    def parse_symbol_from_line(line: str) -> str:
        parts = line.strip().split()
        if len(parts) < 2:
            return ""
        symbol_text = parts[1].strip()
        if "（" in symbol_text:
            symbol_text = symbol_text.split("（", 1)[0].strip()
        elif "(" in symbol_text:
            symbol_text = symbol_text.split("(", 1)[0].strip()
        return symbol_text

    def parse_date_symbol_key(line: str) -> str:
        parts = line.strip().split()
        if len(parts) < 2:
            return line.strip()
        symbol_text = parse_symbol_from_line(line)
        return f"{parts[0]}\t{symbol_text}"

    if txt_path.exists():
        try:
            for raw_line in txt_path.read_text(encoding="utf-8-sig").splitlines():
                line = raw_line.strip()
                if not line:
                    continue
                existing.add(parse_date_symbol_key(line))
                symbol = parse_symbol_from_line(line)
                if symbol:
                    symbol_counts[symbol] = symbol_counts.get(symbol, 0) + 1
        except Exception as exc:
            logging.warning("读取 CROSS.txt 失败，将仅写入本次触发结果：%s", exc)
            existing = set()
            symbol_counts = {}

    new_records: List[Tuple[str, str, str]] = []
    date_text = report_date.strftime("%Y-%m-%d")
    for result in triggered:
        key = f"{date_text}\t{result.ticker}"
        if key in existing:
            continue
        next_count = symbol_counts.get(result.ticker, 0) + 1
        symbol_text = f"{result.ticker}（{next_count}）" if next_count > 1 else result.ticker
        score_text = f"{result.total_score:.1f}分"
        new_records.append((date_text, symbol_text, score_text))
        existing.add(key)
        symbol_counts[result.ticker] = next_count

    if new_records:
        with txt_path.open("a", encoding="utf-8-sig") as f:
            for date_text, symbol_text, score_text in new_records:
                f.write(f"{date_text}\t{symbol_text}\t{score_text}\n")
        logging.info("已写入 CROSS.txt：%s", ", ".join(symbol for _, symbol, _ in new_records))
    else:
        logging.info("CROSS.txt 中已存在本交易日触发记录，未重复写入。")
    return new_records


def append_triggered_symbols_google_sheet(records: List[Tuple[str, str, str]]) -> None:
    if not records:
        return
    credential_path = os.environ.get(GOOGLE_SERVICE_ACCOUNT_ENV, "").strip()
    if not credential_path and GOOGLE_SERVICE_ACCOUNT_FILE.exists():
        credential_path = str(GOOGLE_SERVICE_ACCOUNT_FILE)
    if not credential_path:
        logging.info(
            "未配置 Google Sheets 授权，已跳过云端表格写入。可设置环境变量 %s，或把服务账号 JSON 放到：%s",
            GOOGLE_SERVICE_ACCOUNT_ENV,
            GOOGLE_SERVICE_ACCOUNT_FILE,
        )
        return

    try:
        import gspread
    except Exception as exc:
        logging.warning("未安装或无法导入 gspread，已跳过 Google Sheets 写入：%s", exc)
        return

    try:
        client = gspread.service_account(filename=credential_path)
        spreadsheet = client.open_by_key(GOOGLE_SHEET_ID)
        try:
            worksheet = spreadsheet.worksheet(GOOGLE_WORKSHEET_NAME)
        except gspread.WorksheetNotFound:
            worksheet = spreadsheet.add_worksheet(title=GOOGLE_WORKSHEET_NAME, rows=1000, cols=3)
        if not worksheet.get_all_values():
            worksheet.append_row(["触发日期", "股票代码", "评分"], value_input_option="USER_ENTERED")
        worksheet.append_rows(
            [[date_text, symbol_text, score_text] for date_text, symbol_text, score_text in records],
            value_input_option="USER_ENTERED",
        )
        logging.info("已写入 Google Sheets：%s", ", ".join(symbol for _, symbol, _ in records))
    except Exception as exc:
        logging.warning("Google Sheets 写入失败，已保留本地 CROSS.txt 记录：%s", exc)


def parse_plain_ticker(symbol_text: str) -> str:
    symbol = symbol_text.strip().upper()
    if "（" in symbol:
        symbol = symbol.split("（", 1)[0].strip()
    elif "(" in symbol:
        symbol = symbol.split("(", 1)[0].strip()
    return symbol


def read_cross_trigger_history(output_dir: Path) -> Dict[str, List[date]]:
    txt_path = output_dir / "CROSS.txt"
    history: Dict[str, List[date]] = {}
    if not txt_path.exists():
        return history
    try:
        lines = txt_path.read_text(encoding="utf-8-sig").splitlines()
    except Exception as exc:
        logging.warning("读取 CROSS.txt 失败，无法更新 Excel 跟踪文件：%s", exc)
        return history

    for raw_line in lines:
        parts = raw_line.strip().split()
        if len(parts) < 2:
            continue
        try:
            trigger_date = date.fromisoformat(parts[0])
        except ValueError:
            continue
        ticker = parse_plain_ticker(parts[1])
        if not ticker:
            continue
        history.setdefault(ticker, []).append(trigger_date)
    for ticker in list(history):
        history[ticker] = sorted(set(history[ticker]))
    return history


def safe_tracking_filename(ticker: str, episode_number: int) -> str:
    safe = re.sub(r"[^A-Z0-9._-]+", "_", ticker.upper()).strip("._-")
    return f"{safe or 'UNKNOWN'}_{episode_number}"


def tracking_workbook_path(tracking_dir: Path, ticker: str, episode_number: int, trigger_date: date) -> Path:
    base_name = safe_tracking_filename(ticker, episode_number)
    dated_path = tracking_dir / f"{base_name}_{trigger_date.strftime('%Y%m%d')}.xlsx"
    legacy_path = tracking_dir / f"{base_name}.xlsx"
    if dated_path.exists():
        return dated_path
    if legacy_path.exists():
        return legacy_path
    return dated_path


def tracking_prices_from_df(df: pd.DataFrame, trigger_date: date, data_day: date) -> pd.DataFrame:
    df = df.copy()
    df["TrackClose"] = pd.to_numeric(df["Close"], errors="coerce")
    df["TrackDailyChangePct"] = df["TrackClose"].pct_change() * 100
    df = df[(df.index.date >= trigger_date) & (df.index.date <= data_day)].copy()
    if df.empty:
        return pd.DataFrame()
    df = df.head(TRACKING_DAYS)
    df = df.dropna(subset=["TrackClose"])
    return df


def is_trigger_inside_episode(df: pd.DataFrame, episode_start: date, trigger_date: date) -> bool:
    episode_window = df[(df.index.date >= episode_start) & (df.index.date <= trigger_date)]
    if episode_window.empty:
        return True
    return len(episode_window) <= TRACKING_DAYS


def build_tracking_episodes(df: pd.DataFrame, trigger_dates: List[date], data_day: date) -> List[date]:
    episodes: List[date] = []
    for trigger_date in sorted(d for d in set(trigger_dates) if d <= data_day):
        if not episodes:
            episodes.append(trigger_date)
            continue
        if not is_trigger_inside_episode(df, episodes[-1], trigger_date):
            episodes.append(trigger_date)
    return episodes


def write_tracking_workbook(path: Path, ticker: str, episode_number: int, trigger_date: date, data_day: date, df: pd.DataFrame) -> None:
    if Workbook is None:
        raise RuntimeError("未安装 openpyxl，无法生成 Excel 跟踪文件")
    path.parent.mkdir(parents=True, exist_ok=True)

    wb = Workbook()
    ws = wb.active
    ws.title = "30日跟踪"
    ws.freeze_panes = "A8"

    title = f"{ticker}_{episode_number} 触发后30个交易日股价跟踪"
    ws.merge_cells("A1:D1")
    ws["A1"] = title
    ws["A1"].font = Font(bold=True, size=14)
    ws["A1"].alignment = Alignment(horizontal="center")

    base_close = float(df.iloc[0]["TrackClose"])
    meta_rows = [
        ("股票代码", ticker),
        ("跟踪周期", f"第{episode_number}次"),
        ("首次触发日", trigger_date),
        ("触发日收盘价", base_close),
        ("更新到交易日", data_day),
    ]
    for idx, (label, value) in enumerate(meta_rows, start=2):
        ws.cell(row=idx, column=1, value=label)
        ws.cell(row=idx, column=2, value=value)
        ws.cell(row=idx, column=1).font = Font(bold=True)
    ws["B4"].number_format = "yyyy-mm-dd"
    ws["B5"].number_format = "0.00"
    ws["B6"].number_format = "yyyy-mm-dd"

    headers = ["交易日", "收盘价", "当日涨跌幅%", "相对触发日涨跌幅", "进度"]
    header_row = 7
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=header_row, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E78")
        cell.alignment = Alignment(horizontal="center")

    first_data_row = header_row + 1
    for offset, (idx, row) in enumerate(df.iterrows(), start=0):
        excel_row = first_data_row + offset
        trade_date = idx.date() if hasattr(idx, "date") else idx
        ws.cell(row=excel_row, column=1, value=trade_date)
        ws.cell(row=excel_row, column=2, value=float(row["TrackClose"]))
        daily_change = row.get("TrackDailyChangePct")
        ws.cell(row=excel_row, column=3, value=None if pd.isna(daily_change) else float(daily_change))
        ws.cell(row=excel_row, column=4, value=f"=B{excel_row}/$B${first_data_row}-1")
        ws.cell(row=excel_row, column=5, value=f"第{offset + 1}/{TRACKING_DAYS}个交易日")
        ws.cell(row=excel_row, column=1).number_format = "yyyy-mm-dd"
        ws.cell(row=excel_row, column=2).number_format = "0.00"
        ws.cell(row=excel_row, column=3).number_format = "0.00"
        ws.cell(row=excel_row, column=4).number_format = "0.00%"

    last_data_row = first_data_row + len(df) - 1
    chart = LineChart()
    chart.title = f"{ticker}_{episode_number} 相对触发日涨跌幅"
    chart.y_axis.title = "涨跌幅"
    chart.x_axis.title = "交易日"
    chart.y_axis.numFmt = "0%"
    chart.height = 10
    chart.width = 18
    data = Reference(ws, min_col=4, min_row=header_row, max_row=last_data_row)
    cats = Reference(ws, min_col=1, min_row=first_data_row, max_row=last_data_row)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    chart.legend = None
    ws.add_chart(chart, "F2")

    widths = {"A": 14, "B": 12, "C": 18, "D": 18}
    for col, width in widths.items():
        ws.column_dimensions[col].width = width
    for col in range(5, 15):
        ws.column_dimensions[get_column_letter(col)].width = 12

    wb.save(path)


def update_cross_tracking_excels(output_dir: Path, data_day: date) -> List[Path]:
    history = read_cross_trigger_history(output_dir)
    if not history:
        logging.info("CROSS.txt 暂无触发记录，未生成 Excel 跟踪文件。")
        return []
    if Workbook is None:
        logging.warning("未安装 openpyxl，跳过 Excel 跟踪文件生成。")
        return []

    updated_paths: List[Path] = []
    tracking_dir = EXCEL_TRACKING_DIR
    tracking_dir.mkdir(parents=True, exist_ok=True)

    for ticker, trigger_dates in sorted(history.items()):
        try:
            full_df, source, err = fetch_ohlcv(ticker, data_day)
            if full_df.empty:
                logging.warning("%s Excel 跟踪更新失败：%s", ticker, err or "无法获取行情数据")
                continue
            episodes = build_tracking_episodes(full_df, trigger_dates, data_day)
            for episode_number, trigger_date in enumerate(episodes, start=1):
                df = tracking_prices_from_df(full_df, trigger_date, data_day)
                if df.empty:
                    logging.warning("%s_%s Excel 跟踪更新失败：没有 %s 至 %s 的收盘数据", ticker, episode_number, trigger_date, data_day)
                    continue
                workbook_path = tracking_workbook_path(tracking_dir, ticker, episode_number, trigger_date)
                write_tracking_workbook(workbook_path, ticker, episode_number, trigger_date, data_day, df)
                updated_paths.append(workbook_path)
                logging.info(
                    "已更新 %s_%s 跟踪 Excel：%s，记录 %s/%s 个交易日。",
                    ticker,
                    episode_number,
                    workbook_path,
                    len(df),
                    TRACKING_DAYS,
                )
        except Exception as exc:
            logging.warning("%s Excel 跟踪更新失败：%s", ticker, exc)
    return updated_paths


def write_word_report(path: Path, data_day: date, now_et: datetime, results: List[StockResult], scanned: int, success_count: int, news_searched: bool) -> None:
    if Document is None:
        raise RuntimeError("未安装 python-docx，请运行：pip install python-docx")
    doc = Document()
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.add_heading(f"强势金叉监控报告_{data_day.strftime('%Y-%m-%d')}", level=1)

    doc.add_heading("一、报告概况", level=2)
    for text in [
        f"数据对应的美股交易日：{data_day.strftime('%Y-%m-%d')}",
        f"报告生成时间（美东）：{now_et.strftime('%Y-%m-%d %H:%M:%S %Z')}",
        f"扫描股票总数：{scanned}",
        f"数据获取成功数量：{success_count}",
        f"评分达到{TRIGGER_SCORE}分及以上的触发股票数量：{len(results)}",
        f"是否执行新闻搜索：{'是' if news_searched else '否'}",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    if not results:
        doc.add_paragraph(f"本交易日没有股票达到{TRIGGER_SCORE}分触发阈值，因此未执行新闻搜索。")
        doc.add_heading("免责声明", level=2)
        doc.add_paragraph("本报告仅用于技术指标监控和公开信息整理，不构成任何投资建议。技术指标和新闻归因可能存在滞后或误差，请自行判断投资风险。")
        doc.save(path)
        return

    doc.add_heading("二、触发股票汇总表", level=2)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["股票代码", "最新价", "当日涨跌幅", "总评分", "信号等级", "成交量倍数", "主要原因"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for r in results:
        cells = table.add_row().cells
        values = [
            r.ticker,
            fmt_price(r.latest_close),
            f"{fmt_num(r.latest_change_pct, 2)}%",
            fmt_num(r.total_score, 1),
            r.signal_level,
            fmt_num(r.volume_ratio, 2),
            "、".join(r.reason_categories) if r.reason_categories else "技术性突破",
        ]
        for cell, value in zip(cells, values):
            cell.text = value

    doc.add_heading("三、触发股票简要分析", level=2)
    for r in results:
        doc.add_heading(f"{r.ticker}｜综合评分：{fmt_num(r.total_score, 1)}分｜{r.signal_level}", level=3)
        doc.add_paragraph(
            f"最新收盘价：{fmt_price(r.latest_close)}；当日涨跌幅：{fmt_num(r.latest_change_pct, 2)}%；成交量倍数：{fmt_num(r.volume_ratio, 2)}倍。"
            f"技术面：{r.technical_summary}"
            f"新闻原因：{r.reason_summary}"
            f"原因分类：{'、'.join(r.reason_categories)}。判断置信度：{r.confidence}。"
            f"风险提示：{'；'.join(r.risks)}"
        )
        if r.news:
            doc.add_paragraph("相关新闻：")
            for item in r.news[:3]:
                doc.add_paragraph(
                    f"{item.title}｜{item.published}｜{item.source}｜{item.summary_cn}｜{item.link}",
                    style="List Bullet",
                )
        else:
            doc.add_paragraph("相关新闻：未发现足以解释本次信号的明确重大新闻，更可能属于技术性反弹、板块联动或资金推动，判断置信度较低。")

    doc.add_heading("免责声明", level=2)
    doc.add_paragraph("本报告仅用于技术指标监控和公开信息整理，不构成任何投资建议。技术指标和新闻归因可能存在滞后或误差，请自行判断投资风险。")
    doc.save(path)


def run(force: bool = False) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    setup_logging(OUTPUT_DIR)
    now_et = eastern_now()
    data_day = latest_official_trading_day(now_et)
    out_path = OUTPUT_DIR / f"强势金叉监控报告_{data_day.strftime('%Y-%m-%d')}.docx"
    if out_path.exists() and not force:
        logging.info("报告已存在，未覆盖：%s。若需覆盖请使用 --force", out_path)
        update_cross_tracking_excels(OUTPUT_DIR, data_day)
        return out_path

    technical_results: List[StockResult] = []
    success_count = 0
    logging.info("开始扫描 %s 只股票，正式交易日：%s", len(TICKERS), data_day)
    for i, ticker in enumerate(TICKERS, 1):
        logging.info("[%s/%s] %s", i, len(TICKERS), ticker)
        try:
            result = analyze_ticker(ticker, data_day)
            if not result.data_error:
                success_count += 1
            technical_results.append(result)
            if result.total_score < TRIGGER_SCORE:
                logging.info("%s 综合评分为 %.1f，低于%s分，不触发且不搜索新闻。", ticker, result.total_score, TRIGGER_SCORE)
        except Exception as exc:
            logging.exception("%s 分析失败：%s", ticker, exc)
            technical_results.append(StockResult(ticker=ticker, company_name=get_company_name(ticker), data_error=str(exc)))
        time.sleep(SLEEP_BETWEEN_TICKERS)

    triggered = [r for r in technical_results if not r.data_error and r.total_score >= TRIGGER_SCORE]
    triggered.sort(key=lambda x: x.total_score, reverse=True)
    news_searched = bool(triggered)
    for r in triggered:
        logging.info("%s 达到 %.1f 分，开始搜索新闻。", r.ticker, r.total_score)
        try:
            r.news = search_news_for_result(r)
        except Exception as exc:
            logging.warning("%s 新闻搜索失败：%s", r.ticker, exc)
            r.news = []
        analyze_reason(r)
        time.sleep(0.5)

    write_word_report(out_path, data_day, now_et, triggered, len(TICKERS), success_count, news_searched)
    new_records = append_triggered_symbols_txt(OUTPUT_DIR, data_day, triggered)
    append_triggered_symbols_google_sheet(new_records)
    update_cross_tracking_excels(OUTPUT_DIR, data_day)
    logging.info("报告生成完成：%s", out_path)
    return out_path


def main() -> None:
    parser = argparse.ArgumentParser(description="美股强势 MACD/KDJ/RSI 金叉监控程序")
    parser.add_argument("--force", action="store_true", help="若同一交易日报告已存在，则强制覆盖")
    args = parser.parse_args()
    try:
        path = run(force=args.force)
        print("")
        print("报告生成完成：")
        print(f"Word: {path}")
    except KeyboardInterrupt:
        print("用户中断。")
    except Exception as exc:
        print("运行失败：", exc)
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()

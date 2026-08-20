from __future__ import annotations

import datetime as dt
import json
import logging
import math
import os
import time
import urllib.request
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from zoneinfo import ZoneInfo

import numpy as np
import pandas as pd
import yfinance as yf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


# =========================
# 基础配置
# =========================

TICKERS = [
    "OPTX",
    "ASTS",
    "CRWV",
    "FLY",
    "AAOI",
    "LUNR",
    "CRCL",
    "ONDS",
]

REPORT_BASE_DIR = Path(os.environ.get("REPORT_BASE_DIR", Path(__file__).resolve().parents[1] / "data"))
OUTPUT_DIR = REPORT_BASE_DIR / "drop"
REPORT_PREFIX = "美股技术风险监控"
TRIGGER_TXT_NAME = "drop.txt"

HISTORY_PERIOD = "18mo"
INTERVAL = "1d"
MIN_EXPECTED_ROWS = 300
REQUEST_PAUSE_SECONDS = 0.3
TRIGGER_SCORE = 4


@dataclass
class StockRiskResult:
    """保存单只股票的技术指标、风险评分和报告文字。"""

    ticker: str
    date: Optional[dt.date] = None
    close: Optional[float] = None
    macd_signal: str = "无法判断"
    kdj_signal: str = "无法判断"
    rsi_signal: str = "无法判断"
    ma_signal: str = "无法判断"
    volume_signal: str = "无法判断"
    rsi6: Optional[float] = None
    rsi12: Optional[float] = None
    rsi24: Optional[float] = None
    ma20: Optional[float] = None
    ma50: Optional[float] = None
    ma200: Optional[float] = None
    volume_ratio: Optional[float] = None
    deviation20: Optional[float] = None
    deviation50: Optional[float] = None
    uptrend_score: int = 0
    trend_regime: str = "无法判断"
    drop_risk_score: int = 0
    risk_level: str = "无法判断"
    current_status: str = "无法判断"
    risk_sources: List[str] = field(default_factory=list)
    observation: str = ""
    data_warning: str = ""
    error: str = ""


def setup_logging() -> None:
    """初始化日志输出，方便查看程序运行状态。"""

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    log_path = OUTPUT_DIR / "drop_monitor.log"
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[
            logging.FileHandler(log_path, encoding="utf-8"),
            logging.StreamHandler(),
        ],
    )


def fmt_num(value: Optional[float], digits: int = 2) -> str:
    """把数字格式化为报告可读文本。"""

    if value is None:
        return "-"
    try:
        if pd.isna(value) or math.isinf(float(value)):
            return "-"
        return f"{float(value):.{digits}f}"
    except Exception:
        return "-"


def fmt_price(value: Optional[float]) -> str:
    """把价格格式化为两位小数。"""

    return fmt_num(value, 2)


def safe_float(value) -> Optional[float]:
    """把 pandas/numpy 数值安全转换为 float。"""

    try:
        if value is None or pd.isna(value):
            return None
        return float(value)
    except Exception:
        return None


def flatten_yfinance_columns(df: pd.DataFrame) -> pd.DataFrame:
    """兼容 yfinance 单股票和多级列名返回格式。"""

    if isinstance(df.columns, pd.MultiIndex):
        # 单只股票下载时也可能返回 MultiIndex，优先取第一级 OHLCV 字段。
        if len(df.columns.levels) >= 2:
            try:
                df.columns = df.columns.get_level_values(0)
            except Exception:
                df.columns = [str(col[0]) for col in df.columns]
    return df


def eastern_now() -> dt.datetime:
    """获取当前美东时间，用于判断最近一个已收盘交易日。"""

    return dt.datetime.now(ZoneInfo("America/New_York"))


def previous_weekday(day: dt.date) -> dt.date:
    """返回给定日期之前最近的工作日，遇到周末自动回退。"""

    day = day - dt.timedelta(days=1)
    while day.weekday() >= 5:
        day = day - dt.timedelta(days=1)
    return day


def latest_closed_us_trading_day(now_et: Optional[dt.datetime] = None) -> dt.date:
    """估算最近一个已收盘美股交易日，节假日时由行情数据自动回退到上一个可用交易日。"""

    now_et = now_et or eastern_now()
    today = now_et.date()
    if today.weekday() >= 5:
        return previous_weekday(today + dt.timedelta(days=1))
    if now_et.time() < dt.time(17, 0):
        return previous_weekday(today)
    return today


def normalize_ohlcv(df: pd.DataFrame, target_day: dt.date) -> Tuple[pd.DataFrame, str]:
    """清洗OHLCV并限制到目标交易日及以前，确保使用最近已收盘数据。"""

    if df is None or df.empty:
        return pd.DataFrame(), "未返回数据"

    df = flatten_yfinance_columns(df.copy())
    required = ["Open", "High", "Low", "Close", "Volume"]
    missing = [col for col in required if col not in df.columns]
    if missing:
        return pd.DataFrame(), f"缺少字段：{', '.join(missing)}"

    df = df[required].copy()
    for col in required:
        df[col] = pd.to_numeric(df[col], errors="coerce")
    df = df.dropna(subset=["Open", "High", "Low", "Close", "Volume"])
    df = df.sort_index()
    df = df[df.index.date <= target_day]
    if df.empty:
        return pd.DataFrame(), f"没有 {target_day} 及以前的有效收盘数据"
    return df, ""


def fetch_stock_data_yahoo_chart(ticker: str, target_day: dt.date) -> Tuple[pd.DataFrame, str]:
    """yfinance限流时使用Yahoo Chart API备用获取日线数据。"""

    end_dt = dt.datetime.combine(target_day + dt.timedelta(days=2), dt.time(0, 0), tzinfo=ZoneInfo("America/New_York"))
    start_dt = end_dt - dt.timedelta(days=560)
    url = (
        f"https://query1.finance.yahoo.com/v8/finance/chart/{ticker}"
        f"?period1={int(start_dt.timestamp())}&period2={int(end_dt.timestamp())}"
        f"&interval=1d&events=history&includeAdjustedClose=true"
    )
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=20) as resp:
            payload = json.loads(resp.read().decode("utf-8"))
        result = payload["chart"]["result"][0]
        timestamps = result.get("timestamp") or []
        quote = result["indicators"]["quote"][0]
        rows = []
        for idx, stamp in enumerate(timestamps):
            row = {
                "Date": dt.datetime.fromtimestamp(stamp, ZoneInfo("America/New_York")).date(),
                "Open": quote.get("open", [None] * len(timestamps))[idx],
                "High": quote.get("high", [None] * len(timestamps))[idx],
                "Low": quote.get("low", [None] * len(timestamps))[idx],
                "Close": quote.get("close", [None] * len(timestamps))[idx],
                "Volume": quote.get("volume", [None] * len(timestamps))[idx],
            }
            rows.append(row)
        df = pd.DataFrame(rows)
        if df.empty:
            return pd.DataFrame(), "Yahoo Chart API未返回数据"
        df = df.drop_duplicates("Date").sort_values("Date")
        df.index = pd.to_datetime(df["Date"])
        return normalize_ohlcv(df, target_day)
    except Exception as exc:
        return pd.DataFrame(), f"Yahoo Chart API失败：{exc}"


def fetch_stock_data(ticker: str, target_day: dt.date) -> Tuple[pd.DataFrame, str]:
    """优先使用 yfinance 获取日线数据，失败或限流时使用Yahoo Chart API备用。"""

    try:
        df = yf.download(
            ticker,
            period=HISTORY_PERIOD,
            interval=INTERVAL,
            auto_adjust=False,
            progress=False,
            threads=False,
            timeout=15,
        )
    except Exception as exc:
        df = pd.DataFrame()
        yfinance_error = f"yfinance下载失败：{exc}"
    else:
        df, yfinance_error = normalize_ohlcv(df, target_day)
        if not df.empty:
            return df, ""

    fallback_df, fallback_error = fetch_stock_data_yahoo_chart(ticker, target_day)
    if not fallback_df.empty:
        return fallback_df, ""
    return pd.DataFrame(), f"{yfinance_error}；备用数据源：{fallback_error}"


def calculate_moving_averages(df: pd.DataFrame) -> pd.DataFrame:
    """计算 MA5、MA10、MA20、MA50、MA200 和均线偏离率。"""

    df = df.copy()
    for window in [5, 10, 20, 50, 200]:
        df[f"MA{window}"] = df["Close"].rolling(window).mean()
    df["Deviation20"] = (df["Close"] - df["MA20"]) / df["MA20"] * 100
    df["Deviation50"] = (df["Close"] - df["MA50"]) / df["MA50"] * 100
    return df


def calculate_macd(df: pd.DataFrame) -> pd.DataFrame:
    """计算 MACD：DIF、DEA 和 MACD柱。"""

    df = df.copy()
    ema12 = df["Close"].ewm(span=12, adjust=False).mean()
    ema26 = df["Close"].ewm(span=26, adjust=False).mean()
    df["DIF"] = ema12 - ema26
    df["DEA"] = df["DIF"].ewm(span=9, adjust=False).mean()
    df["MACD_Hist"] = df["DIF"] - df["DEA"]
    return df


def calculate_kdj(df: pd.DataFrame, n: int = 9) -> pd.DataFrame:
    """计算 KDJ 指标，初始 K 和 D 设为 50。"""

    df = df.copy()
    low_n = df["Low"].rolling(n).min()
    high_n = df["High"].rolling(n).max()
    denominator = (high_n - low_n).replace(0, np.nan)
    rsv = (df["Close"] - low_n) / denominator * 100
    rsv = rsv.fillna(50)

    k_values = []
    d_values = []
    k_prev = 50.0
    d_prev = 50.0
    for value in rsv:
        k_today = k_prev * 2 / 3 + float(value) / 3
        d_today = d_prev * 2 / 3 + k_today / 3
        k_values.append(k_today)
        d_values.append(d_today)
        k_prev = k_today
        d_prev = d_today

    df["K"] = k_values
    df["D"] = d_values
    df["J"] = 3 * df["K"] - 2 * df["D"]
    return df


def calculate_rsi(series: pd.Series, period: int) -> pd.Series:
    """使用 Wilder 平滑方法计算 RSI。"""

    delta = series.diff()
    gain = delta.clip(lower=0)
    loss = -delta.clip(upper=0)
    avg_gain = gain.ewm(alpha=1 / period, adjust=False, min_periods=period).mean()
    avg_loss = loss.ewm(alpha=1 / period, adjust=False, min_periods=period).mean()
    rs = avg_gain / avg_loss.replace(0, np.nan)
    rsi = 100 - (100 / (1 + rs))
    return rsi.fillna(50)


def calculate_rsi_set(df: pd.DataFrame) -> pd.DataFrame:
    """计算 RSI6、RSI12 和 RSI24。"""

    df = df.copy()
    df["RSI6"] = calculate_rsi(df["Close"], 6)
    df["RSI12"] = calculate_rsi(df["Close"], 12)
    df["RSI24"] = calculate_rsi(df["Close"], 24)
    return df


def calculate_volume_metrics(df: pd.DataFrame) -> pd.DataFrame:
    """计算成交量均线、成交量倍数和日涨跌幅。"""

    df = df.copy()
    df["Volume_MA20"] = df["Volume"].rolling(20).mean()
    df["Volume_Ratio"] = df["Volume"] / df["Volume_MA20"]
    df["Daily_Return_Pct"] = df["Close"].pct_change() * 100
    return df


def calculate_all_indicators(df: pd.DataFrame) -> pd.DataFrame:
    """集中计算所有技术指标，方便后续扩展 BOLL、资金流等模块。"""

    df = calculate_moving_averages(df)
    df = calculate_macd(df)
    df = calculate_kdj(df)
    df = calculate_rsi_set(df)
    df = calculate_volume_metrics(df)
    return df


def is_macd_hist_shrinking(df: pd.DataFrame) -> bool:
    """判断 MACD 柱是否连续三天缩短。"""

    hist = df["MACD_Hist"].dropna().tail(4)
    if len(hist) < 4:
        return False
    return bool((hist.diff().dropna().tail(3) < 0).all())


def detect_macd_bearish_divergence(df: pd.DataFrame, lookback: int = 60) -> bool:
    """判断过去20-60个交易日是否出现价格创新高但MACD未创新高。"""

    valid = df.dropna(subset=["Close", "DIF"]).tail(lookback)
    if len(valid) < 20:
        return False
    latest = valid.iloc[-1]
    previous = valid.iloc[:-1]
    price_new_high = latest["Close"] >= previous["Close"].max()
    macd_not_new_high = latest["DIF"] < previous["DIF"].max()
    return bool(price_new_high and macd_not_new_high)


def score_macd(df: pd.DataFrame, result: StockRiskResult) -> int:
    """根据 MACD 死叉、零轴、柱体缩短和顶背离计算风险分。"""

    score = 0
    latest = df.iloc[-1]
    prev = df.iloc[-2]
    macd_notes = []

    dead_cross = latest["DIF"] < latest["DEA"] and prev["DIF"] >= prev["DEA"]
    if dead_cross:
        score += 1
        macd_notes.append("MACD死叉")
        result.risk_sources.append("MACD死叉")

    if latest["DIF"] < 0 and latest["DIF"] < latest["DEA"]:
        score += 2
        macd_notes.append("DIF位于零轴下方")
        result.risk_sources.append("MACD处于零轴下方")

    if is_macd_hist_shrinking(df):
        score += 1
        macd_notes.append("MACD柱连续3天缩短")
        result.risk_sources.append("MACD动能衰减")

    if detect_macd_bearish_divergence(df):
        score += 2
        macd_notes.append("疑似顶背离")
        result.risk_sources.append("MACD顶背离")

    result.macd_signal = "；".join(macd_notes) if macd_notes else "MACD暂无明显转弱"
    return score


def score_kdj(df: pd.DataFrame, result: StockRiskResult) -> int:
    """根据 KDJ 死叉、过热和 J 线回落计算风险分。"""

    score = 0
    latest = df.iloc[-1]
    prev = df.iloc[-2]
    notes = []

    dead_cross = latest["K"] < latest["D"] and prev["K"] >= prev["D"]
    if dead_cross:
        score += 1
        notes.append("KDJ死叉")
        result.risk_sources.append("KDJ死叉")

    if dead_cross and latest["K"] > 80 and latest["D"] > 80:
        score += 2
        notes.append("KDJ高位死叉")
        result.risk_sources.append("KDJ高位死叉")

    if latest["J"] > 100:
        score += 2
        notes.append("J线大于100，短线过热")
        result.risk_sources.append("KDJ极端过热")

    if prev["J"] > 100 and latest["J"] <= 100:
        score += 1
        notes.append("J线从100以上回落")
        result.risk_sources.append("KDJ高位回落")

    result.kdj_signal = "；".join(notes) if notes else "KDJ暂无明显风险"
    return score


def score_rsi(df: pd.DataFrame, result: StockRiskResult) -> int:
    """根据 RSI 死叉、结构破坏、跌破关键位和过热计算风险分。"""

    score = 0
    latest = df.iloc[-1]
    prev = df.iloc[-2]
    notes = []

    rsi_dead_cross = latest["RSI6"] < latest["RSI12"] and prev["RSI6"] >= prev["RSI12"]
    if rsi_dead_cross:
        score += 1
        notes.append("RSI6跌破RSI12")
        result.risk_sources.append("RSI短线死叉")

    if latest["RSI6"] < latest["RSI12"] < latest["RSI24"]:
        score += 1
        notes.append("RSI空头排列")
        result.risk_sources.append("RSI多头结构破坏")

    if latest["RSI12"] < 40:
        score += 3
        notes.append("RSI12跌破40")
        result.risk_sources.append("RSI12跌破40")
    elif latest["RSI12"] < 50:
        score += 2
        notes.append("RSI12跌破50")
        result.risk_sources.append("RSI12跌破50")

    if latest["RSI12"] > 80:
        score += 2
        notes.append("RSI12大于80，明显过热")
        result.risk_sources.append("RSI明显过热")
    elif latest["RSI12"] > 70:
        score += 1
        notes.append("RSI12大于70，偏过热")
        result.risk_sources.append("RSI偏过热")

    result.rsi_signal = "；".join(notes) if notes else "RSI暂无明显风险"
    return score


def score_ma_trend(df: pd.DataFrame, result: StockRiskResult) -> int:
    """根据股价和均线位置判断短中长期趋势风险。"""

    score = 0
    latest = df.iloc[-1]
    notes = []

    close = latest["Close"]
    if pd.notna(latest["MA20"]) and close < latest["MA20"]:
        score += 1
        notes.append("跌破MA20")
        result.risk_sources.append("跌破20日均线")

    if pd.notna(latest["MA50"]) and close < latest["MA50"]:
        score += 2
        notes.append("跌破MA50")
        result.risk_sources.append("跌破50日均线")

    if pd.notna(latest["MA20"]) and pd.notna(latest["MA50"]) and latest["MA20"] < latest["MA50"]:
        score += 1
        notes.append("MA20低于MA50")
        result.risk_sources.append("均线趋势转弱")

    if pd.notna(latest["MA200"]) and close < latest["MA200"]:
        score += 2
        notes.append("跌破MA200")
        result.risk_sources.append("跌破200日均线")

    # 均线偏离主要用于过热识别，作为轻度风险提示。
    if pd.notna(latest["Deviation20"]) and latest["Deviation20"] >= 20:
        score += 1
        notes.append("偏离MA20超过20%")
        result.risk_sources.append("短期涨幅偏大")

    if pd.notna(latest["Deviation50"]) and latest["Deviation50"] >= 35:
        score += 1
        notes.append("偏离MA50超过35%")
        result.risk_sources.append("中期涨幅偏大")

    result.ma_signal = "；".join(notes) if notes else "均线结构暂无明显风险"
    return score


def score_volume(df: pd.DataFrame, result: StockRiskResult) -> int:
    """根据放量下跌条件计算成交量风险分。"""

    latest = df.iloc[-1]
    notes = []
    score = 0

    if pd.notna(latest["Daily_Return_Pct"]) and pd.notna(latest["Volume_Ratio"]):
        if latest["Daily_Return_Pct"] < -3 and latest["Volume_Ratio"] > 1.5:
            score += 2
            notes.append("放量下跌")
            result.risk_sources.append("放量下跌")

    if not notes and pd.notna(latest["Volume_Ratio"]):
        notes.append(f"成交量为20日均量的{latest['Volume_Ratio']:.2f}倍")

    result.volume_signal = "；".join(notes) if notes else "成交量暂无明显异常"
    return score


def is_slope_up(series: pd.Series, days: int = 10) -> bool:
    """判断均线最近一段时间是否向上。"""

    values = series.dropna().tail(days)
    if len(values) < days:
        return False
    return bool(values.iloc[-1] > values.iloc[0])


def calculate_uptrend_score(df: pd.DataFrame) -> Tuple[int, str, List[str]]:
    """先判断股票是否处于上升趋势，避免把已下跌股票误判为见顶风险。"""

    latest = df.iloc[-1]
    close = latest["Close"]
    score = 0
    reasons: List[str] = []

    if pd.notna(latest["MA20"]) and close > latest["MA20"]:
        score += 1
        reasons.append("收盘价站上MA20")
    if pd.notna(latest["MA50"]) and close > latest["MA50"]:
        score += 2
        reasons.append("收盘价站上MA50")
    if pd.notna(latest["MA20"]) and pd.notna(latest["MA50"]) and latest["MA20"] > latest["MA50"]:
        score += 2
        reasons.append("MA20高于MA50")
    if pd.notna(latest["MA200"]) and close > latest["MA200"]:
        score += 1
        reasons.append("收盘价站上MA200")
    if pd.notna(latest["MA50"]) and pd.notna(latest["MA200"]) and latest["MA50"] > latest["MA200"]:
        score += 1
        reasons.append("MA50高于MA200")
    if is_slope_up(df["MA20"], 10):
        score += 1
        reasons.append("MA20斜率向上")
    if is_slope_up(df["MA50"], 10):
        score += 1
        reasons.append("MA50斜率向上")

    close_20_ago = df["Close"].dropna().iloc[-21] if len(df["Close"].dropna()) >= 21 else np.nan
    close_60_ago = df["Close"].dropna().iloc[-61] if len(df["Close"].dropna()) >= 61 else np.nan
    if pd.notna(close_20_ago) and close > close_20_ago:
        score += 1
        reasons.append("近20日收益为正")
    if pd.notna(close_60_ago) and close > close_60_ago:
        score += 1
        reasons.append("近60日收益为正")

    required_structure = (
        pd.notna(latest["MA20"])
        and pd.notna(latest["MA50"])
        and close > latest["MA50"]
        and latest["MA20"] > latest["MA50"]
    )
    if score >= 6 and required_structure:
        regime = "上升趋势"
    elif score >= 4:
        regime = "趋势修复/震荡偏强"
    else:
        regime = "非上升趋势"
    return score, regime, reasons


def risk_level_from_score(score: int) -> str:
    """根据总分映射风险等级。"""

    if score >= 10:
        return "高风险，可能进入下降趋势"
    if score >= 7:
        return "明显下跌风险"
    if score >= 4:
        return "短线转弱"
    return "正常波动"


def judge_current_status(result: StockRiskResult) -> str:
    """综合风险等级和过热信号生成当前状态描述。"""

    overheat = any("过热" in item or "涨幅偏大" in item for item in result.risk_sources)
    if result.drop_risk_score >= 10:
        return "下跌风险"
    if result.drop_risk_score >= 7:
        return "趋势转弱"
    if result.drop_risk_score >= 4:
        return "短线调整风险"
    if overheat:
        return "高位过热"
    return "健康上涨或正常波动"


def build_observation_text(result: StockRiskResult) -> str:
    """生成关键支撑和后续观察位置说明。"""

    lines = [
        f"当前价格 {fmt_price(result.close)}；MA20 {fmt_price(result.ma20)}；MA50 {fmt_price(result.ma50)}；MA200 {fmt_price(result.ma200)}。",
    ]
    if result.ma50 is not None:
        lines.append("若后续有效跌破MA50，趋势进一步恶化的概率上升。")
    elif result.ma20 is not None:
        lines.append("当前数据不足以计算MA50，优先观察MA20支撑。")
    if result.ma200 is None:
        lines.append("MA200暂不可用，可能是上市时间较短或历史数据不足。")
    return "".join(lines)


def analyze_stock(ticker: str, target_day: dt.date) -> StockRiskResult:
    """下载并分析单只股票，失败时返回错误信息但不中断整体程序。"""

    result = StockRiskResult(ticker=ticker)
    df, err = fetch_stock_data(ticker, target_day)
    if df.empty:
        result.error = err
        result.risk_level = "数据获取失败"
        result.current_status = "无法判断"
        return result

    if len(df) < MIN_EXPECTED_ROWS:
        result.data_warning = f"历史数据不足{MIN_EXPECTED_ROWS}个交易日，当前仅{len(df)}个交易日；部分长期均线可能为空。"

    if len(df) < 60:
        result.error = result.data_warning or f"有效数据过短：{len(df)}个交易日"
        result.risk_level = "数据不足"
        result.current_status = "无法判断"
        return result

    df = calculate_all_indicators(df)
    latest = df.iloc[-1]
    result.date = latest.name.date() if hasattr(latest.name, "date") else None
    result.close = safe_float(latest["Close"])
    result.rsi6 = safe_float(latest["RSI6"])
    result.rsi12 = safe_float(latest["RSI12"])
    result.rsi24 = safe_float(latest["RSI24"])
    result.ma20 = safe_float(latest["MA20"])
    result.ma50 = safe_float(latest["MA50"])
    result.ma200 = safe_float(latest["MA200"])
    result.volume_ratio = safe_float(latest["Volume_Ratio"])
    result.deviation20 = safe_float(latest["Deviation20"])
    result.deviation50 = safe_float(latest["Deviation50"])
    result.uptrend_score, result.trend_regime, trend_reasons = calculate_uptrend_score(df)

    if result.trend_regime != "上升趋势":
        result.drop_risk_score = 0
        result.risk_level = "非上升趋势，跳过过热见顶评分"
        result.current_status = result.trend_regime
        result.ma_signal = "；".join(trend_reasons) if trend_reasons else "未满足上升趋势过滤条件"
        result.macd_signal = "未进入过热见顶评分"
        result.kdj_signal = "未进入过热见顶评分"
        result.rsi_signal = "未进入过热见顶评分"
        result.volume_signal = (
            f"成交量为20日均量的{latest['Volume_Ratio']:.2f}倍"
            if pd.notna(latest["Volume_Ratio"])
            else "成交量暂无明显异常"
        )
        result.risk_sources = [f"{result.trend_regime}：不按过热见顶模型触发"]
        result.observation = build_observation_text(result)
        return result

    score = 0
    score += score_macd(df, result)
    score += score_kdj(df, result)
    score += score_rsi(df, result)
    score += score_ma_trend(df, result)
    score += score_volume(df, result)

    result.drop_risk_score = int(score)
    result.risk_level = risk_level_from_score(result.drop_risk_score)
    result.current_status = judge_current_status(result)
    result.observation = build_observation_text(result)
    if not result.risk_sources:
        result.risk_sources.append("暂无明显技术风险")
    return result


def parse_trigger_symbol(symbol_text: str) -> str:
    """从 drop.txt 的 LUNR(2) 这类文本中提取股票代码。"""

    symbol = symbol_text.strip().upper()
    if "(" in symbol:
        symbol = symbol.split("(", 1)[0].strip()
    if "（" in symbol:
        symbol = symbol.split("（", 1)[0].strip()
    return symbol


def append_triggered_drop_txt(results: List[StockRiskResult], report_date: dt.date) -> Path:
    """把达到触发阈值的股票写入 drop.txt，格式：日期 股票代码(次数) 得分。"""

    txt_path = OUTPUT_DIR / TRIGGER_TXT_NAME
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    existing_keys = set()
    symbol_counts: Dict[str, int] = {}

    if txt_path.exists():
        try:
            for raw_line in txt_path.read_text(encoding="utf-8-sig").splitlines():
                parts = raw_line.strip().split()
                if len(parts) < 2:
                    continue
                symbol = parse_trigger_symbol(parts[1])
                if not symbol:
                    continue
                existing_keys.add(f"{parts[0]}\t{symbol}")
                symbol_counts[symbol] = symbol_counts.get(symbol, 0) + 1
        except Exception as exc:
            logging.warning("读取 drop.txt 失败，将仅按本次结果写入：%s", exc)
            existing_keys = set()
            symbol_counts = {}

    date_text = report_date.strftime("%Y-%m-%d")
    triggered = [
        r for r in results
        if not r.error and r.date == report_date and r.drop_risk_score >= TRIGGER_SCORE
    ]
    triggered.sort(key=lambda r: (-r.drop_risk_score, r.ticker))

    new_lines = []
    for result in triggered:
        key = f"{date_text}\t{result.ticker}"
        if key in existing_keys:
            continue
        next_count = symbol_counts.get(result.ticker, 0) + 1
        symbol_text = f"{result.ticker}({next_count})"
        new_lines.append(f"{date_text}\t{symbol_text}\t{result.drop_risk_score}分")
        existing_keys.add(key)
        symbol_counts[result.ticker] = next_count

    if new_lines:
        with txt_path.open("a", encoding="utf-8-sig") as file:
            for line in new_lines:
                file.write(line + "\n")
        logging.info("已写入 drop.txt：%s", "；".join(new_lines))
    else:
        logging.info("本交易日无新增触发记录，或 drop.txt 已存在对应记录。")
    return txt_path


def results_to_dataframe(results: List[StockRiskResult]) -> pd.DataFrame:
    """把分析结果转换为 CSV 所需表格。"""

    rows = []
    for r in results:
        rows.append(
            {
                "ticker": r.ticker,
                "date": r.date.isoformat() if r.date else "",
                "close": r.close,
                "MACD_signal": r.macd_signal,
                "KDJ_signal": r.kdj_signal,
                "RSI6": r.rsi6,
                "RSI12": r.rsi12,
                "RSI24": r.rsi24,
                "MA20": r.ma20,
                "MA50": r.ma50,
                "MA200": r.ma200,
                "volume_ratio": r.volume_ratio,
                "Uptrend_Score": r.uptrend_score,
                "Trend_Regime": r.trend_regime,
                "Drop_Risk_Score": r.drop_risk_score,
                "Risk_Level": r.risk_level,
                "Current_Status": r.current_status,
                "Risk_Sources": "；".join(r.risk_sources),
                "Data_Warning": r.data_warning,
                "Error": r.error,
            }
        )
    return pd.DataFrame(rows)


def get_report_date(results: List[StockRiskResult]) -> dt.date:
    """使用结果中最新交易日作为报告日期，若无数据则使用当天日期。"""

    dates = [r.date for r in results if r.date is not None]
    if dates:
        return max(dates)
    return dt.date.today()


def add_table_header(table, headers: List[str]) -> None:
    """给 Word 表格添加表头。"""

    hdr_cells = table.rows[0].cells
    for cell, text in zip(hdr_cells, headers):
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def write_word_report(results: List[StockRiskResult], report_date: dt.date) -> Path:
    """生成 Word 风险监控报告。"""

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / f"{REPORT_PREFIX}_{report_date.strftime('%Y-%m-%d')}.docx"
    doc = Document()

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_heading("美股技术风险监控报告", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(f"报告日期：{report_date.strftime('%Y-%m-%d')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sorted_results = sorted(results, key=lambda x: x.drop_risk_score, reverse=True)
    doc.add_heading("一、总体风险排名", level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    add_table_header(table, ["股票", "趋势状态", "风险评分", "风险等级", "当前价格", "主要风险"])
    for r in sorted_results:
        cells = table.add_row().cells
        cells[0].text = r.ticker
        cells[1].text = f"{r.trend_regime}({r.uptrend_score})"
        cells[2].text = str(r.drop_risk_score)
        cells[3].text = r.risk_level
        cells[4].text = fmt_price(r.close)
        cells[5].text = "；".join(r.risk_sources[:4])

    doc.add_paragraph("")
    doc.add_heading("二、逐股票分析", level=2)
    for r in sorted_results:
        doc.add_heading(r.ticker, level=3)
        if r.error:
            doc.add_paragraph(f"数据异常：{r.error}")
            continue

        if r.data_warning:
            doc.add_paragraph(f"数据提示：{r.data_warning}")

        doc.add_paragraph(f"当前价格：{fmt_price(r.close)}")
        doc.add_paragraph(f"趋势过滤：{r.trend_regime}；上升趋势评分：{r.uptrend_score}")
        doc.add_paragraph(f"当前状态：{r.current_status}")
        doc.add_paragraph(f"风险评分：总分 {r.drop_risk_score}；风险等级：{r.risk_level}")

        doc.add_paragraph("技术指标：")
        for text in [
            f"MACD状态：{r.macd_signal}",
            f"KDJ状态：{r.kdj_signal}",
            f"RSI：RSI6={fmt_num(r.rsi6, 1)}；RSI12={fmt_num(r.rsi12, 1)}；RSI24={fmt_num(r.rsi24, 1)}；{r.rsi_signal}",
            f"MA20/50/200：MA20={fmt_price(r.ma20)}；MA50={fmt_price(r.ma50)}；MA200={fmt_price(r.ma200)}；{r.ma_signal}",
            f"成交量变化：Volume/MA20={fmt_num(r.volume_ratio, 2)}倍；{r.volume_signal}",
        ]:
            doc.add_paragraph(text, style="List Bullet")

        doc.add_paragraph("主要风险来源：")
        for source in r.risk_sources:
            doc.add_paragraph(source, style="List Bullet")

        doc.add_paragraph(f"判断：{r.current_status}。{r.macd_signal}；{r.kdj_signal}；{r.rsi_signal}。")
        doc.add_paragraph(f"建议观察位置：{r.observation}")

    doc.add_paragraph("本报告仅用于技术指标监控和公开行情数据整理，不构成任何投资建议。")

    doc.save(path)
    return path


def cleanup_csv_files() -> None:
    """清理历史CSV文件，保持drop文件夹只保留Word、txt、脚本和日志。"""

    for csv_path in OUTPUT_DIR.glob("*.csv"):
        try:
            csv_path.unlink()
        except Exception as exc:
            logging.warning("删除CSV文件失败：%s，原因：%s", csv_path, exc)


def run_monitor() -> Tuple[Path, Path, List[StockRiskResult]]:
    """执行完整监控流程：下载数据、计算指标、评分、生成文件。"""

    setup_logging()
    results: List[StockRiskResult] = []
    target_day = latest_closed_us_trading_day()
    logging.info("开始美股技术风险监控，共 %s 只股票；目标交易日：%s。", len(TICKERS), target_day)
    for idx, ticker in enumerate(TICKERS, start=1):
        logging.info("[%s/%s] 分析 %s", idx, len(TICKERS), ticker)
        try:
            result = analyze_stock(ticker, target_day)
        except Exception as exc:
            logging.exception("%s 分析失败：%s", ticker, exc)
            result = StockRiskResult(ticker=ticker, error=str(exc), risk_level="分析失败")
        results.append(result)
        time.sleep(REQUEST_PAUSE_SECONDS)

    report_date = get_report_date(results)
    cleanup_csv_files()
    word_path = write_word_report(results, report_date)
    txt_path = append_triggered_drop_txt(results, report_date)
    logging.info("监控完成：Word=%s TXT=%s", word_path, txt_path)
    return word_path, txt_path, results


def main() -> None:
    """程序入口，供 Spyder、命令行和计划任务直接运行。"""

    word_path, txt_path, _ = run_monitor()
    print("")
    print("美股技术风险监控完成：")
    print(f"Word报告：{word_path}")
    print(f"触发记录：{txt_path}")


if __name__ == "__main__":
    main()
